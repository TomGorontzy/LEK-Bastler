import os
import sys
import tempfile
import unittest
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT


ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'src')
if SRC not in sys.path:
    sys.path.insert(0, SRC)

from word_processor import WordProcessor
from import_wizard import ImportSession, is_blocking_warning


class RegressionCoreTests(unittest.TestCase):
    def setUp(self):
        self.wp = WordProcessor()
        self.base_rules = deepcopy(self.wp.rules)

    def tearDown(self):
        self.wp.rules = deepcopy(self.base_rules)

    def _make_table_doc(self, path, rows):
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        for key, value in rows:
            row = table.add_row()
            row.cells[0].text = key
            row.cells[1].text = value
        doc.save(path)

    def _extract_first_task(self, path):
        tasks = self.wp.extract_tasks(path)
        self.assertTrue(tasks, 'Es wurde keine Aufgabe erkannt.')
        return tasks[0]

    def _paragraph_num_id(self, paragraph):
        nodes = paragraph._element.xpath('./w:pPr/w:numPr/w:numId')
        if not nodes:
            return None
        return nodes[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')

    def test_table_with_intro_hint_points_preview_order(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, 'table_full.docx')
            self._make_table_doc(path, [
                ('ID', 'A-001'),
                ('Titel', 'Komplettaufgabe'),
                ('Intro/Einleitung (optional)', 'Kontextblock vorhanden'),
                ('Aufgabenstellung (Pflicht)', 'Bitte bearbeiten'),
                ('Lösungsmöglichkeit/Hinweis (optional)', 'Achten Sie auf X'),
                ('Punkte', '10'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
            ])

            task = self._extract_first_task(path)
            lines = self.wp.build_task_flow_preview_lines(task)

            joined = '\n'.join(lines)
            expected_order = [
                'Titel:',
                'Intro/Einleitung:',
                'Aufgabenstellung:',
                'Hinweis:',
                'Punkte:',
            ]
            last_pos = -1
            for marker in expected_order:
                pos = joined.find(marker)
                self.assertGreater(pos, last_pos, f"Abschnittsreihenfolge fehlerhaft bei '{marker}'")
                last_pos = pos

    def test_table_without_optional_blocks_delta_check(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, 'table_minimal.docx')
            self._make_table_doc(path, [
                ('ID', 'A-002'),
                ('Titel', 'Minimalaufgabe'),
                ('Aufgabenstellung (Pflicht)', 'Nur Pflichtinhalt'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
            ])

            task = self._extract_first_task(path)
            delta = self.wp.analyze_task_flow_preview_delta(task)

            self.assertTrue(delta.get('is_structured'))
            missing = set(delta.get('missing_optional_sections') or [])
            self.assertTrue({'Intro/Einleitung', 'Hinweis', 'Punkte'}.issubset(missing))

    def test_missing_category_warning_and_block_rule(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, 'table_missing_category.docx')
            self._make_table_doc(path, [
                ('ID', 'A-003'),
                ('Titel', 'Ohne Kategorie'),
                ('Aufgabenstellung (Pflicht)', 'Inhalt'),
                ('Kategorie', ''),
                ('Schwierigkeitsgrad', 'mittel'),
            ])

            task = self._extract_first_task(path)
            diagnostics = self.wp._build_task_diagnostic(task)

            warnings = diagnostics.get('warnings') or []
            self.assertTrue(any('Kategorie fehlt' in w for w in warnings))
            self.assertTrue(bool(self.wp.get_import_rule('category_rules.block_export_on_missing', True)))

    def test_missing_title_fallback_and_required_warning(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, 'table_missing_title.docx')
            self._make_table_doc(path, [
                ('ID', 'A-004'),
                ('Aufgabenstellung (Pflicht)', 'Eine ausführliche Aufgabenstellung für Titelableitung'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
            ])

            self.wp.rules = deepcopy(self.base_rules)
            self.wp.rules.setdefault('template_rules', {})['required_fields'] = ['id', 'titel', 'aufgabenstellungpflicht']

            task = self._extract_first_task(path)
            self.assertNotEqual((task.get('title') or '').strip(), '')
            self.assertTrue(any('Pflichtfeld fehlt: Titel' in w for w in (task.get('pre_warnings') or [])))

    def test_inconsistent_difficulty_warning_and_block(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, 'table_inconsistent_diff.docx')
            self._make_table_doc(path, [
                ('ID', 'A-005'),
                ('Titel', 'Uneinheitliche Schwierigkeit'),
                ('Aufgabenstellung (Pflicht)', 'Inhalt'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'leicht und schwer'),
            ])

            task = self._extract_first_task(path)
            self.assertTrue(any('Inkonsistenter Schwierigkeitsgrad' in w for w in (task.get('pre_warnings') or [])))
            self.assertTrue(bool(self.wp.get_import_rule('difficulty_rules.block_export_on_inconsistent', True)))

    def test_partial_approval_exports_exact_approved(self):
        raw_tasks = [
            {'number': 1, 'number_display': '1', 'category': 'A', 'title': 'T1', 'content': ['c1'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 2, 'number_display': '2', 'category': 'A', 'title': 'T2', 'content': ['c2'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 3, 'number_display': '3', 'category': 'A', 'title': 'T3', 'content': ['c3'], 'difficulty': 'Mittel', 'keywords': []},
        ]

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=raw_tasks,
        )

        session.clear_approvals()
        session.set_task_approval(1, True)
        session.set_task_approval(3, True)

        approved = session.get_approved_raw_tasks()
        approved_numbers = [int(t.get('number', 0)) for t in approved]
        self.assertEqual(approved_numbers, [1, 3])

    def test_grouped_subtasks_are_approved_together(self):
        raw_tasks = [
            {'number': 1, 'number_display': '1.0', 'category': 'A', 'title': 'Intro', 'content': ['c1'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 2, 'number_display': '1.1', 'category': 'A', 'title': 'Teil 1', 'content': ['c2'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 3, 'number_display': '1.2', 'category': 'A', 'title': 'Teil 2', 'content': ['c3'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 4, 'number_display': '2.0', 'category': 'B', 'title': 'Andere Gruppe', 'content': ['c4'], 'difficulty': 'Mittel', 'keywords': []},
        ]

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=raw_tasks,
        )

        session.clear_approvals()
        effective_ids = session.set_task_approvals([2], True)

        self.assertEqual(effective_ids, [1, 2, 3])
        self.assertEqual(session.approved_task_ids, {1, 2, 3})

        approved = session.get_approved_raw_tasks()
        approved_numbers = [int(t.get('number', 0)) for t in approved]
        self.assertEqual(approved_numbers, [1, 2, 3])

    def test_grouped_subtasks_are_removed_together(self):
        raw_tasks = [
            {'number': 1, 'number_display': '1.0', 'category': 'A', 'title': 'Intro', 'content': ['c1'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 2, 'number_display': '1.1', 'category': 'A', 'title': 'Teil 1', 'content': ['c2'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 3, 'number_display': '1.2', 'category': 'A', 'title': 'Teil 2', 'content': ['c3'], 'difficulty': 'Mittel', 'keywords': []},
            {'number': 4, 'number_display': '2.0', 'category': 'B', 'title': 'Andere Gruppe', 'content': ['c4'], 'difficulty': 'Mittel', 'keywords': []},
        ]

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=raw_tasks,
        )

        session.approve_all()
        effective_ids = session.set_task_approvals([2], False)

        self.assertEqual(effective_ids, [1, 2, 3])
        self.assertEqual(session.approved_task_ids, {4})

        approved = session.get_approved_raw_tasks()
        approved_numbers = [int(t.get('number', 0)) for t in approved]
        self.assertEqual(approved_numbers, [4])

    def test_import_task_normalizes_scalar_list_fields(self):
        raw_task = {
            'number': 1,
            'number_display': '1',
            'category': 'A',
            'title': 'T1',
            'content': ['c1'],
            'difficulty': 'Mittel',
            'keywords': 'netzwerk',
            'intro': 'Einleitungstext',
            'warnings': 'Bitte prüfen',
        }

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=[raw_task],
        )

        self.assertEqual(len(session.tasks), 1)
        task = session.tasks[0]
        self.assertEqual(task.keywords, ['netzwerk'])
        self.assertEqual(task.intro, ['Einleitungstext'])
        self.assertEqual(task.warnings, ['Bitte prüfen'])

    def test_set_task_approval_ignores_unknown_id(self):
        raw_tasks = [
            {'number': 1, 'number_display': '1', 'category': 'A', 'title': 'T1', 'content': ['c1'], 'difficulty': 'Mittel', 'keywords': []},
        ]

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=raw_tasks,
        )

        self.assertEqual(session.approved_task_ids, {1})
        session.set_task_approval(999, True)
        self.assertEqual(session.approved_task_ids, {1})

    def test_stats_include_blocking_warning_count(self):
        raw_tasks = [
            {
                'number': 1,
                'number_display': '1',
                'category': 'A',
                'title': 'T1',
                'content': ['c1'],
                'difficulty': 'Mittel',
                'keywords': [],
                'warnings': ['Pflichtfeld fehlt: Titel'],
            },
            {
                'number': 2,
                'number_display': '2',
                'category': 'A',
                'title': 'T2',
                'content': ['c2'],
                'difficulty': 'Mittel',
                'keywords': [],
                'warnings': ['Keine Schlüsselwörter erkannt.'],
            },
        ]

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=raw_tasks,
        )

        stats = session.get_stats()
        self.assertEqual(stats.get('warnings'), 2)
        self.assertEqual(stats.get('blocking'), 1)

    def test_blocking_summary_groups_warning_types(self):
        raw_tasks = [
            {
                'number': 1,
                'number_display': '1',
                'category': 'A',
                'title': 'T1',
                'content': ['c1'],
                'difficulty': 'Mittel',
                'keywords': [],
                'warnings': ['Pflichtfeld fehlt: Titel', 'Kategorie fehlt (Pflichtfeld).'],
            },
            {
                'number': 2,
                'number_display': '2',
                'category': 'A',
                'title': 'T2',
                'content': ['c2'],
                'difficulty': 'Mittel',
                'keywords': [],
                'warnings': ['Inkonsistenter Schwierigkeitsgrad erkannt.'],
            },
        ]

        session = ImportSession.from_raw_tasks(
            source_file='dummy.docx',
            source_filename='dummy.docx',
            lek_theme='Demo',
            raw_tasks=raw_tasks,
        )

        summary = session.get_blocking_summary()
        self.assertEqual(summary.get('missing_title'), 1)
        self.assertEqual(summary.get('missing_category'), 1)
        self.assertEqual(summary.get('inconsistent_difficulty'), 1)
        self.assertEqual(summary.get('total'), 3)

    def test_is_blocking_warning_classification(self):
        self.assertTrue(is_blocking_warning('Pflichtfeld fehlt: Titel'))
        self.assertTrue(is_blocking_warning('Kategorie fehlt (Pflichtfeld).'))
        self.assertTrue(is_blocking_warning('Inkonsistenter Schwierigkeitsgrad erkannt.'))
        self.assertFalse(is_blocking_warning('Keine Schlüsselwörter erkannt.'))

    def test_structured_export_copies_nested_tables_from_task_cell(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_nested_table.docx')

            doc = Document()
            table = doc.add_table(rows=0, cols=2)
            for key, value in [
                ('ID', 'A-100'),
                ('Titel', 'Mit Untertabelle'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
                ('Aufgabenstellung (Pflicht)', ''),
            ]:
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = value

            task_cell = table.rows[-1].cells[1]
            task_cell.text = ''
            nested = task_cell.add_table(rows=1, cols=2)
            nested.cell(0, 0).text = 'Antwort'
            nested.cell(0, 1).text = 'Freitext'
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, [task])
            self.wp.append_task_content_for_lek(out_doc, task)

            self.assertGreaterEqual(len(out_doc.tables), 1, 'Verschachtelte Tabelle wurde nicht in die LEK übernommen.')

            all_table_texts = []
            for out_table in out_doc.tables:
                for row in out_table.rows:
                    for cell in row.cells:
                        text = str(cell.text or '').strip()
                        if text:
                            all_table_texts.append(text)

            joined = '\n'.join(all_table_texts)
            self.assertIn('Antwort', joined)
            self.assertIn('Freitext', joined)

    def test_export_resets_style_based_numbering_per_task(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_numbering_reset.docx')

            doc = Document()
            for task_id, title, first_item, second_item in [
                ('A-110', 'Liste Eins', 'Alpha 1', 'Alpha 2'),
                ('A-111', 'Liste Zwei', 'Beta 1', 'Beta 2'),
            ]:
                table = doc.add_table(rows=0, cols=2)
                for key, value in [
                    ('ID', task_id),
                    ('Titel', title),
                    ('Kategorie', 'Demo'),
                    ('Schwierigkeitsgrad', 'mittel'),
                    ('Aufgabenstellung (Pflicht)', ''),
                ]:
                    row = table.add_row()
                    row.cells[0].text = key
                    row.cells[1].text = value

                task_cell = table.rows[-1].cells[1]
                task_cell.text = ''
                first_para = task_cell.paragraphs[0]
                first_para.style = 'List Number'
                first_para.add_run(first_item)
                second_para = task_cell.add_paragraph(second_item)
                second_para.style = 'List Number'
                doc.add_paragraph('')

            doc.save(src_path)

            tasks = self.wp.extract_tasks(src_path)
            self.assertEqual(len(tasks), 2)

            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, tasks)
            for idx, task in enumerate(tasks):
                self.wp.append_task_to_lek_document(out_doc, task)
                if idx < len(tasks) - 1:
                    out_doc.add_paragraph('')

            para_by_text = {
                str(p.text or '').strip(): p
                for p in out_doc.paragraphs
                if str(p.text or '').strip() in {'Alpha 1', 'Alpha 2', 'Beta 1', 'Beta 2'}
            }

            alpha_1 = self._paragraph_num_id(para_by_text['Alpha 1'])
            alpha_2 = self._paragraph_num_id(para_by_text['Alpha 2'])
            beta_1 = self._paragraph_num_id(para_by_text['Beta 1'])
            beta_2 = self._paragraph_num_id(para_by_text['Beta 2'])

            self.assertIsNotNone(alpha_1)
            self.assertEqual(alpha_1, alpha_2)
            self.assertIsNotNone(beta_1)
            self.assertEqual(beta_1, beta_2)
            self.assertNotEqual(alpha_1, beta_1)

    def test_structured_export_materializes_numbering_in_nested_table_lists(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_nested_numbering.docx')

            doc = Document()
            table = doc.add_table(rows=0, cols=2)
            for key, value in [
                ('ID', 'A-112'),
                ('Titel', 'Untertabelle mit Liste'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
                ('Aufgabenstellung (Pflicht)', ''),
            ]:
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = value

            task_cell = table.rows[-1].cells[1]
            task_cell.text = ''
            nested = task_cell.add_table(rows=2, cols=1)
            nested_first = nested.cell(0, 0).paragraphs[0]
            nested_first.style = 'List Number'
            nested_first.add_run('Unterpunkt 1')
            nested_second = nested.cell(1, 0).paragraphs[0]
            nested_second.style = 'List Number'
            nested_second.add_run('Unterpunkt 2')
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, [task])
            self.wp.append_task_content_for_lek(out_doc, task)

            nested_paragraphs = []
            for out_table in out_doc.tables:
                for row in out_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = str(paragraph.text or '').strip()
                            if text in {'Unterpunkt 1', 'Unterpunkt 2'}:
                                nested_paragraphs.append(paragraph)

            self.assertEqual(len(nested_paragraphs), 2)
            first_num = self._paragraph_num_id(nested_paragraphs[0])
            second_num = self._paragraph_num_id(nested_paragraphs[1])
            self.assertIsNotNone(first_num)
            self.assertEqual(first_num, second_num)

    def test_structured_export_copies_sdt_controls_from_task_cell(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_sdt.docx')

            doc = Document()
            table = doc.add_table(rows=0, cols=2)
            for key, value in [
                ('ID', 'A-101'),
                ('Titel', 'Mit Content Control'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
                ('Aufgabenstellung (Pflicht)', 'Basistext'),
            ]:
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = value

            task_cell = table.rows[-1].cells[1]
            tc = task_cell._tc

            sdt = OxmlElement('w:sdt')
            sdt_pr = OxmlElement('w:sdtPr')
            sdt_content = OxmlElement('w:sdtContent')
            sdt_para = OxmlElement('w:p')
            sdt_run = OxmlElement('w:r')
            sdt_text = OxmlElement('w:t')
            sdt_text.text = '☒ Kontrollkästchen gesetzt'

            sdt_run.append(sdt_text)
            sdt_para.append(sdt_run)
            sdt_content.append(sdt_para)
            sdt.append(sdt_pr)
            sdt.append(sdt_content)
            tc.append(sdt)
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, [task])
            self.wp.append_task_content_for_lek(out_doc, task)

            sdt_nodes = out_doc._element.xpath('.//w:sdt')
            self.assertTrue(sdt_nodes, 'Inhaltssteuerelement (w:sdt) wurde nicht übernommen.')

            sdt_texts = [str(node.text or '') for node in out_doc._element.xpath('.//w:sdt//w:t')]
            self.assertIn('☒ Kontrollkästchen gesetzt', '\n'.join(sdt_texts))

    def test_export_writes_title_as_heading1_with_right_bordered_points(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_points.docx')
            self._make_table_doc(src_path, [
                ('ID', 'A-201'),
                ('Titel', 'Subnetzplanung Filiale Nord'),
                ('Aufgabenstellung (Pflicht)', 'Planen Sie das Subnetz.'),
                ('Punkte', '12'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
            ])

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp.append_task_to_lek_document(out_doc, task)

            self.assertTrue(out_doc.paragraphs, 'Es wurden keine Absätze erzeugt.')
            heading_para = out_doc.paragraphs[0]
            style_name = str(getattr(getattr(heading_para, 'style', None), 'name', '') or '').lower()
            self.assertTrue('heading 1' in style_name or 'überschrift 1' in style_name)

            heading_text = str(heading_para.text or '')
            self.assertIn('Subnetzplanung Filiale Nord', heading_text)
            self.assertIn('12', heading_text)
            self.assertNotIn('Punkte', heading_text)

            # Punkte-Label hat Innenabstand und Mindestbreite
            self.assertIn(' 12 ', heading_text)

            # Nach Heading 1 folgt grundsätzlich eine Leerzeile
            self.assertGreaterEqual(len(out_doc.paragraphs), 2)
            self.assertEqual(str(out_doc.paragraphs[1].text or ''), '')

            borders = heading_para._element.xpath('.//w:rPr/w:bdr')
            self.assertTrue(borders, 'Punkte-Run ohne Rahmen erkannt.')

    def test_export_does_not_duplicate_heading_title_in_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'heading_source.docx')

            doc = Document()
            doc.add_heading('Netzwerktechnik', level=1)
            doc.add_heading('Aufgabe 1 VLAN-Konzept erstellen', level=2)
            doc.add_paragraph('Beschreiben Sie ein VLAN-Konzept für zwei Standorte.')
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp.append_task_to_lek_document(out_doc, task)

            all_text = '\n'.join(p.text for p in out_doc.paragraphs if p.text)
            self.assertEqual(all_text.count('VLAN-Konzept erstellen'), 1)

    def test_points_box_uses_export_rule_width_and_padding(self):
        self.wp.rules = deepcopy(self.base_rules)
        self.wp.rules.setdefault('export_rules', {})
        self.wp.rules['export_rules'].setdefault('title_points_box', {})
        self.wp.rules['export_rules']['title_points_box']['min_inner_width'] = 16
        self.wp.rules['export_rules']['title_points_box']['padding_spaces'] = 2

        label = self.wp._format_points_label('12')
        self.assertTrue(label.startswith('  '))
        self.assertTrue(label.endswith('  '))

        stripped = label.strip()
        self.assertEqual(stripped, '12')
        self.assertGreaterEqual(len(label) - 4, 16)  # 2+2 padding spaces

    def test_intro_is_reported_as_hint_not_warning(self):
        task = {
            'title': 'Topologie planen',
            'content': [
                'Einleitung: Ausgangslage mit zwei Filialen',
                'Erstellen Sie ein Netzkonzept.',
            ],
            'keywords': ['netzwerk'],
            'difficulty': 'Mittel',
            'category': 'Demo',
            'pre_warnings': [],
        }

        diagnostic = self.wp._build_task_diagnostic(task)
        warnings = diagnostic.get('warnings') or []
        hints = diagnostic.get('hints') or []

        self.assertFalse(any('Einleitungs-/Kontexttext erkannt' in w for w in warnings))
        self.assertIn('Einleitung für Aufgabe und Folgeaufgaben erforderlich.', hints)

    def test_structured_export_omits_points_line_below_heading(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_points_omit_body.docx')
            self._make_table_doc(src_path, [
                ('ID', 'A-301'),
                ('Titel', 'IPv6 Routing'),
                ('Aufgabenstellung (Pflicht)', 'Konfigurieren Sie statisches Routing.'),
                ('Punkte', '15 Punkte'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
            ])

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp.append_task_to_lek_document(out_doc, task)

            # Heading enthält die Punktezahl, aber ohne das Wort "Punkte"
            heading_text = str(out_doc.paragraphs[0].text or '')
            self.assertIn('15', heading_text)
            self.assertNotIn('Punkte', heading_text)

            body_text = '\n'.join(str(p.text or '') for p in out_doc.paragraphs[1:])
            self.assertNotIn('Punkte:', body_text)

    def test_export_adds_blank_lines_and_sets_table_to_full_width(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'heading_with_table.docx')

            doc = Document()
            doc.add_heading('Netzwerktechnik', level=1)
            doc.add_heading('Aufgabe 1 Router-Topologie', level=2)
            doc.add_paragraph('Beschreiben Sie die Topologie.')
            tbl = doc.add_table(rows=1, cols=2)
            tbl.cell(0, 0).text = 'Spalte A'
            tbl.cell(0, 1).text = 'Spalte B'
            doc.add_paragraph('Nach der Tabelle folgt Text.')
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, [task])
            self.wp.append_task_to_lek_document(out_doc, task)

            # Leerzeile direkt nach Heading 1
            self.assertGreaterEqual(len(out_doc.paragraphs), 2)
            self.assertEqual(str(out_doc.paragraphs[1].text or ''), '')

            # Tabelle wurde übernommen und auf 100% Breite gesetzt
            self.assertGreaterEqual(len(out_doc.tables), 1)
            out_tbl = out_doc.tables[0]._tbl
            tbl_w_nodes = out_tbl.xpath('./w:tblPr/w:tblW')
            self.assertTrue(tbl_w_nodes, 'Keine Tabellenbreite (w:tblW) gesetzt.')
            self.assertEqual(tbl_w_nodes[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'), 'pct')
            self.assertEqual(tbl_w_nodes[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'), '5000')

            # Vor und nach Tabelle steht jeweils eine Leerzeile
            body = out_doc._element.body
            elems = list(body)
            tbl_index = next(i for i, e in enumerate(elems) if e.tag.endswith('}tbl'))

            self.assertGreater(tbl_index, 0)
            prev_elem = elems[tbl_index - 1]
            self.assertTrue(prev_elem.tag.endswith('}p'))
            prev_text = ''.join((n.text or '') for n in prev_elem.xpath('.//w:t')).strip()
            self.assertEqual(prev_text, '')

            self.assertLess(tbl_index + 1, len(elems))
            next_elem = elems[tbl_index + 1]
            self.assertTrue(next_elem.tag.endswith('}p'))
            next_text = ''.join((n.text or '') for n in next_elem.xpath('.//w:t')).strip()
            self.assertEqual(next_text, '')

    def test_structured_export_preserves_blank_paragraphs_in_task_cell(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'structured_blank_lines.docx')

            doc = Document()
            table = doc.add_table(rows=0, cols=2)
            for key, value in [
                ('ID', 'A-401'),
                ('Titel', 'Leerzeilen erhalten'),
                ('Kategorie', 'Demo'),
                ('Schwierigkeitsgrad', 'mittel'),
                ('Aufgabenstellung (Pflicht)', ''),
            ]:
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = value

            task_cell = table.rows[-1].cells[1]
            task_cell.text = ''
            task_cell.paragraphs[0].add_run('Erste Zeile')
            task_cell.add_paragraph('')
            task_cell.add_paragraph('Dritte Zeile')
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, [task])
            self.wp.append_task_to_lek_document(out_doc, task)

            texts = [str(p.text or '') for p in out_doc.paragraphs]
            self.assertIn('Erste Zeile', texts)
            self.assertIn('Dritte Zeile', texts)

            first_idx = texts.index('Erste Zeile')
            third_idx = texts.index('Dritte Zeile')
            self.assertGreater(third_idx, first_idx)
            self.assertTrue(any(t == '' for t in texts[first_idx + 1:third_idx]))

    def test_export_sets_keep_rules_on_task_heading_and_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            src_path = os.path.join(tmp, 'keep_rules.docx')

            doc = Document()
            doc.add_heading('Kategorie', level=1)
            doc.add_heading('Aufgabe 1 Zusammenhalt', level=2)
            doc.add_paragraph('Erster Inhaltsabsatz.')
            doc.add_paragraph('Zweiter Inhaltsabsatz.')
            doc.save(src_path)

            task = self._extract_first_task(src_path)
            out_doc = Document()
            self.wp._prepare_target_document_context(out_doc, [task])
            self.wp.append_task_to_lek_document(out_doc, task)

            self.assertTrue(out_doc.paragraphs, 'Keine Absätze im Exportdokument gefunden.')

            heading = out_doc.paragraphs[0]
            self.assertTrue(bool(heading.paragraph_format.keep_with_next))
            self.assertTrue(bool(heading.paragraph_format.keep_together))

            nonempty = [p for p in out_doc.paragraphs if str(p.text or '').strip()]
            self.assertGreaterEqual(len(nonempty), 2)
            self.assertTrue(any(bool(p.paragraph_format.keep_together) for p in nonempty[1:]))

    def test_external_table_reference_resolves_in_collection_subfolder(self):
        with tempfile.TemporaryDirectory() as tmp:
            aufgaben_dir = os.path.join(tmp, 'data', 'Aufgaben')
            os.makedirs(aufgaben_dir, exist_ok=True)

            collection_path = os.path.join(aufgaben_dir, 'Aufgaben_Auftragssteuerung und -koordination.docx')
            subfolder = os.path.join(aufgaben_dir, 'Auftragssteuerung und -koordination')
            os.makedirs(subfolder, exist_ok=True)
            external_path = os.path.join(subfolder, 'Kalkulationsschema.docx')

            doc = Document()
            doc.add_heading('Auftragssteuerung und -koordination', level=1)
            doc.add_heading('Aufgabe 1 Verkaufsgespräch', level=2)
            doc.add_paragraph('Beschreiben Sie das Vorgehen.')
            doc.add_paragraph('<<tabelle=Kalkulationsschema>>')
            doc.save(collection_path)

            ext_doc = Document()
            ext_table = ext_doc.add_table(rows=1, cols=2)
            ext_table.cell(0, 0).text = 'Spalte A'
            ext_table.cell(0, 1).text = 'Spalte B'
            ext_doc.save(external_path)

            tasks = self.wp.extract_tasks(collection_path)
            self.assertEqual(len(tasks), 1)
            task = tasks[0]

            self.assertEqual(task.get('external_table_reference'), 'Kalkulationsschema')
            self.assertEqual(os.path.normpath(task.get('external_table_path', '')), os.path.normpath(external_path))
            self.assertFalse(task.get('external_table_missing'))
            self.assertFalse(any('tabelle=' in line.lower() for line in (task.get('content') or [])))

    def test_external_table_reference_inserts_landscape_document(self):
        with tempfile.TemporaryDirectory() as tmp:
            aufgaben_dir = os.path.join(tmp, 'data', 'Aufgaben')
            os.makedirs(aufgaben_dir, exist_ok=True)

            collection_path = os.path.join(aufgaben_dir, 'Aufgaben_Auftragssteuerung und -koordination.docx')
            subfolder = os.path.join(aufgaben_dir, 'Auftragssteuerung und -koordination')
            os.makedirs(subfolder, exist_ok=True)
            external_path = os.path.join(subfolder, 'BreiteTabelle.docx')

            doc = Document()
            doc.add_heading('Auftragssteuerung und -koordination', level=1)
            doc.add_heading('Aufgabe 1 Kalkulation', level=2)
            doc.add_paragraph('Nutzen Sie die Tabelle für die Berechnung.')
            doc.add_paragraph('<<tabelle=BreiteTabelle>>')
            doc.save(collection_path)

            ext_doc = Document()
            ext_section = ext_doc.sections[0]
            ext_section.orientation = WD_ORIENT.LANDSCAPE
            ext_section.page_width, ext_section.page_height = ext_section.page_height, ext_section.page_width
            ext_table = ext_doc.add_table(rows=1, cols=6)
            for idx in range(6):
                ext_table.cell(0, idx).text = f'Spalte {idx + 1}'
            ext_doc.save(external_path)

            tasks = self.wp.extract_tasks(collection_path)
            task = tasks[0]

            out_doc = Document()
            self.wp.append_task_to_lek_document(out_doc, task)

            all_table_text = '\n'.join(cell.text for table in out_doc.tables for row in table.rows for cell in row.cells)
            self.assertIn('Spalte 1', all_table_text)
            self.assertTrue(any(section.orientation == WD_ORIENT.LANDSCAPE for section in out_doc.sections))


if __name__ == '__main__':
    unittest.main()