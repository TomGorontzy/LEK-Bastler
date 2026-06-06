"""
Template Manager - Verwaltung von LEK-Vorlagen
"""

import os
import re
from docx import Document
from docx.shared import Pt
import glob

class TemplateManager:
    """Klasse für die Verwaltung und Auswahl von LEK-Vorlagen"""
    
    def __init__(self, template_folder="Vorlagen"):
        self.template_folder = template_folder
    
    def find_matching_template(self, lek_theme):
        """
        Findet eine passende Vorlage basierend auf dem LEK-Thema
        
        Args:
            lek_theme (str): Das extrahierte LEK-Thema
            
        Returns:
            str: Pfad zur passenden Vorlage oder None
        """
        if not os.path.exists(self.template_folder):
            return None

        # Feste Basisvorlage bevorzugen
        fixed_template = os.path.join(self.template_folder, "LEK-Vorlage.docx")
        if os.path.exists(fixed_template):
            return fixed_template
        
        # Suche nach .docx-Dateien im Vorlagen-Ordner
        template_files = glob.glob(os.path.join(self.template_folder, "*.docx"))
        
        if not lek_theme:
            # Wenn kein Thema erkannt, nimm die erste verfügbare Vorlage
            return template_files[0] if template_files else None
        
        # Suche nach Vorlagen, die das Thema im Namen enthalten
        theme_words = re.split(r'[-_\s]+', lek_theme.lower())
        
        best_match = None
        best_score = 0
        
        for template_path in template_files:
            template_name = os.path.basename(template_path).lower()
            
            # Zähle übereinstimmende Wörter
            score = 0
            for word in theme_words:
                if word in template_name:
                    score += 1
            
            # Spezielle Behandlung für bekannte Begriffe
            if any(keyword in template_name for keyword in ['auftragssteuerung', 'koordination']):
                if any(keyword in lek_theme.lower() for keyword in ['auftragssteuerung', 'koordination']):
                    score += 10
            
            if score > best_score:
                best_score = score
                best_match = template_path
        
        return best_match if best_score > 0 else (template_files[0] if template_files else None)
    
    def get_available_templates(self):
        """
        Gibt eine Liste aller verfügbaren Vorlagen zurück
        
        Returns:
            list: Liste von Tupeln (Dateiname, Vollpfad)
        """
        if not os.path.exists(self.template_folder):
            return []
        
        template_files = glob.glob(os.path.join(self.template_folder, "*.docx"))
        return [(os.path.basename(f), f) for f in template_files]
    
    def load_template_as_base(self, template_path):
        """
        Lädt eine Vorlage als Basis-Dokument und ersetzt Platzhalter
        
        Args:
            template_path (str): Pfad zur Vorlage
            
        Returns:
            Document: Geladenes Word-Dokument
        """
        try:
            doc = Document(template_path)
            return doc
        except Exception as e:
            raise Exception(f"Fehler beim Laden der Vorlage: {str(e)}")
    
    def replace_template_placeholders(self, doc, lek_theme=""):
        """
        Ersetzt Platzhalter-Texte in der Vorlage durch tatsächliche Werte
        
        Args:
            doc: Word-Dokument
            lek_theme (str): LEK-Thema zum Ersetzen von 'TitelFürThema'
        """
        if not lek_theme:
            lek_theme = "Aufgabensammlung"  # Fallback
        
        # Durchsuche alle Paragraphen nach Platzhaltern
        for paragraph in doc.paragraphs:
            if 'TitelFürThema' in paragraph.text:
                # Ersetze den Platzhalter im Text
                self._replace_text_in_paragraph(paragraph, 'TitelFürThema', lek_theme)
        
        # Durchsuche auch Tabellen nach Platzhaltern
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'TitelFürThema' in paragraph.text:
                            self._replace_text_in_paragraph(paragraph, 'TitelFürThema', lek_theme)
        
        # Durchsuche Header und Footer
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if 'TitelFürThema' in paragraph.text:
                        self._replace_text_in_paragraph(paragraph, 'TitelFürThema', lek_theme)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if 'TitelFürThema' in paragraph.text:
                        self._replace_text_in_paragraph(paragraph, 'TitelFürThema', lek_theme)
    
    def _replace_text_in_paragraph(self, paragraph, old_text, new_text):
        """
        Ersetzt Text in einem Paragraph unter Beibehaltung der Formatierung
        
        Args:
            paragraph: Word-Paragraph-Objekt
            old_text (str): Zu ersetzender Text
            new_text (str): Neuer Text
        """
        # Einfache Textersetzung - behält Grundformatierung bei
        if old_text in paragraph.text:
            # Sammle alle Runs und deren Text
            full_text = paragraph.text
            
            if old_text in full_text:
                # Ersetze den Text
                new_full_text = full_text.replace(old_text, new_text)
                
                # Einfache Lösung: Ersetze den ersten Run mit dem neuen Text
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    first_run.text = new_full_text
                    
                    # Entferne alle anderen Runs
                    for i in range(len(paragraph.runs) - 1, 0, -1):
                        paragraph._element.remove(paragraph.runs[i]._element)
                else:
                    # Falls keine Runs vorhanden, erstelle einen neuen
                    paragraph.add_run(new_full_text)
    
    def _copy_paragraphs_with_formatting(self, target_doc, source_paragraphs):
        """
        Kopiert Paragraphen mit vollständiger Formatierung aus dem Quell- ins Zieldokument
        """
        for source_para in source_paragraphs:
            new_para = target_doc.add_paragraph()
            
            # Paragraph-Stil übernehmen
            if source_para.style:
                try:
                    new_para.style = source_para.style
                except:
                    # Falls Stil nicht verfügbar, verwende Standard
                    pass
            
            # Paragraph-Eigenschaften übernehmen
            para_format = new_para.paragraph_format
            source_format = source_para.paragraph_format
            
            if source_format.alignment:
                para_format.alignment = source_format.alignment
            if source_format.left_indent:
                para_format.left_indent = source_format.left_indent
            if source_format.right_indent:
                para_format.right_indent = source_format.right_indent
            if source_format.first_line_indent:
                para_format.first_line_indent = source_format.first_line_indent
            if source_format.space_before:
                para_format.space_before = source_format.space_before
            if source_format.space_after:
                para_format.space_after = source_format.space_after
            
            # Runs (Text-Segmente) kopieren
            for source_run in source_para.runs:
                new_run = new_para.add_run(source_run.text)
                
                # Font-Eigenschaften übernehmen
                if source_run.font.name:
                    new_run.font.name = source_run.font.name
                if source_run.font.size:
                    new_run.font.size = source_run.font.size
                if source_run.font.color.rgb:
                    new_run.font.color.rgb = source_run.font.color.rgb
                
                # Text-Formatierung übernehmen
                new_run.bold = source_run.bold
                new_run.italic = source_run.italic
                new_run.underline = source_run.underline
                
                # Erweiterte Formatierung
                if source_run.font.subscript:
                    new_run.font.subscript = source_run.font.subscript
                if source_run.font.superscript:
                    new_run.font.superscript = source_run.font.superscript
                if source_run.font.strike:
                    new_run.font.strike = source_run.font.strike
                if source_run.font.small_caps:
                    new_run.font.small_caps = source_run.font.small_caps
                if source_run.font.all_caps:
                    new_run.font.all_caps = source_run.font.all_caps
    
    def _copy_elements_with_formatting(self, target_doc, source_elements):
        """
        Kopiert alle Arten von Word-Elementen (Paragraphen, Tabellen, etc.) mit Formatierung
        
        Args:
            target_doc: Ziel-Word-Dokument
            source_elements: Liste von XML-Elementen
        """
        from copy import deepcopy
        
        for element in source_elements:
            try:
                if element.tag.endswith('p'):  # Paragraph
                    # Element direkt kopieren
                    new_element = deepcopy(element)
                    target_doc._element.body.append(new_element)
                    
                elif element.tag.endswith('tbl'):  # Tabelle
                    # Tabelle direkt kopieren
                    new_table_element = deepcopy(element)
                    target_doc._element.body.append(new_table_element)
                
                else:
                    # Andere Strukturen (Listen etc.)
                    new_element = deepcopy(element)
                    target_doc._element.body.append(new_element)
                    
            except Exception as e:
                # Fallback: Fehlermeldung als Paragraph
                target_doc.add_paragraph(f"[Struktur konnte nicht kopiert werden: {type(element).__name__}]")
    
    def insert_tasks_from_page_2(self, doc, tasks, lek_theme=""):
        """
        Fügt Aufgaben ab der zweiten Seite in das Dokument ein.
        
        Args:
            doc: Word-Dokument (Vorlage)
            tasks: Liste der Aufgaben
            lek_theme: LEK-Thema
        """
        # Seitenumbruch: Aufgaben starten auf Seite 2
        doc.add_page_break()
        
        # Seitennummerierung ab Seite 2 einrichten (beginnt mit 1)
        self._add_page_numbering_from_page_2(doc)

        from word_processor import WordProcessor
        wp = WordProcessor()
        
        # Aufgaben hinzufügen
        for i, task in enumerate(tasks, 1):
            # Zentrale Exportlogik: Heading1 + Punkte, strukturierte Tabellen als Fließtext
            wp.append_task_to_lek_document(doc, task)
            
            # Metadaten (optional, auskommentiert für sauberes Layout)
            # if task.get('difficulty') or task.get('keywords'):
            #     meta_info = []
            #     if task.get('difficulty'):
            #         meta_info.append(f"Schwierigkeit: {task['difficulty']}")
            #     if task.get('keywords'):
            #         meta_info.append(f"Schlüsselwörter: {', '.join(task['keywords'][:3])}")
            #     
            #     meta_paragraph = doc.add_paragraph()
            #     meta_run = meta_paragraph.add_run(' | '.join(meta_info))
            #     meta_run.italic = True
            #     meta_run.font.size = Pt(9)
            
            # Abstand zwischen Aufgaben
            if i < len(tasks):
                doc.add_paragraph()  # Leerer Absatz als Trenner
        
        # Dokument speichern nicht hier, wird von der aufrufenden Funktion gemacht
        
    def _add_page_numbering_from_page_2(self, doc):
        """
        Fügt Seitennummerierung ab Seite 2 hinzu (beginnend mit 1)
        Format: "Seite/Gesamtseiten" rechtsbündig in der Fußzeile
        
        Args:
            doc: Word-Dokument
        """
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        
        # Aktuelle Sektion konfigurieren
        current_section = doc.sections[-1]
        
        # Fußzeile für die aktuelle Sektion konfigurieren
        footer = current_section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Seitennummerierung hinzufügen (Feld-Code für PAGE und NUMPAGES)
        run = footer_paragraph.add_run()
        
        # XML für Seitennummerierung erstellen
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE \\* MERGEFORMAT'
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # " / " Trenner
        footer_paragraph.add_run(' / ')
        
        # Gesamtseitenzahl
        run2 = footer_paragraph.add_run()
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES \\* MERGEFORMAT'
        run2._r.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar4)
        
        # Formatierung der Fußzeile
        for run in footer_paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Aptos'

    # Rückwärtskompatibilität
    def insert_tasks_from_page_3(self, doc, tasks, lek_theme=""):
        self.insert_tasks_from_page_2(doc, tasks, lek_theme)

    # Rückwärtskompatibilität
    def _add_page_numbering_from_page_3(self, doc):
        self._add_page_numbering_from_page_2(doc)