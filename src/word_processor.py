"""
WordProcessor - Klasse für die Verarbeitung von Word-Dokumenten
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from template_manager import TemplateManager
from difflib import SequenceMatcher
import re
import os
import sys
import json
import shutil
from pathlib import Path
from copy import deepcopy
from datetime import datetime

class WordProcessor:
    """Klasse für das Lesen und Schreiben von Word-Dokumenten"""
    
    def __init__(self):
        if getattr(sys, 'frozen', False):
            runtime_base = Path(sys.executable).resolve().parent
        else:
            runtime_base = Path(__file__).resolve().parents[1]

        template_folder = runtime_base / 'data' / 'Vorlagen'
        self.template_manager = TemplateManager(str(template_folder))
        self.rules = self._load_import_rules(runtime_base)
        self._num_id_map = {}
        self._last_context_report = None

    def _default_import_rules(self):
        """Liefert Standardregeln für Import/Diagnostik."""
        return {
            'duplicate_rules': {
                'mode': 'normal',
                'similarity_thresholds': {
                    'strict': 0.82,
                    'normal': 0.70,
                    'relaxed': 0.58,
                },
                'fallback_similarity_threshold': 0.70,
            },
            'duplicate_mode': 'normal',
            'duplicate_similarity_thresholds': {
                'strict': 0.82,
                'normal': 0.70,
                'relaxed': 0.58,
            },
            'duplicate_similarity_threshold': 0.70,
            'max_preview_blocks': 10,
            'bulk_max_errors': 5,
            'preview_rules': {
                'show_optional_missing_sections': True,
                'optional_sections': ['Intro/Einleitung', 'Hinweis', 'Punkte'],
                'task_flow_sections': [
                    {
                        'name': 'title',
                        'label': 'Titel',
                        'fields': ['titel'],
                        'optional': False,
                    },
                    {
                        'name': 'intro',
                        'label': 'Intro/Einleitung',
                        'fields': ['introeinleitungoptional'],
                        'optional': True,
                    },
                    {
                        'name': 'task',
                        'label': 'Aufgabenstellung',
                        'fields': ['aufgabenstellungpflicht'],
                        'optional': False,
                    },
                    {
                        'name': 'hint',
                        'label': 'Hinweis',
                        'fields': ['loesungsmoeglichkeithinweisoptional'],
                        'optional': True,
                    },
                    {
                        'name': 'points',
                        'label': 'Punkte',
                        'fields': ['punkte'],
                        'optional': True,
                    },
                ],
            },
            'field_alias_rules': {
                'labels': {
                    'id': 'ID',
                    'titel': 'Titel',
                    'kategorie': 'Kategorie',
                    'introeinleitungoptional': 'Intro/Einleitung (optional)',
                    'aufgabenstellungpflicht': 'Aufgabenstellung (Pflicht)',
                    'loesungsmoeglichkeithinweisoptional': 'Lösungsmöglichkeit/Hinweis (optional)',
                    'schlagwortekommagetrennt': 'Schlagworte (kommagetrennt)',
                    'schwierigkeitsgrad': 'Schwierigkeitsgrad',
                    'punkte': 'Punkte',
                },
                'structured_task_fields': {
                    'id': ['id'],
                    'titel': ['titel'],
                    'kategorie': ['kategorie'],
                    'introeinleitungoptional': ['introeinleitungoptional', 'introeinleitung', 'einleitung', 'intro', 'kontext'],
                    'aufgabenstellungpflicht': ['aufgabenstellungpflicht', 'aufgabenstellung'],
                    'loesungsmoeglichkeithinweisoptional': ['loesungsmoeglichkeithinweisoptional', 'loesungsmoeglichkeithinweis', 'hinweis'],
                    'schlagwortekommagetrennt': ['schlagwortekommagetrennt', 'schlagworte', 'schluesselwoerter', 'schlusselwoerter'],
                    'schwierigkeitsgrad': ['schwierigkeitsgrad', 'schwierigkeit'],
                    'punkte': ['punkte', 'punktzahl', 'bewertung', 'bewertungpunkte'],
                },
            },
            'collection_governance': {
                'recommended_heading_structure': {
                    'level1': 'Kategorie',
                    'level2': 'Aufgabe',
                },
                'notes': [
                    'Neue Sammlungen bevorzugt als strukturierte 2-Spalten-Tabellen pflegen.',
                    'Pflichtfelder über template_rules.required_fields steuern.',
                    'Alias-Erweiterungen ausschließlich in field_alias_rules pflegen.',
                ],
            },
            'parser_rules': {
                'mode': 'auto',
                'prefer_mode_on_mixed': 'headings',
                'fallback_to_secondary_on_empty': True,
                'include_secondary_in_mixed': False,
            },
            'export_rules': {
                'title_points_box': {
                    'min_inner_width': 12,
                    'padding_spaces': 1,
                },
            },
            'external_table_rules': {
                'marker_name': 'tabelle',
                'orientation_marker_name': 'tabelle_format',
                'default_extension': '.docx',
                'search_subfolder_from_collection_name': True,
                'block_export_on_missing_reference': True,
            },
            'difficulty_rules': {
                'allowed_values': ['leicht', 'mittel', 'schwer'],
                'aliases': {
                    'easy': 'leicht',
                    'einfach': 'leicht',
                    'medium': 'mittel',
                    'normal': 'mittel',
                    'hard': 'schwer',
                    'difficult': 'schwer',
                    'komplex': 'schwer',
                },
                'block_export_on_inconsistent': True,
            },
            'template_rules': {
                'required_fields': ['id', 'aufgabenstellungpflicht', 'titel'],
                'block_export_on_missing_required': True,
            },
            'default_import_metadata': {
                'category': '',
                'difficulty': 'mittel',
                'keywords': '',
                'title': '',
            },
            'category_rules': {
                'required': True,
                'block_export_on_missing': True,
                'missing_values': ['', 'ohne kategorie', '-', 'n/a', 'keine'],
            },
        }

    def _deep_merge_dict(self, base, override):
        """Merged zwei Dicts rekursiv (override gewinnt)."""
        result = dict(base)
        for key, value in (override or {}).items():
            if isinstance(value, dict) and isinstance(result.get(key), dict):
                result[key] = self._deep_merge_dict(result[key], value)
            else:
                result[key] = value
        return result

    def _load_import_rules(self, runtime_base):
        """Lädt optionale Regelkonfiguration aus data/config/import_rules.json."""
        defaults = self._default_import_rules()
        cfg_path = Path(runtime_base) / 'data' / 'config' / 'import_rules.json'

        if not cfg_path.exists():
            return defaults

        try:
            with cfg_path.open('r', encoding='utf-8') as f:
                raw = json.load(f)
            if not isinstance(raw, dict):
                return defaults
            return self._deep_merge_dict(defaults, raw)
        except Exception:
            return defaults

    def _rule_value(self, key, default=None):
        """Liefert einen Regelwert aus self.rules mit Fallback."""
        return self.rules.get(key, default)

    def get_import_rule(self, key, default=None):
        """Öffentliche Regelabfrage mit optionaler Punktnotation (z. B. a.b.c)."""
        if not key:
            return default

        if '.' not in key:
            return self.rules.get(key, default)

        current = self.rules
        for part in key.split('.'):
            if not isinstance(current, dict) or part not in current:
                return default
            current = current[part]
        return current

    def _resolve_duplicate_threshold(self):
        """Ermittelt den aktiven Duplikat-Schwellwert anhand des konfigurierten Modus."""
        mode = str(
            self.get_import_rule('duplicate_mode', self.get_import_rule('duplicate_rules.mode', 'normal')) or 'normal'
        ).lower()
        thresholds = self.get_import_rule('duplicate_similarity_thresholds', None)
        if not isinstance(thresholds, dict) or not thresholds:
            thresholds = self.get_import_rule('duplicate_rules.similarity_thresholds', {}) or {}

        if mode in thresholds:
            try:
                return float(thresholds[mode])
            except (TypeError, ValueError):
                pass

        # Fallback: direkter Schwellwert
        try:
            fallback_threshold = self.get_import_rule(
                'duplicate_similarity_threshold',
                self.get_import_rule('duplicate_rules.fallback_similarity_threshold', 0.70)
            )
            return float(fallback_threshold or 0.70)
        except (TypeError, ValueError):
            return 0.70

    def _is_missing_category(self, value):
        """Prüft, ob ein Kategorienwert als fehlend gilt (regelbasiert)."""
        raw = str(value or '').strip().lower()
        if not raw:
            return True

        configured = self.get_import_rule('category_rules.missing_values', []) or []
        normalized_config = {str(v).strip().lower() for v in configured}
        return raw in normalized_config

    def _is_missing_required_value(self, value):
        """Prüft, ob ein Pflichtfeldwert als fehlend gilt."""
        return str(value or '').strip() == ''

    def _field_labels(self):
        """Liefert lesbare Feldlabels für strukturierte Aufgabenfelder."""
        return self.get_import_rule('field_alias_rules.labels', {}) or {}

    def _field_alias_map(self):
        """Liefert Alias-Definitionen je kanonischem Feldschlüssel."""
        return self.get_import_rule('field_alias_rules.structured_task_fields', {}) or {}

    def _field_aliases(self, canonical_key, fallback_aliases=None):
        """Liefert normalisierte Aliasliste für ein Feld inkl. kanonischem Schlüssel."""
        aliases_map = self._field_alias_map()
        configured = aliases_map.get(canonical_key, []) if isinstance(aliases_map, dict) else []
        raw_aliases = list(configured or [])
        if fallback_aliases:
            raw_aliases.extend(list(fallback_aliases))
        raw_aliases.append(canonical_key)

        normalized = []
        seen = set()
        for alias in raw_aliases:
            key = self._normalize_table_key(alias)
            if key and key not in seen:
                seen.add(key)
                normalized.append(key)
        return normalized

    def _value_by_alias(self, values_by_key, canonical_key, fallback_aliases=None):
        """Liefert den ersten nicht-leeren Feldwert über Aliasauflösung."""
        for alias in self._field_aliases(canonical_key, fallback_aliases=fallback_aliases):
            value = values_by_key.get(alias)
            if str(value or '').strip():
                return str(value).strip()
        return ''

    def _task_flow_sections(self):
        """Liefert die konfigurierten Bereichsdefinitionen für Vorschau/Export."""
        configured = self.get_import_rule('preview_rules.task_flow_sections', []) or []
        sections = []
        for section in configured:
            if not isinstance(section, dict):
                continue
            label = str(section.get('label') or '').strip()
            fields = [str(field).strip() for field in (section.get('fields') or []) if str(field).strip()]
            if not label or not fields:
                continue
            sections.append({
                'name': str(section.get('name') or '').strip().lower(),
                'label': label,
                'fields': fields,
                'optional': bool(section.get('optional', False)),
            })
        return sections

    def persist_import_rules(self, runtime_base=None):
        """Speichert die aktuell aktiven Importregeln dauerhaft nach data/config/import_rules.json.

        Args:
            runtime_base (Path|str|None): Basisverzeichnis der App. Bei None wird es automatisch ermittelt.

        Returns:
            str: Vollständiger Pfad der geschriebenen Konfigurationsdatei.
        """
        if runtime_base is None:
            if getattr(sys, 'frozen', False):
                runtime_base = Path(sys.executable).resolve().parent
            else:
                runtime_base = Path(__file__).resolve().parents[1]
        else:
            runtime_base = Path(runtime_base)

        cfg_path = runtime_base / 'data' / 'config' / 'import_rules.json'
        cfg_path.parent.mkdir(parents=True, exist_ok=True)

        merged = self._deep_merge_dict(self._default_import_rules(), self.rules if isinstance(self.rules, dict) else {})
        with cfg_path.open('w', encoding='utf-8') as f:
            json.dump(merged, f, ensure_ascii=False, indent=2)
            f.write('\n')

        self.rules = merged
        return str(cfg_path)
    
    def extract_tasks(self, file_path):
        """
        Extrahiert Aufgaben aus einer Word-Datei.

        Struktur-Regel:
        - Überschrift 1 = Kategorie
        - Überschrift 2 = neue Aufgabe
        - Alle folgenden Elemente bis zur nächsten Überschrift 2 oder 1
          gehören zur aktuellen Aufgabe.
        
        Args:
            file_path (str): Pfad zur Word-Datei
            
        Returns:
            list: Liste von Aufgaben-Dictionaries mit kompletter Struktur
        """
        try:
            doc = Document(file_path)
            tasks = self._extract_tasks_with_parser_mode(doc)
            self._normalize_task_collection(tasks)
            self._annotate_external_table_references(tasks, file_path)
            
            return tasks
            
        except Exception as e:
            raise Exception(f"Fehler beim Lesen der Word-Datei: {str(e)}")

    def _extract_tasks_with_parser_mode(self, doc):
        """Extrahiert Aufgaben gemäß parser_rules (auto/headings/tables/mixed)."""
        mode = str(self.get_import_rule('parser_rules.mode', 'auto') or 'auto').strip().lower()
        prefer_mode = str(self.get_import_rule('parser_rules.prefer_mode_on_mixed', 'headings') or 'headings').strip().lower()
        fallback_on_empty = bool(self.get_import_rule('parser_rules.fallback_to_secondary_on_empty', True))
        include_secondary_in_mixed = bool(self.get_import_rule('parser_rules.include_secondary_in_mixed', False))

        heading_tasks = self._extract_tasks_from_headings(doc)
        table_tasks = self._extract_tasks_from_structured_tables(doc)

        if mode == 'headings':
            return heading_tasks

        if mode == 'tables':
            return table_tasks

        if mode == 'mixed':
            primary = heading_tasks if prefer_mode != 'tables' else table_tasks
            secondary = table_tasks if prefer_mode != 'tables' else heading_tasks

            if primary and secondary and include_secondary_in_mixed:
                combined = list(primary) + list(secondary)
                self._reindex_task_numbers(combined)
                return combined

            if primary:
                return primary
            return secondary if fallback_on_empty else primary

        # auto (Default): H1/H2 bevorzugen, auf Tabellen zurückfallen
        if heading_tasks:
            return heading_tasks
        return table_tasks if fallback_on_empty else []

    def _extract_tasks_from_headings(self, doc):
        """Extrahiert Aufgaben aus H1/H2-Struktur (Kategorie/Aufgabe)."""
        tasks = []
        all_elements = self._get_all_body_elements(doc)

        current_category = "Ohne Kategorie"
        current_task = None

        def _finalize_current_task():
            if current_task is None:
                return

            content = []
            for task_element in current_task['all_elements']:
                if task_element['type'] == 'paragraph' and task_element['content'].text.strip():
                    content.append(task_element['content'].text.strip())
                elif task_element['type'] == 'table':
                    table_text = self._extract_table_text_for_keywords(task_element['content'])
                    if table_text:
                        content.append(table_text)

            current_task['content'] = content
            current_task['original_paragraphs'] = [
                elem['content'] for elem in current_task['all_elements'] if elem['type'] == 'paragraph'
            ]
            tasks.append(current_task)

        for element in all_elements:
            if element['type'] != 'paragraph':
                if current_task is not None:
                    current_task['all_elements'].append(element)
                continue

            paragraph = element['content']
            paragraph_text = paragraph.text.strip()

            if self._is_heading1(paragraph):
                _finalize_current_task()
                current_task = None
                category_text = paragraph_text
                current_category = category_text if category_text else "Ohne Kategorie"
                continue

            if self._is_heading2(paragraph):
                _finalize_current_task()
                task_title = paragraph_text or f"Aufgabe {len(tasks) + 1}"
                _main_no, _sub_no, number_label = self._extract_main_sub_number(task_title)
                current_task = {
                    'number': len(tasks) + 1,
                    'number_display': number_label or str(len(tasks) + 1),
                    'category': current_category,
                    'title': self._extract_task_title(task_title),
                    'content': [],
                    'all_elements': [element],
                    'original_paragraphs': [],
                    'difficulty': 'Mittel',
                    'keywords': []
                }
                continue

            if current_task is not None:
                current_task['all_elements'].append(element)

        _finalize_current_task()
        return tasks

    def _normalize_task_collection(self, tasks):
        """Normalisiert Metadaten und Nummern aller erkannten Aufgaben."""
        for task in tasks:
            full_content = ' '.join(task.get('content', []))
            if full_content:
                explicit_difficulty = self._extract_explicit_difficulty(full_content)
                if explicit_difficulty:
                    task['difficulty'] = explicit_difficulty
                else:
                    task['difficulty'] = self._extract_difficulty(full_content)

                explicit_keywords = self._extract_explicit_keywords(full_content)
                if explicit_keywords:
                    task['keywords'] = explicit_keywords
                else:
                    task['keywords'] = self._extract_keywords(full_content)

            number_display_raw = task.get('number_display') or task.get('number')
            is_intro_context = self._is_intro_context(
                task.get('content', []),
                task.get('title', ''),
            )
            normalized_number = self._normalize_intro_number_display(number_display_raw, is_intro_context)
            if normalized_number:
                task['number_display'] = normalized_number

    def _reindex_task_numbers(self, tasks):
        """Vergibt laufende interne Nummern (1..n) für kombinierte Parser-Ergebnisse."""
        for idx, task in enumerate(tasks, 1):
            task['number'] = idx

    def _external_table_marker_name(self):
        """Liefert den Marker-Namen für externe Tabellenreferenzen."""
        return str(self.get_import_rule('external_table_rules.marker_name', 'tabelle') or 'tabelle').strip().lower()

    def _external_table_orientation_marker_name(self):
        """Liefert den Marker-Namen für optionale Orientierungs-Overrides."""
        return str(self.get_import_rule('external_table_rules.orientation_marker_name', 'tabelle_format') or 'tabelle_format').strip().lower()

    def _parse_external_table_reference_line(self, text):
        """Parst Marker wie <<tabelle=datei>> oder <<tabelle_format=landscape>>."""
        raw = str(text or '').strip()
        if not raw:
            return None

        table_marker = re.escape(self._external_table_marker_name())
        orientation_marker = re.escape(self._external_table_orientation_marker_name())

        table_match = re.match(
            rf'^\s*(?:<<\s*)?{table_marker}\s*=\s*(.+?)(?:\s*>>)?\s*$',
            raw,
            flags=re.IGNORECASE,
        )
        if table_match:
            return {
                'type': 'table',
                'value': str(table_match.group(1) or '').strip().strip('>'),
            }

        orientation_match = re.match(
            rf'^\s*(?:<<\s*)?{orientation_marker}\s*=\s*(.+?)(?:\s*>>)?\s*$',
            raw,
            flags=re.IGNORECASE,
        )
        if orientation_match:
            return {
                'type': 'orientation',
                'value': str(orientation_match.group(1) or '').strip().strip('>').lower(),
            }

        return None

    def _strip_external_table_reference_from_lines(self, lines):
        """Entfernt Referenzmarker aus Inhaltszeilen und extrahiert deren Metadaten."""
        filtered_lines = []
        table_reference = ''
        orientation = 'auto'

        for line in (lines or []):
            text = str(line or '').strip()
            marker = self._parse_external_table_reference_line(text)
            if not marker:
                if text:
                    filtered_lines.append(text)
                continue

            if marker.get('type') == 'table' and not table_reference:
                table_reference = str(marker.get('value') or '').strip()
            elif marker.get('type') == 'orientation':
                candidate = str(marker.get('value') or '').strip().lower()
                if candidate in {'auto', 'landscape', 'portrait'}:
                    orientation = candidate

        return {
            'lines': filtered_lines,
            'reference': table_reference,
            'orientation': orientation,
        }

    def _collection_subfolder_name_from_path(self, collection_path):
        """Leitet den lernbereichsspezifischen Unterordner aus dem Sammlungsnamen ab."""
        stem = str(Path(collection_path).stem or '').strip()
        if not stem:
            return ''

        normalized = re.sub(r'^Aufgaben[_\-\s]*', '', stem, flags=re.IGNORECASE).strip()
        return normalized or stem

    def _resolve_external_table_reference(self, reference, collection_path):
        """Löst einen Tabellenverweis relativ zum Sammlungsordner/Lernbereich auf."""
        ref = str(reference or '').strip()
        if not ref or not collection_path:
            return ''

        ref_name = Path(ref).name
        ref_path = Path(ref_name)
        default_ext = str(self.get_import_rule('external_table_rules.default_extension', '.docx') or '.docx').strip() or '.docx'
        if not ref_path.suffix:
            ref_path = ref_path.with_suffix(default_ext)

        collection_file = Path(collection_path)
        base_dir = collection_file.parent
        candidates = []

        if bool(self.get_import_rule('external_table_rules.search_subfolder_from_collection_name', True)):
            subfolder_name = self._collection_subfolder_name_from_path(collection_path)
            if subfolder_name:
                candidates.append(base_dir / subfolder_name / ref_path.name)

        candidates.append(base_dir / ref_path.name)

        for candidate in candidates:
            if candidate.exists() and candidate.is_file():
                return str(candidate)

        return ''

    def _annotate_external_table_references(self, tasks, collection_path):
        """Erkennt Tabellenmarker in Aufgaben und ergänzt Pfad-/Orientierungsmetadaten."""
        for task in tasks or []:
            parsed = self._strip_external_table_reference_from_lines(task.get('content', []) or [])
            task['content'] = list(parsed.get('lines') or [])
            task['source_collection_path'] = str(collection_path)

            reference = str(parsed.get('reference') or '').strip()
            orientation = str(parsed.get('orientation') or 'auto').strip().lower() or 'auto'
            resolved_path = self._resolve_external_table_reference(reference, collection_path) if reference else ''

            task['external_table_reference'] = reference
            task['external_table_orientation'] = orientation
            task['external_table_path'] = resolved_path
            task['external_table_missing'] = bool(reference) and not bool(resolved_path)

    def _task_external_table_label(self, task):
        """Liefert eine kompakte Label-Darstellung für Preview/Diagnostik."""
        reference = str((task or {}).get('external_table_reference') or '').strip()
        if not reference:
            return ''

        path = str((task or {}).get('external_table_path') or '').strip()
        filename = Path(path).name if path else reference
        orientation = str((task or {}).get('external_table_orientation') or 'auto').strip().lower() or 'auto'
        if orientation == 'landscape':
            orientation_label = 'Querformat'
        elif orientation == 'portrait':
            orientation_label = 'Hochformat'
        else:
            orientation_label = 'Auto'

        return f"Externe Tabelle: {filename} ({orientation_label})"

    def _extract_tasks_from_structured_tables(self, doc):
        """
        Extrahiert Aufgaben aus strukturierten 2-Spalten-Tabellen (Label/Wert).

        Erwartetes Muster je Tabelle (mindestens):
        - ID
        - Aufgabenstellung (Pflicht)
        Optional:
        - Intro/Einleitung
        - Lösungsmöglichkeit/Hinweis
        - Schlagworte
        - Schwierigkeitsgrad
        - Kategorie
        """
        tasks = []
        task_number = 1

        for table in doc.tables:
            parsed_task = self._parse_structured_task_table(table, task_number)
            if parsed_task is None:
                continue

            tasks.append(parsed_task)
            task_number += 1

        return tasks

    def _extract_main_sub_number(self, text):
        """Extrahiert Haupt-/Nebennummer am Anfang eines Textes (z. B. 1, 1.1, Aufgabe 2.3)."""
        raw = str(text or '').strip()
        if not raw:
            return None, None, None

        match = re.match(
            r'^\s*(?:aufgabe\s*)?(\d+)(?:\s*[\.,]\s*(\d+))?(?:\s*[:\)\-]|\s+|$)',
            raw,
            flags=re.IGNORECASE,
        )
        if not match:
            return None, None, None

        main_number = match.group(1)
        sub_number = match.group(2)
        number_label = f"{main_number}.{sub_number}" if sub_number else main_number
        return main_number, sub_number, number_label

    def _is_intro_context(self, content_lines, title=''):
        """Erkennt, ob es sich inhaltlich um einen Einleitungs-/Kontextblock handelt."""
        markers = (
            'einleitung',
            'intro',
            'kontext',
            'ausgangslage',
            'vorbemerkung',
            'einfuehr',
            'einführ',
        )

        candidates = []
        if str(title or '').strip():
            candidates.append(str(title).strip().lower())

        for line in (content_lines or [])[:3]:
            text = str(line or '').strip().lower()
            if text:
                candidates.append(text)

        return any(any(marker in text for marker in markers) for text in candidates)

    def _normalize_intro_number_display(self, number_display, is_intro_context):
        """Normalisiert Einleitungen auf Nebennummer .0 (z. B. 1 -> 1.0)."""
        label = str(number_display or '').strip()
        if not label or not is_intro_context:
            return label

        main_number, sub_number, _label = self._extract_main_sub_number(label)
        if main_number and sub_number is None:
            return f"{main_number}.0"

        return label

    def append_task_from_document(
        self,
        source_doc_path,
        target_collection_path,
        category,
        difficulty='Mittel',
        keywords='',
        intro='',
        hint='',
        task_id='',
        title=''
    ):
        """
        Übernimmt eine neue Aufgabe aus einer Word-Datei in eine tabellenbasierte Aufgabensammlung.

        Die gesamte Quelle (Paragraphen, Tabellen, Formatierungen; best effort inkl. eingebetteter Inhalte)
        wird in das Feld "Aufgabenstellung (Pflicht)" der neuen Tabellen-Aufgabe übernommen.

        Returns:
            dict: Informationen zur übernommenen Aufgabe (id, target_file)
        """
        if not os.path.exists(source_doc_path):
            raise FileNotFoundError(f"Quelldatei nicht gefunden: {source_doc_path}")
        if not os.path.exists(target_collection_path):
            raise FileNotFoundError(f"Zieldatei nicht gefunden: {target_collection_path}")

        source_doc = Document(source_doc_path)
        target_doc = Document(target_collection_path)

        template_table = self._find_first_structured_task_table(target_doc)
        if template_table is None:
            raise ValueError(
                "Die Zieldatei enthält keine erkannte Aufgaben-Tabellenstruktur. "
                "Bitte eine Sammlung auf Basis des Aufgaben-Gerüsts verwenden."
            )

        new_table = self._append_table_clone(target_doc, template_table)

        resolved_id = task_id.strip() if str(task_id or '').strip() else self._generate_next_task_id(target_doc)
        resolved_category = str(category or '').strip() or 'Ohne Kategorie'
        resolved_difficulty = str(difficulty or '').strip() or 'Mittel'
        resolved_title = str(title or '').strip() or self._derive_title_from_source_document(source_doc)

        self._set_structured_table_value(new_table, 'id', resolved_id)
        self._set_structured_table_value(new_table, 'titel', resolved_title)
        self._set_structured_table_value(new_table, 'kategorie', resolved_category)
        self._set_structured_table_value(new_table, 'introeinleitungoptional', str(intro or '').strip())
        self._set_structured_table_value(new_table, 'loesungsmoeglichkeithinweisoptional', str(hint or '').strip())
        self._set_structured_table_value(new_table, 'schlagwortekommagetrennt', str(keywords or '').strip())
        self._set_structured_table_value(new_table, 'schwierigkeitsgrad', resolved_difficulty)

        task_cell = self._get_structured_table_value_cell(new_table, 'aufgabenstellungpflicht')
        if task_cell is None:
            task_cell = self._ensure_structured_row(new_table, 'Aufgabenstellung (Pflicht)')

        self._replace_cell_content_from_document(target_doc, task_cell, source_doc)

        backup_path = self._create_collection_backup(target_collection_path)
        target_doc.save(target_collection_path)
        return {
            'id': resolved_id,
            'title': resolved_title,
            'target_file': target_collection_path,
            'backup_file': backup_path,
        }

    def preview_task_append(
        self,
        source_doc_path,
        target_collection_path,
        category,
        difficulty='Mittel',
        keywords='',
        title='',
    ):
        """
        Liefert eine Vorschau für die Aufgabenübernahme ohne Dateiänderung.

        Returns:
            dict mit Ziel-ID, Quellstruktur, Metadaten und ggf. Validierungshinweisen.
        """
        if not os.path.exists(source_doc_path):
            raise FileNotFoundError(f"Quelldatei nicht gefunden: {source_doc_path}")
        if not os.path.exists(target_collection_path):
            raise FileNotFoundError(f"Zieldatei nicht gefunden: {target_collection_path}")

        source_doc = Document(source_doc_path)
        target_doc = Document(target_collection_path)

        next_id = self._generate_next_task_id(target_doc)
        resolved_title = str(title or '').strip() or self._derive_title_from_source_document(source_doc)
        nonempty_paragraphs = [p.text.strip() for p in source_doc.paragraphs if p.text and p.text.strip()]
        preview_lines = nonempty_paragraphs[:3]
        source_blocks = []

        for child in source_doc._element.body:
            tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_local == 'p':
                text = ''
                for para in source_doc.paragraphs:
                    if para._element == child:
                        text = para.text.strip()
                        break
                if text:
                    source_blocks.append({'type': 'paragraph', 'text': text})
            elif tag_local == 'tbl':
                table_obj = None
                for table in source_doc.tables:
                    if table._element == child:
                        table_obj = table
                        break

                if table_obj is not None:
                    rows = len(table_obj.rows)
                    cols = len(table_obj.columns) if table_obj.columns else 0
                    first_cell = ''
                    if rows > 0 and cols > 0:
                        first_cell = table_obj.cell(0, 0).text.strip()
                    source_blocks.append({
                        'type': 'table',
                        'rows': rows,
                        'cols': cols,
                        'first_cell': first_cell,
                    })

        normalized_difficulty = self._extract_explicit_difficulty(f"Schwierigkeit: {difficulty}")
        if not normalized_difficulty:
            normalized_difficulty = self._extract_difficulty(str(difficulty or ''))

        duplicate_check = self._detect_potential_duplicate(source_doc, target_doc)

        max_preview_blocks = int(self._rule_value('max_preview_blocks', 10) or 10)

        return {
            'next_id': next_id,
            'title': resolved_title,
            'source_paragraph_count': len(nonempty_paragraphs),
            'source_table_count': len(source_doc.tables),
            'source_preview_lines': preview_lines,
            'source_preview_blocks': source_blocks[:max_preview_blocks],
            'category': str(category or '').strip() or 'Ohne Kategorie',
            'difficulty_input': str(difficulty or '').strip(),
            'difficulty_normalized': normalized_difficulty,
            'keywords': str(keywords or '').strip(),
            'difficulty_inconsistent': self._has_inconsistent_difficulty(difficulty),
            'duplicate_check': duplicate_check,
        }

    def _build_source_signature(self, source_doc):
        """Erstellt eine normalisierte Textsignatur der Quellaufgabe."""
        chunks = []

        for child in source_doc._element.body:
            tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_local == 'p':
                for para in source_doc.paragraphs:
                    if para._element == child:
                        text = para.text.strip()
                        if text:
                            chunks.append(text)
                        break
            elif tag_local == 'tbl':
                for table in source_doc.tables:
                    if table._element == child:
                        table_text = self._extract_table_text_for_keywords(table)
                        if table_text:
                            chunks.append(table_text)
                        break

        return self._normalize_signature_text(' '.join(chunks))

    def _task_signature(self, task):
        """Erstellt eine normalisierte Vergleichssignatur aus einem Task-Dictionary."""
        parts = []
        title = str(task.get('title') or '').strip()
        if title:
            parts.append(title)

        content_lines = task.get('content') or []
        parts.extend(str(line).strip() for line in content_lines if str(line).strip())

        return self._normalize_signature_text(' '.join(parts))

    def _normalize_signature_text(self, text):
        """Normalisiert Text für robuste Ähnlichkeitsvergleiche."""
        value = (text or '').lower().strip()
        value = re.sub(r'\s+', ' ', value)
        value = re.sub(r'[^a-z0-9äöüß ]+', ' ', value)
        value = re.sub(r'\s+', ' ', value).strip()
        return value

    def _detect_potential_duplicate(self, source_doc, target_doc, threshold=None):
        """
        Prüft die Quellaufgabe gegen vorhandene Aufgaben der Zielsammlung.

        Returns:
            dict: {is_duplicate, similarity, matched_title, matched_id, matched_category}
        """
        if threshold is None:
            threshold = self._resolve_duplicate_threshold()

        source_sig = self._build_source_signature(source_doc)
        if not source_sig:
            return {
                'is_duplicate': False,
                'similarity': 0.0,
                'matched_title': '',
                'matched_id': '',
                'matched_category': '',
            }

        existing_tasks = self._extract_tasks_from_structured_tables(target_doc)
        best_similarity = 0.0
        best_task = None

        for task in existing_tasks:
            candidate_sig = self._task_signature(task)
            if not candidate_sig:
                continue
            similarity = SequenceMatcher(None, source_sig, candidate_sig).ratio()
            if similarity > best_similarity:
                best_similarity = similarity
                best_task = task

        if best_task is None:
            return {
                'is_duplicate': False,
                'similarity': 0.0,
                'matched_title': '',
                'matched_id': '',
                'matched_category': '',
            }

        return {
            'is_duplicate': best_similarity >= threshold,
            'similarity': round(best_similarity, 3),
            'matched_title': best_task.get('title', ''),
            'matched_id': best_task.get('id', ''),
            'matched_category': best_task.get('category', ''),
        }

    def _normalize_table_key(self, value):
        """Normalisiert Tabellen-Keys für robuste Label-Erkennung."""
        text = (value or '').strip().lower()
        replacements = {
            'ä': 'ae',
            'ö': 'oe',
            'ü': 'ue',
            'ß': 'ss',
        }
        for src, dst in replacements.items():
            text = text.replace(src, dst)

        text = re.sub(r'\([^\)]*\)', '', text)
        text = re.sub(r'[^a-z0-9]+', '', text)
        return text

    def _find_first_structured_task_table(self, doc):
        """Liefert die erste erkannte Aufgaben-Tabelle (ID + Aufgabenstellung) oder None."""
        for table in doc.tables:
            keys = set()
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                keys.add(self._normalize_table_key(row.cells[0].text))

            if 'id' in keys and ('aufgabenstellungpflicht' in keys or 'aufgabenstellung' in keys):
                return table
        return None

    def _append_table_clone(self, target_doc, source_table):
        """Kloniert eine Tabelle ans Ende des Dokuments und liefert die neue Tabelle zurück."""
        body = target_doc._element.body
        sect_pr = body.sectPr
        cloned_tbl = deepcopy(source_table._tbl)

        if sect_pr is not None:
            body.insert(list(body).index(sect_pr), cloned_tbl)
        else:
            body.append(cloned_tbl)

        return target_doc.tables[-1]

    def _iter_table_rows_by_norm_key(self, table):
        """Iteriert über Tabellenzeilen als Tupel (norm_key, row)."""
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = self._normalize_table_key(row.cells[0].text)
            if key:
                yield key, row

    def _clear_cell_content(self, cell):
        """Entfernt alle Blockinhalte aus einer Tabellenzelle (bis auf tcPr)."""
        tc = cell._tc
        for child in list(tc):
            # tcPr erhalten, Inhaltsknoten löschen
            tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_local != 'tcPr':
                tc.remove(child)

    def _set_structured_table_value(self, table, key_norm, value):
        """Setzt den Wert (2. Spalte) für einen strukturierten Tabellen-Key."""
        target_cell = self._get_structured_table_value_cell(table, key_norm)
        if target_cell is None:
            label = self._key_norm_to_label(key_norm)
            target_cell = self._ensure_structured_row(table, label)

        self._clear_cell_content(target_cell)
        target_cell.text = str(value or '')

    def _get_structured_table_value_cell(self, table, key_norm):
        """Liefert die Wertzelle (2. Spalte) für key_norm oder None."""
        for row_key, row in self._iter_table_rows_by_norm_key(table):
            if row_key == key_norm:
                return row.cells[1]
        return None

    def _ensure_structured_row(self, table, label):
        """Stellt sicher, dass eine Label-Zeile existiert, und liefert die Wertzelle zurück."""
        wanted = self._normalize_table_key(label)
        for row_key, row in self._iter_table_rows_by_norm_key(table):
            if row_key == wanted:
                return row.cells[1]

        new_row = table.add_row()
        new_row.cells[0].text = label
        new_row.cells[1].text = ''
        return new_row.cells[1]

    def _key_norm_to_label(self, key_norm):
        """Mappt normalisierte Keys auf lesbare Standard-Labels."""
        mapping = self._field_labels()
        return mapping.get(key_norm, key_norm)

    def _derive_title_from_source_document(self, source_doc):
        """Leitet einen robusten Titel aus der Quellaufgabe ab (erste nicht-leere Zeile)."""
        for para in source_doc.paragraphs:
            text = str(para.text or '').strip()
            if text:
                return self._extract_task_title(text)

        # Fallback über Tabelleninhalt
        for table in source_doc.tables:
            table_text = self._extract_table_text_for_keywords(table)
            if table_text:
                first_line = table_text.split('\n', 1)[0].strip() if '\n' in table_text else table_text.strip()
                if first_line:
                    return self._extract_task_title(first_line)

        return 'Ohne Titel'

    def _replace_cell_content_from_document(self, target_doc, target_cell, source_doc):
        """
        Ersetzt den Zelleninhalt durch den Inhalt eines Quell-Dokuments.

        Hinweis: Beziehungen werden best effort in das Zieldokument übernommen.
        """
        self._clear_cell_content(target_cell)
        tc = target_cell._tc

        # Sicherstellen, dass Beziehungen pro old_rId nur einmal neu angelegt werden
        rel_map = {}

        for child in source_doc._element.body:
            tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_local not in ('p', 'tbl'):
                continue

            cloned = deepcopy(child)
            self._remap_relationship_ids_in_xml(
                xml_element=cloned,
                source_part=source_doc.part,
                target_part=target_doc.part,
                rel_map=rel_map,
            )
            tc.append(cloned)

        # Word erwartet mindestens einen Absatz in der Zelle
        has_block = any((c.tag.split('}')[-1] if '}' in c.tag else c.tag) in ('p', 'tbl') for c in tc)
        if not has_block:
            target_cell.text = ''

    def _remap_relationship_ids_in_xml(self, xml_element, source_part, target_part, rel_map):
        """
        Remappt r:* Beziehungs-IDs in einem XML-Element von source_part auf target_part.
        """
        rel_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

        for node in xml_element.iter():
            for attr_name, old_rid in list(node.attrib.items()):
                if not attr_name.startswith(rel_ns):
                    continue

                if old_rid in rel_map:
                    node.set(attr_name, rel_map[old_rid])
                    continue

                new_rid = self._clone_relationship(source_part, target_part, old_rid)
                if new_rid:
                    rel_map[old_rid] = new_rid
                    node.set(attr_name, new_rid)

    def _clone_relationship(self, source_part, target_part, old_rid):
        """
        Klont eine Relationship von source_part nach target_part und liefert neue rId.

        Best effort: Bei nicht unterstützten Relationship-Typen wird alte rId zurückgegeben.
        """
        try:
            rel = source_part.rels[old_rid]
        except Exception:
            return old_rid

        try:
            if getattr(rel, 'is_external', False):
                return target_part.relate_to(rel.target_ref, rel.reltype, is_external=True)

            target_obj = getattr(rel, 'target_part', None) or getattr(rel, '_target', None)
            if target_obj is None:
                return old_rid

            return target_part.relate_to(target_obj, rel.reltype)
        except Exception:
            return old_rid

    def _generate_next_task_id(self, target_doc):
        """Erzeugt eine nächste Aufgaben-ID auf Basis vorhandener IDs (A-001, A-001.1, ...)."""
        ids = []

        for table in target_doc.tables:
            id_cell = self._get_structured_table_value_cell(table, 'id')
            if id_cell is None:
                continue
            val = id_cell.text.strip()
            if val:
                ids.append(val)

        if not ids:
            return 'A-001'

        nums = []
        for val in ids:
            m = re.search(r'(\d+)(?!.*\d)', val)
            if m:
                try:
                    nums.append(int(m.group(1)))
                except ValueError:
                    pass

        if not nums:
            return 'A-001'

        return f"A-{max(nums) + 1:03d}"

    def _create_collection_backup(self, target_collection_path):
        """Erstellt eine zeitgestempelte Sicherungskopie der Ziel-Aufgabensammlung."""
        source_path = Path(target_collection_path)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_name = f"{source_path.stem}_backup_{timestamp}{source_path.suffix}"
        backup_path = source_path.with_name(backup_name)
        shutil.copy2(source_path, backup_path)
        return str(backup_path)

    def migrate_missing_titles_in_collection(self, collection_path):
        """Ergänzt fehlende Titel in einer tabellenbasierten Aufgabensammlung automatisch.

        Für jede erkannte Aufgaben-Tabelle (ID + Aufgabenstellung) wird bei leerem/fehlendem
        Feld `Titel` ein Titel aus der Aufgabenstellung abgeleitet.

        Returns:
            dict: {
                changed_tasks: int,
                scanned_tasks: int,
                backup_file: str|None,
                updated_ids: list[str],
                updated_entries: list[dict],
            }
        """
        if not os.path.exists(collection_path):
            raise FileNotFoundError(f"Datei nicht gefunden: {collection_path}")

        doc = Document(collection_path)
        scanned = 0
        changed = 0
        updated_ids = []
        updated_entries = []

        for table in doc.tables:
            keys = {key for key, _row in self._iter_table_rows_by_norm_key(table)}
            is_task_table = 'id' in keys and ('aufgabenstellungpflicht' in keys or 'aufgabenstellung' in keys)
            if not is_task_table:
                continue

            scanned += 1

            title_cell = self._get_structured_table_value_cell(table, 'titel')
            current_title = str(title_cell.text or '').strip() if title_cell is not None else ''
            if current_title:
                continue

            task_cell = self._get_structured_table_value_cell(table, 'aufgabenstellungpflicht')
            if task_cell is None:
                task_cell = self._get_structured_table_value_cell(table, 'aufgabenstellung')

            task_text = str(task_cell.text or '').strip() if task_cell is not None else ''
            derived_title = self._extract_task_title(task_text.split('\n', 1)[0].strip()) if task_text else ''
            if not derived_title:
                derived_title = 'Ohne Titel'

            self._set_structured_table_value(table, 'titel', derived_title)
            changed += 1

            id_cell = self._get_structured_table_value_cell(table, 'id')
            task_id = str(id_cell.text or '').strip() if id_cell is not None else ''
            if task_id:
                updated_ids.append(task_id)

            updated_entries.append({
                'id': task_id,
                'title': derived_title,
            })

        backup_path = None
        if changed > 0:
            backup_path = self._create_collection_backup(collection_path)
            doc.save(collection_path)

        return {
            'changed_tasks': changed,
            'scanned_tasks': scanned,
            'backup_file': backup_path,
            'updated_ids': updated_ids,
            'updated_entries': updated_entries,
        }

    def _parse_structured_task_table(self, table, task_number):
        """Parst eine einzelne strukturierte Aufgaben-Tabelle in ein Task-Dictionary."""
        if not table.rows or len(table.columns) < 2:
            return None

        values_by_key = {}

        for row in table.rows:
            if len(row.cells) < 2:
                continue

            raw_key = row.cells[0].text.strip()
            raw_value = self._extract_cell_value_text(row.cells[1])
            norm_key = self._normalize_table_key(raw_key)
            if not norm_key:
                continue

            values_by_key[norm_key] = raw_value

        task_text = self._value_by_alias(values_by_key, 'aufgabenstellungpflicht', fallback_aliases=['aufgabenstellung'])
        if not task_text:
            return None

        required_fields = {
            str(v).strip().lower()
            for v in (self.get_import_rule('template_rules.required_fields', []) or [])
            if str(v).strip()
        }
        pre_warnings = []

        # Pflichtfelder prüfen (mit Fallback-Werten bleibt Verarbeitung robust, aber Warnung wird erzeugt)
        raw_id = self._value_by_alias(values_by_key, 'id')
        explicit_title = self._value_by_alias(values_by_key, 'titel')
        raw_task = self._value_by_alias(values_by_key, 'aufgabenstellungpflicht', fallback_aliases=['aufgabenstellung'])
        raw_difficulty = self._value_by_alias(values_by_key, 'schwierigkeitsgrad', fallback_aliases=['schwierigkeit'])
        raw_keywords = self._value_by_alias(
            values_by_key,
            'schlagwortekommagetrennt',
            fallback_aliases=['schlagworte', 'schluesselwoerter', 'schlusselwoerter'],
        )
        raw_category = self._value_by_alias(values_by_key, 'kategorie')

        labels = self._field_labels()
        required_mapping = {
            'id': (labels.get('id', 'ID'), raw_id),
            'titel': (labels.get('titel', 'Titel'), explicit_title),
            'aufgabenstellungpflicht': (labels.get('aufgabenstellungpflicht', 'Aufgabenstellung (Pflicht)'), raw_task),
            'schwierigkeitsgrad': (labels.get('schwierigkeitsgrad', 'Schwierigkeitsgrad'), raw_difficulty),
            'schlagwortekommagetrennt': (labels.get('schlagwortekommagetrennt', 'Schlagworte (kommagetrennt)'), raw_keywords),
            'kategorie': (labels.get('kategorie', 'Kategorie'), raw_category),
        }

        for req_key, (req_label, req_value) in required_mapping.items():
            if req_key in required_fields and self._is_missing_required_value(req_value):
                pre_warnings.append(f"Pflichtfeld fehlt: {req_label}")

        intro_text = self._value_by_alias(
            values_by_key,
            'introeinleitungoptional',
            fallback_aliases=['introeinleitung', 'einleitung', 'intro', 'kontext'],
        )
        hint_text = self._value_by_alias(
            values_by_key,
            'loesungsmoeglichkeithinweisoptional',
            fallback_aliases=['loesungsmoeglichkeithinweis', 'hinweis'],
        )
        difficulty_raw = raw_difficulty
        keywords_raw = raw_keywords

        category_required = bool(self.get_import_rule('category_rules.required', True))
        if self._is_missing_category(raw_category):
            category = 'Ohne Kategorie'
            if category_required:
                pre_warnings.append(
                    "Kategorie fehlt. Bitte vor Export in der Quelle eine gültige Kategorie pflegen."
                )
        else:
            category = str(raw_category).strip()

        if self._has_inconsistent_difficulty(difficulty_raw):
            pre_warnings.append(
                "Inkonsistenter Schwierigkeitsgrad erkannt (mehrere Werte). Bitte vor Export in der Quelle bereinigen."
            )

        if difficulty_raw:
            explicit_difficulty = self._extract_explicit_difficulty(f"Schwierigkeit: {difficulty_raw}")
            difficulty = explicit_difficulty if explicit_difficulty else self._extract_difficulty(difficulty_raw)
        else:
            difficulty = self._extract_difficulty(task_text)

        if keywords_raw:
            keywords = [k.strip().lower() for k in keywords_raw.split(',') if k.strip()]
        else:
            keywords = self._extract_keywords(task_text)

        content_lines = []
        if intro_text:
            content_lines.append(f"Einleitung: {intro_text}")
        content_lines.append(task_text)
        if hint_text:
            content_lines.append(f"Hinweis: {hint_text}")
        if difficulty_raw:
            content_lines.append(f"Schwierigkeit: {difficulty_raw}")
        if keywords_raw:
            content_lines.append(f"Schlüsselwörter: {keywords_raw}")

        task_id = values_by_key.get('id') or str(task_number)
        _main_no, _sub_no, number_label = self._extract_main_sub_number(task_id)
        base_number_display = number_label or str(task_number)
        is_intro_context = bool(str(intro_text or '').strip()) or self._is_intro_context([task_text], explicit_title)
        number_display = self._normalize_intro_number_display(base_number_display, is_intro_context)
        title_source = explicit_title or task_text.split('\n', 1)[0].strip()

        return {
            'number': task_number,
            'number_display': number_display or base_number_display,
            'id': task_id,
            'category': category,
            'title': self._extract_task_title(title_source),
            'content': content_lines,
            'all_elements': [{'type': 'table', 'content': table}],
            'original_paragraphs': [],
            'difficulty': difficulty,
            'keywords': keywords,
            'pre_warnings': pre_warnings,
        }

    def _extract_cell_value_text(self, cell):
        """Liest robusten Zelltext inkl. verschachtelter Tabellen und SDT-Texte aus."""
        if cell is None:
            return ''

        raw = str(cell.text or '').strip()
        if raw:
            return raw

        text_nodes = []
        try:
            for node in cell._tc.xpath('.//w:t'):
                value = str(node.text or '').strip()
                if value:
                    text_nodes.append(value)
        except Exception:
            return ''

        return '\n'.join(text_nodes).strip()

    def extract_tasks_with_diagnostics(self, file_path):
        """
        Extrahiert Aufgaben und ergänzt Diagnoseinformationen pro Aufgabe.

        Returns:
            tuple[list, dict]: (tasks, diagnostics_report)
        """
        tasks = self.extract_tasks(file_path)

        task_diagnostics = []
        low_confidence_count = 0
        warning_count = 0

        for task in tasks:
            diagnostic = self._build_task_diagnostic(task)
            task['intro'] = diagnostic['intro']
            task['warnings'] = diagnostic['warnings']
            task['hints'] = diagnostic.get('hints', [])
            task['confidence'] = diagnostic['confidence']

            if diagnostic['confidence'] == 'low':
                low_confidence_count += 1
            if diagnostic['warnings']:
                warning_count += 1

            task_diagnostics.append({
                'number': task.get('number', 0),
                'title': task.get('title', ''),
                'confidence': diagnostic['confidence'],
                'warnings': list(diagnostic['warnings']),
            })

        global_warnings = []
        if len(tasks) == 0:
            global_warnings.append("Keine Aufgaben erkannt. Prüfen Sie die Formatierung (Überschrift 2 für Aufgaben).")

        report = {
            'task_count': len(tasks),
            'low_confidence_count': low_confidence_count,
            'warning_task_count': warning_count,
            'global_warnings': global_warnings,
            'task_diagnostics': task_diagnostics,
        }

        return tasks, report

    def _build_task_diagnostic(self, task):
        """
        Erstellt Diagnosemetadaten (intro/warnings/confidence) für eine Aufgabe.
        """
        title = (task.get('title') or '').strip()
        content_lines = [str(line).strip() for line in (task.get('content') or []) if str(line).strip()]
        keywords = task.get('keywords') or []
        difficulty = (task.get('difficulty') or '').strip()

        warnings = list(task.get('pre_warnings') or [])
        hints = []

        if not content_lines:
            warnings.append('Aufgabe enthält keinen verwertbaren Inhalt.')

        if title and len(title) < 5:
            warnings.append('Aufgabentitel ist sehr kurz.')

        if not keywords:
            warnings.append('Keine Schlüsselwörter erkannt.')

        if difficulty in ('', 'Unbekannt'):
            warnings.append('Schwierigkeit ist nicht eindeutig.')

        category_required = bool(self.get_import_rule('category_rules.required', True))
        if category_required and self._is_missing_category(task.get('category')):
            warnings.append('Kategorie fehlt (Pflichtfeld).')

        if str(task.get('external_table_reference') or '').strip() and bool(task.get('external_table_missing')):
            warnings.append('Externe Tabelle konnte nicht gefunden werden.')

        intro_lines = self._extract_intro_lines(content_lines)
        if intro_lines:
            hints.append('Einleitung für Aufgabe und Folgeaufgaben erforderlich.')

        # Confidence-Heuristik (Sprint 1)
        score = 100
        for warning in warnings:
            if 'keinen verwertbaren Inhalt' in warning:
                score -= 45
            elif 'Einleitungs-/Kontexttext' in warning:
                score -= 20
            else:
                score -= 10

        if score >= 80:
            confidence = 'high'
        elif score >= 50:
            confidence = 'medium'
        else:
            confidence = 'low'

        return {
            'intro': intro_lines,
            'warnings': warnings,
            'hints': hints,
            'confidence': confidence,
        }

    def _extract_intro_lines(self, content_lines):
        """
        Erkennt mögliche Intro-/Kontext-Zeilen am Aufgabenanfang.
        """
        if not content_lines:
            return []

        intro_markers = (
            'einleitung',
            'intro',
            'hinweis',
            'kontext',
            'ausgangslage',
            'vorbemerkung',
        )

        intro = []
        for line in content_lines[:2]:  # nur die ersten Zeilen als Intro-Kandidaten
            line_lower = line.lower()
            if any(marker in line_lower for marker in intro_markers):
                intro.append(line)

        return intro
    
    def _get_heading_level(self, style_name):
        """
        Bestimmt das korrekte Heading-Level basierend auf dem Style-Namen
        
        Args:
            style_name: Name des Word-Styles
            
        Returns:
            int: Das entsprechende Heading-Level (1-9)
        """
        if not style_name:
            return 1
            
        style_lower = style_name.lower()
        
        # Direkte Zuordnung für deutsche und englische Styles
        heading_map = {
            'heading 1': 1, 'überschrift 1': 1,
            'heading 2': 2, 'überschrift 2': 2,
            'heading 3': 3, 'überschrift 3': 3,
            'heading 4': 4, 'überschrift 4': 4,
            'heading 5': 5, 'überschrift 5': 5,
            'heading 6': 6, 'überschrift 6': 6,
            'heading 7': 7, 'überschrift 7': 7,
            'heading 8': 8, 'überschrift 8': 8,
            'heading 9': 9, 'überschrift 9': 9
        }
        
        if style_lower in heading_map:
            return heading_map[style_lower]
        
        # Fallback: Versuche Zahl aus Style-Namen zu extrahieren
        import re
        match = re.search(r'\d+', style_name)
        if match:
            level = int(match.group())
            return min(max(level, 1), 9)  # Begrenze auf 1-9
        
        return 1  # Standard-Fallback

    def _is_heading1(self, paragraph):
        """
        Überprüft, ob ein Paragraph eine Überschrift 1 ist
        
        Args:
            paragraph: Word-Paragraph-Objekt
            
        Returns:
            bool: True wenn es eine Überschrift 1 ist
        """
        try:
            if paragraph.style:
                style_name = paragraph.style.name
                return (style_name == 'Heading 1' or 
                        style_name == 'Überschrift 1' or
                        style_name.lower() == 'heading 1' or
                        style_name.lower() == 'überschrift 1')
        except:
            # Fallback: Prüfe auf Textmuster wenn Style nicht verfügbar
            text = paragraph.text.strip()
            if text:
                return self._is_task_start(text, paragraph)
        
        return False

    def _is_heading2(self, paragraph):
        """
        Überprüft, ob ein Paragraph eine Überschrift 2 ist.

        Args:
            paragraph: Word-Paragraph-Objekt

        Returns:
            bool: True wenn es eine Überschrift 2 ist
        """
        try:
            if paragraph.style:
                style_name = paragraph.style.name
                return (
                    style_name == 'Heading 2'
                    or style_name == 'Überschrift 2'
                    or style_name.lower() == 'heading 2'
                    or style_name.lower() == 'überschrift 2'
                )
        except:
            return False

        return False

    def _get_all_body_elements(self, doc):
        """
        Sammelt alle Body-Elemente (Paragraphen und Tabellen) in der richtigen Reihenfolge
        
        Args:
            doc: Word-Dokument
            
        Returns:
            list: Liste von Element-Dictionaries mit 'type' und 'content'
        """
        elements = []
        
        # Nutze XML-Struktur für korrekte Reihenfolge
        body = doc._element.body
        
        for child in body:
            tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag_name == 'p':  # Paragraph
                # Finde entsprechendes Paragraph-Objekt
                for para in doc.paragraphs:
                    if para._element == child:
                        elements.append({
                            'type': 'paragraph',
                            'content': para
                        })
                        break
                        
            elif tag_name == 'tbl':  # Table
                # Finde entsprechendes Table-Objekt
                for table in doc.tables:
                    if table._element == child:
                        elements.append({
                            'type': 'table',
                            'content': table
                        })
                        break

            elif tag_name == 'sdt':  # Inhaltssteuerelement (Content Control)
                elements.append({
                    'type': 'sdt',
                    'content': child  # Kein python-docx-Wrapper – raw XML-Element
                })
        
        return elements

    def _extract_table_text_for_keywords(self, table):
        """
        Extrahiert Text aus Tabelle für Keyword-Analyse
        
        Args:
            table: Word-Table-Objekt
            
        Returns:
            str: Tabellen-Text zusammengefasst
        """
        try:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' '.join(row_text))
            return ' '.join(table_text)
        except:
            return ""
    
    def _create_cover_page(self, doc, lek_theme=""):
        """
        Erstellt ein Deckblatt (Semaphor) mit Farbschema basierend auf #00a499
        
        Args:
            doc: Das Word-Dokument
            lek_theme (str): LEK-Thema für das Deckblatt
        """
        # Farbpalette basierend auf #00a499
        primary_color = RGBColor(0, 164, 153)      # #00a499
        light_color = RGBColor(102, 204, 195)      # Hellere Version
        dark_color = RGBColor(0, 120, 112)         # Dunklere Version
        
        # Haupttitel
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run("LERNERFOLGSKONTROLLE")
        title_run.font.name = 'Aptos'
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = primary_color
        title_run.bold = True
        
        # Leerzeilen
        for _ in range(3):
            doc.add_paragraph()
        
        # Semaphor-Design (vereinfacht mit Text-Elementen)
        semaphor_paragraph = doc.add_paragraph()
        semaphor_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Oberer Balken
        top_run = semaphor_paragraph.add_run("━━━━━━━━━━━━━━━━━━━━")
        top_run.font.size = Pt(16)
        top_run.font.color.rgb = dark_color
        top_run.bold = True
        
        doc.add_paragraph()
        
        # Mittlerer Bereich
        middle_paragraph = doc.add_paragraph()
        middle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if lek_theme:
            middle_run = middle_paragraph.add_run(f"◆  {lek_theme.upper()}  ◆")
        else:
            middle_run = middle_paragraph.add_run("◆  AUFGABENSAMMLUNG  ◆")
            
        middle_run.font.name = 'Aptos'
        middle_run.font.size = Pt(18)
        middle_run.font.color.rgb = light_color
        middle_run.bold = True
        
        doc.add_paragraph()
        
        # Unterer Balken
        bottom_paragraph = doc.add_paragraph()
        bottom_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom_run = bottom_paragraph.add_run("━━━━━━━━━━━━━━━━━━━━")
        bottom_run.font.size = Pt(16)
        bottom_run.font.color.rgb = dark_color
        bottom_run.bold = True
        
        # Weitere Leerzeilen
        for _ in range(8):
            doc.add_paragraph()
        
        # Datum
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from datetime import datetime
        date_run = date_paragraph.add_run(f"Erstellt am {datetime.now().strftime('%d.%m.%Y')}")
        date_run.font.name = 'Aptos'
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = primary_color
        date_run.italic = True
    
    def _copy_elements_with_formatting(self, doc, elements):
        """
        Kopiert alle Elemente (Paragraphen und Tabellen) mit ihrer ursprünglichen Formatierung
        
        Args:
            doc: Ziel-Word-Dokument
            elements: Liste von Element-Dictionaries mit 'type' und 'content'
        """
        for element in elements:
            if element['type'] == 'paragraph':
                # Verwende bestehende Paragraph-Kopierfunktion
                self._copy_paragraphs_with_formatting(doc, [element['content']])
                
            elif element['type'] == 'table':
                # Kopiere Tabelle mit vollständiger Formatierung
                self._copy_table_with_formatting(doc, element['content'])

    def _copy_elements_for_lek(self, doc, elements):
        """
        Spezielle Kopierfunktion für LEK-Erstellung mit angepassten Überschrift-Levels
        
        Args:
            doc: Ziel-Word-Dokument (LEK)
            elements: Liste von Element-Dictionaries mit 'type' und 'content'
        """
        from copy import deepcopy
        from docx.oxml.ns import qn as _qn

        body = doc._element.body
        sect_pr = body.find(_qn('w:sectPr'))

        for element in elements:
            try:
                etype = element['type']
                if etype in ('paragraph', 'table'):
                    xml_elem = element['content']._element
                elif etype == 'sdt':
                    xml_elem = element['content']
                else:
                    continue

                new_elem = deepcopy(xml_elem)
                self._remap_num_ids_in_element(new_elem, self._num_id_map)

                if etype == 'table':
                    # Vor Tabellen im LEK immer eine Leerzeile
                    doc.add_paragraph()
                    self._set_table_full_width_xml(new_elem)

                if sect_pr is not None:
                    body.insert(list(body).index(sect_pr), new_elem)
                else:
                    body.append(new_elem)

                if etype == 'table':
                    # Nach Tabellen im LEK immer eine Leerzeile
                    doc.add_paragraph()
            except Exception as e:
                # Fallback: Fehlermeldung als Absatz einfügen
                doc.add_paragraph(f"[Element konnte nicht kopiert werden: {str(e)[:80]}]")

    def _is_landscape_section(self, section):
        """Prüft, ob eine Sektion im Querformat ist."""
        try:
            if getattr(section, 'orientation', None) == WD_ORIENT.LANDSCAPE:
                return True
            return int(section.page_width) > int(section.page_height)
        except Exception:
            return False

    def _detect_external_document_orientation(self, source_doc, override='auto'):
        """Ermittelt die geeignetste Orientierung für ein extern referenziertes Dokument."""
        choice = str(override or 'auto').strip().lower()
        if choice in {'landscape', 'portrait'}:
            return choice

        try:
            if any(self._is_landscape_section(section) for section in source_doc.sections):
                return 'landscape'
        except Exception:
            pass

        max_columns = 0
        max_row_text = 0
        for table in source_doc.tables:
            try:
                max_columns = max(max_columns, len(table.columns) if table.columns else 0)
                for row in table.rows:
                    max_row_text = max(
                        max_row_text,
                        sum(len(str(cell.text or '').strip()) for cell in row.cells),
                    )
            except Exception:
                continue

        if max_columns >= 6 or max_row_text >= 150:
            return 'landscape'

        return 'portrait'

    def _apply_section_orientation(self, section, orientation):
        """Setzt Hoch-/Querformat robust auf eine Sektion."""
        desired = str(orientation or 'portrait').strip().lower()
        try:
            current_width = section.page_width
            current_height = section.page_height
            if desired == 'landscape':
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = max(current_width, current_height)
                section.page_height = min(current_width, current_height)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = min(current_width, current_height)
                section.page_height = max(current_width, current_height)
        except Exception:
            pass

    def _collect_used_style_and_num_ids_from_xml_elements(self, xml_elements):
        """Sammelt verwendete Style- und NumIDs direkt aus XML-Elementen."""
        style_ids = set()
        num_ids = set()

        for xml_elem in xml_elements or []:
            if xml_elem is None:
                continue
            for n in xml_elem.xpath('.//w:pStyle|.//w:rStyle|.//w:tblStyle'):
                val = n.get(qn('w:val'))
                if val:
                    style_ids.add(val)
            for n in xml_elem.xpath('.//w:numPr/w:numId'):
                val = n.get(qn('w:val'))
                if val:
                    num_ids.add(val)

        return style_ids, num_ids

    def _append_external_document_content(self, target_doc, external_doc):
        """Übernimmt die relevanten Body-Elemente eines externen Dokuments ins Ziel."""
        xml_elements = []
        for child in list(external_doc._element.body):
            tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_local in ('p', 'tbl', 'sdt'):
                xml_elements.append(child)

        used_style_ids, used_num_ids = self._collect_used_style_and_num_ids_from_xml_elements(xml_elements)
        self._merge_required_styles(external_doc.part, target_doc.part, used_style_ids)
        external_num_map = self._merge_required_numbering(external_doc.part, target_doc.part, used_num_ids)

        body = target_doc._element.body
        sect_pr = body.find(qn('w:sectPr'))
        rel_map = {}

        for xml_elem in xml_elements:
            cloned = deepcopy(xml_elem)
            self._remap_relationship_ids_in_xml(
                xml_element=cloned,
                source_part=external_doc.part,
                target_part=target_doc.part,
                rel_map=rel_map,
            )
            self._remap_num_ids_in_element(cloned, external_num_map)

            if sect_pr is not None:
                body.insert(list(body).index(sect_pr), cloned)
            else:
                body.append(cloned)

    def _insert_external_table_reference(self, doc, task):
        """Fügt eine referenzierte externe Tabellen-/Dokumentdatei in die LEK ein."""
        path = str((task or {}).get('external_table_path') or '').strip()
        if not path or not Path(path).exists():
            return False

        source_doc = Document(path)
        orientation = self._detect_external_document_orientation(
            source_doc,
            override=(task or {}).get('external_table_orientation', 'auto'),
        )

        inserted_landscape = False
        if orientation == 'landscape':
            landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
            self._apply_section_orientation(landscape_section, 'landscape')
            inserted_landscape = True

        self._append_external_document_content(doc, source_doc)

        if inserted_landscape:
            portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
            self._apply_section_orientation(portrait_section, 'portrait')

        return True

    def _copy_paragraphs_for_lek(self, doc, paragraphs):
        """
        Kopiert Paragraphen für LEK-Erstellung mit angepassten Überschrift-Levels
        
        Args:
            doc: Ziel-Word-Dokument (LEK)
            paragraphs: Liste von Quell-Paragraphen
        """
        for paragraph in paragraphs:
            if not paragraph.text.strip():
                # Leere Absätze als solche übernehmen
                doc.add_paragraph()
                continue
            
            # Prüfe ob es eine Überschrift ist
            if (paragraph.style and 
                (paragraph.style.name.startswith('Heading') or 
                 paragraph.style.name.startswith('Überschrift') or
                 'Heading' in paragraph.style.name or
                 'Überschrift' in paragraph.style.name)):
                
                # Extrahiere ursprüngliches Level und behalte es bei (KEIN Shift zu Level 2!)
                original_level = self._get_heading_level(paragraph.style.name)
                
                # Erstelle Heading mit ORIGINALEM Level
                new_paragraph = doc.add_heading(paragraph.text, level=original_level)
                
            else:
                # Normaler Absatz - kopiere mit bestehender Logik
                new_paragraph = doc.add_paragraph()
                
                # Absatzformatierung kopieren (vereinfacht für LEK)
                try:
                    if paragraph.style and not paragraph.style.name.startswith('Heading') and not paragraph.style.name.startswith('Überschrift'):
                        try:
                            new_paragraph.style = paragraph.style
                        except:
                            pass
                except:
                    pass
                
                # Runs mit Formatierung kopieren
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    try:
                        if run.font:
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                    except:
                        pass

    def _copy_table_with_formatting(self, doc, source_table):
        """
        Kopiert eine Tabelle mit vollständiger Formatierung
        
        Args:
            doc: Ziel-Word-Dokument
            source_table: Quell-Tabelle
        """
        try:
            # Erstelle neue Tabelle mit gleicher Struktur
            rows = len(source_table.rows)
            cols = len(source_table.columns) if source_table.columns else 1
            
            new_table = doc.add_table(rows=rows, cols=cols)
            
            # Kopiere Zellinhalt und Formatierung
            for row_idx, source_row in enumerate(source_table.rows):
                target_row = new_table.rows[row_idx]
                
                for cell_idx, source_cell in enumerate(source_row.cells):
                    if cell_idx < len(target_row.cells):
                        target_cell = target_row.cells[cell_idx]
                        
                        # Kopiere Text-Inhalt
                        target_cell.text = source_cell.text
                        
                        # Kopiere Zellenformatierung (Basic)
                        try:
                            # Paragraph-Formatierung in der Zelle übernehmen
                            if source_cell.paragraphs:
                                source_para = source_cell.paragraphs[0]
                                target_para = target_cell.paragraphs[0]
                                
                                # Alignment kopieren
                                if source_para.paragraph_format.alignment:
                                    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
                                
                                # Text-Formatierung kopieren
                                if source_para.runs:
                                    # Lösche bestehenden Text
                                    target_para.clear()
                                    for run in source_para.runs:
                                        new_run = target_para.add_run(run.text)
                                        if run.bold:
                                            new_run.bold = True
                                        if run.italic:
                                            new_run.italic = True
                                        if run.underline:
                                            new_run.underline = True
                                        if run.font.size:
                                            new_run.font.size = run.font.size
                        except:
                            # Formatierung fehlgeschlagen, aber Text ist da
                            pass
            
        except Exception as e:
            # Fallback: Einfache Tabellen-Ersetzung
            doc.add_paragraph(f"[Tabelle mit {len(source_table.rows)} Zeilen konnte nicht vollständig kopiert werden]")

    def _copy_paragraphs_with_formatting(self, doc, paragraphs):
        """
        Kopiert Paragraphen mit ihrer ursprünglichen Formatierung
        
        Args:
            doc: Ziel-Word-Dokument
            paragraphs: Liste von Quell-Paragraphen
        """
        for paragraph in paragraphs:
            if not paragraph.text.strip():
                # Leere Absätze als solche übernehmen
                doc.add_paragraph()
                continue
            
            # Neuen Absatz erstellen - prüfe zuerst ob es eine Überschrift ist
            if (paragraph.style and 
                (paragraph.style.name.startswith('Heading') or 
                 paragraph.style.name.startswith('Überschrift') or
                 'Heading' in paragraph.style.name or
                 'Überschrift' in paragraph.style.name)):
                
                # Extrahiere Heading-Level präzise basierend auf Style-Namen
                level = self._get_heading_level(paragraph.style.name)
                
                # Erstelle Heading mit Text
                new_paragraph = doc.add_heading(paragraph.text, level=level)
                
            else:
                # Normaler Absatz
                new_paragraph = doc.add_paragraph()
            
            # Absatzformatierung kopieren
            try:
                if paragraph.style:
                    # Versuche den Style zu übernehmen, außer bei Überschriften
                    if not paragraph.style.name.startswith('Heading') and not paragraph.style.name.startswith('Überschrift'):
                        try:
                            new_paragraph.style = paragraph.style
                        except:
                            # Falls Style nicht verfügbar, verwende Normal
                            pass
                
                # Absatz-Formatierung übernehmen
                if paragraph.paragraph_format:
                    try:
                        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
                        new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
                        new_paragraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
                        new_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
                        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
                        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
                        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
                        new_paragraph.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
                    except:
                        # Falls Formatierung nicht kopiert werden kann, ignorieren
                        pass
            except:
                # Falls Style-Zugriff fehlschlägt, ignorieren
                pass
            
            # Runs mit Formatierung kopieren (nur wenn es kein Heading ist, da dort Text schon gesetzt wurde)
            if not (paragraph.style and 
                    (paragraph.style.name.startswith('Heading') or 
                     paragraph.style.name.startswith('Überschrift') or
                     'Heading' in paragraph.style.name or
                     'Überschrift' in paragraph.style.name)):
                
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    
                    try:
                        # Font-Formatierung kopieren
                        if run.font:
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.color and run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                            
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            
                            # Weitere Font-Eigenschaften
                            try:
                                new_run.font.subscript = run.font.subscript
                                new_run.font.superscript = run.font.superscript
                                new_run.font.strike = run.font.strike
                                new_run.font.small_caps = run.font.small_caps
                                new_run.font.all_caps = run.font.all_caps
                            except:
                                # Ignoriere erweiterte Formatierungen falls nicht verfügbar
                                pass
                    except:
                        # Falls Font-Formatierung nicht kopiert werden kann, ignorieren
                        pass

    def _copy_single_paragraph_with_formatting(self, doc, paragraph):
        """Kopiert einen einzelnen Paragraph mit Grundformatierung in das Zieldokument."""
        if paragraph is None:
            doc.add_paragraph()
            return

        text = str(paragraph.text or '')
        if not text.strip() and not paragraph.runs:
            doc.add_paragraph()
            return

        new_paragraph = doc.add_paragraph()

        try:
            if paragraph.style and not (
                paragraph.style.name.startswith('Heading')
                or paragraph.style.name.startswith('Überschrift')
            ):
                try:
                    new_paragraph.style = paragraph.style
                except:
                    pass

            if paragraph.paragraph_format:
                try:
                    new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
                    new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
                    new_paragraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
                    new_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
                    new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
                    new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
                    new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
                    new_paragraph.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
                except:
                    pass
        except:
            pass

        if paragraph.runs:
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                try:
                    if run.font:
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                except:
                    pass
        elif text:
            new_paragraph.add_run(text)

    def _is_structured_task_table_for_export(self, table):
        """Prüft, ob eine Tabelle einer strukturierten Aufgaben-Tabelle entspricht."""
        if table is None:
            return False

        keys = set()
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            keys.add(self._normalize_table_key(row.cells[0].text))

        return 'id' in keys and ('aufgabenstellungpflicht' in keys or 'aufgabenstellung' in keys)

    def _structured_table_value_cell_by_aliases(self, table, aliases):
        """Liefert die rechte Wertzelle für den ersten passenden Alias-Key."""
        alias_set = {self._normalize_table_key(a) for a in aliases if str(a).strip()}
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = self._normalize_table_key(row.cells[0].text)
            if key in alias_set:
                return row.cells[1]
        return None

    def _append_cell_as_flow_text(self, doc, cell, fallback_text='', task=None):
        """Überträgt den Inhalt einer Tabellenzelle als Fließtext-/Blockelemente ins Dokument."""
        written = False
        if cell is not None:
            source_part = getattr(cell, 'part', None)
            written = self._append_cell_block_elements_for_lek(doc, cell, source_part=source_part, task=task)

        if not written and str(fallback_text or '').strip():
            fallback_paragraph = doc.add_paragraph(str(fallback_text).strip())
            self._set_paragraph_keep_rules(fallback_paragraph, keep_with_next=True, keep_together=True)
            written = True

        return written

    def _append_cell_block_elements_for_lek(self, doc, cell, source_part=None, task=None):
        """Kopiert Blockelemente einer Tabellenzelle (p/tbl/sdt) in ein LEK-Dokument."""
        if cell is None:
            return False

        tc = cell._tc
        body = doc._element.body
        sect_pr = body.find(qn('w:sectPr'))
        rel_map = {}
        appended_any = False

        for child in list(tc):
            tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_local == 'tcPr':
                continue

            if tag_local == 'p':
                paragraph_text = ''.join((node.text or '') for node in child.xpath('.//w:t')).strip()
                marker = self._parse_external_table_reference_line(paragraph_text)
                if marker and marker.get('type') == 'table':
                    appended_any = self._insert_external_table_reference(doc, task) or appended_any
                    continue
                if marker and marker.get('type') == 'orientation':
                    continue
            elif tag_local not in ('tbl', 'sdt'):
                continue

            cloned = deepcopy(child)

            if source_part is not None:
                self._remap_relationship_ids_in_xml(
                    xml_element=cloned,
                    source_part=source_part,
                    target_part=doc.part,
                    rel_map=rel_map,
                )

            self._remap_num_ids_in_element(cloned, self._num_id_map)

            if tag_local == 'p':
                self._set_keep_rules_on_paragraph_xml(cloned, keep_next=True, keep_lines=True)
            elif tag_local == 'sdt':
                for node in cloned.iter():
                    node_tag = node.tag.split('}')[-1] if '}' in node.tag else node.tag
                    if node_tag == 'p':
                        self._set_keep_rules_on_paragraph_xml(node, keep_next=True, keep_lines=True)

            if tag_local == 'tbl':
                # Vor Tabellen immer eine Leerzeile
                doc.add_paragraph()
                self._set_table_full_width_xml(cloned)

            if sect_pr is not None:
                body.insert(list(body).index(sect_pr), cloned)
            else:
                body.append(cloned)

            if tag_local == 'tbl':
                # Nach Tabellen immer eine Leerzeile
                doc.add_paragraph()

            appended_any = True

        return appended_any

    def _append_structured_task_as_flow_text(self, doc, task, table, include_title=True):
        """Schreibt strukturierte Aufgaben als Fließtext (ohne Tabelle) in fester Reihenfolge."""
        section_cells = []
        for section in self._task_flow_sections():
            section_name = str(section.get('name', '')).strip().lower()
            if not include_title and section_name == 'title':
                continue
            # Punkte nur im Titel (Heading), nicht zusätzlich im Fließtext ausgeben
            if section_name == 'points':
                continue

            aliases = []
            for field_key in section.get('fields', []):
                aliases.extend(self._field_aliases(field_key))

            fallback_text = task.get('title', '') if section.get('name') == 'title' else ''
            section_cells.append((
                section.get('name', ''),
                bool(section.get('optional', False)),
                self._structured_table_value_cell_by_aliases(table, aliases),
                fallback_text,
            ))

        blocks_written = 0
        for _name, optional_block, cell, fallback in section_cells:
            has_content = self._append_cell_as_flow_text(doc, cell, fallback_text=fallback, task=task)
            if not has_content and optional_block:
                continue

            if has_content:
                blocks_written += 1
                # Leerzeile zwischen allen geschriebenen Blöcken
                doc.add_paragraph()

        if self._insert_external_table_reference(doc, task):
            blocks_written += 1
            doc.add_paragraph()

        if blocks_written == 0:
            # Fallback: zumindest Titel/Content aus Task übernehmen
            fallback_title = str(task.get('title') or '').strip()
            if fallback_title:
                doc.add_paragraph(fallback_title)
                doc.add_paragraph()

            for line in task.get('content', []) or []:
                line_text = str(line or '').strip()
                if line_text:
                    doc.add_paragraph(line_text)

    def _strip_leading_task_heading_element(self, all_elements, task_title=''):
        """Entfernt einen führenden Aufgaben-Heading-Absatz aus all_elements (falls vorhanden)."""
        elements = list(all_elements or [])
        if not elements:
            return elements

        first = elements[0]
        if first.get('type') != 'paragraph':
            return elements

        para = first.get('content')
        if para is None:
            return elements

        text = str(getattr(para, 'text', '') or '').strip()
        title_norm = str(task_title or '').strip().lower()
        extracted_norm = self._extract_task_title(text).strip().lower() if text else ''

        try:
            is_heading = self._is_heading1(para) or self._is_heading2(para)
        except Exception:
            is_heading = False

        if is_heading or (title_norm and extracted_norm == title_norm):
            return elements[1:]
        return elements

    def _append_task_elements_for_lek(self, doc, task, elements):
        """Kopiert Aufgabenelemente und ersetzt Tabellenmarker durch externe Tabellenblöcke."""
        elements = list(elements or [])
        if not elements:
            return False

        body = doc._element.body
        sect_pr = body.find(qn('w:sectPr'))
        rel_map = {}
        wrote_any = False
        inserted_external = False

        for element in elements:
            try:
                etype = element.get('type')
                content = element.get('content')

                if etype == 'paragraph':
                    paragraph_text = str(getattr(content, 'text', '') or '').strip()
                    marker = self._parse_external_table_reference_line(paragraph_text)
                    if marker and marker.get('type') == 'table':
                        inserted_external = self._insert_external_table_reference(doc, task) or inserted_external
                        wrote_any = inserted_external or wrote_any
                        continue
                    if marker and marker.get('type') == 'orientation':
                        continue
                    xml_elem = content._element
                elif etype == 'table':
                    xml_elem = content._element
                elif etype == 'sdt':
                    xml_elem = content
                else:
                    continue

                cloned = deepcopy(xml_elem)

                if hasattr(content, 'part'):
                    self._remap_relationship_ids_in_xml(
                        xml_element=cloned,
                        source_part=content.part,
                        target_part=doc.part,
                        rel_map=rel_map,
                    )

                self._remap_num_ids_in_element(cloned, self._num_id_map)

                if etype == 'paragraph':
                    self._set_keep_rules_on_paragraph_xml(cloned, keep_next=True, keep_lines=True)
                elif etype == 'sdt':
                    for node in cloned.iter():
                        node_tag = node.tag.split('}')[-1] if '}' in node.tag else node.tag
                        if node_tag == 'p':
                            self._set_keep_rules_on_paragraph_xml(node, keep_next=True, keep_lines=True)

                if etype == 'table':
                    doc.add_paragraph()
                    self._set_table_full_width_xml(cloned)

                if sect_pr is not None:
                    body.insert(list(body).index(sect_pr), cloned)
                else:
                    body.append(cloned)

                if etype == 'table':
                    doc.add_paragraph()

                wrote_any = True
            except Exception as e:
                doc.add_paragraph(f"[Element konnte nicht kopiert werden: {str(e)[:80]}]")
                wrote_any = True

        if not inserted_external:
            inserted_external = self._insert_external_table_reference(doc, task)
            wrote_any = inserted_external or wrote_any

        return wrote_any

    def _extract_points_text(self, task, table=None):
        """Ermittelt die Punkteangabe einer Aufgabe (z. B. '10')."""
        points_raw = ''

        if table is not None:
            aliases = self._field_aliases('punkte', fallback_aliases=['punktzahl', 'bewertung'])
            cell = self._structured_table_value_cell_by_aliases(table, aliases)
            lines = self._cell_text_lines(cell)
            if lines:
                points_raw = str(lines[0] or '').strip()

        if not points_raw:
            for line in (task.get('content', []) or []):
                text = str(line or '').strip()
                if not text:
                    continue
                match = re.match(r'^\s*punkte\s*:\s*(.+)$', text, flags=re.IGNORECASE)
                if match:
                    points_raw = str(match.group(1) or '').strip()
                    break

        points_raw = str(points_raw or '').strip()
        if not points_raw:
            return ''

        # Das Wort "Punkte"/"Punkt"/"Pkt." soll entfallen; nur die Anzahl bleibt.
        normalized = re.sub(r'\b(punkte?|pkt\.?)\b', '', points_raw, flags=re.IGNORECASE)
        normalized = re.sub(r'\s+', ' ', normalized).strip(' :-;,.')
        normalized = normalized.strip()

        return normalized or points_raw

    def _format_points_label(self, points_text, min_inner_width=None):
        """Formatiert die Punkte-Box mit Innenabstand und Mindestbreite."""
        text = str(points_text or '').strip()
        if not text:
            return ''

        if min_inner_width is None:
            configured_width = self.get_import_rule('export_rules.title_points_box.min_inner_width', 12)
        else:
            configured_width = min_inner_width

        configured_padding = self.get_import_rule('export_rules.title_points_box.padding_spaces', 1)

        try:
            min_width = max(1, int(configured_width))
        except Exception:
            min_width = 12

        try:
            padding_spaces = max(0, int(configured_padding))
        except Exception:
            padding_spaces = 1

        inner_width = max(min_width, len(text))
        padded = text.ljust(inner_width)
        pad = ' ' * padding_spaces
        return f"{pad}{padded}{pad}"

    def _apply_run_border(self, run):
        """Fügt einem Run einen einfachen Rahmen hinzu."""
        try:
            r = run._r
            rPr = r.get_or_add_rPr()
            bdr = OxmlElement('w:bdr')
            bdr.set(qn('w:val'), 'single')
            bdr.set(qn('w:sz'), '8')
            bdr.set(qn('w:space'), '2')
            bdr.set(qn('w:color'), '000000')
            rPr.append(bdr)
        except Exception:
            pass

    def _set_paragraph_keep_rules(self, paragraph, keep_with_next=False, keep_together=True):
        """Setzt Absatzregeln, um Seitenumbrüche innerhalb einer Aufgabe zu reduzieren."""
        try:
            paragraph.paragraph_format.keep_with_next = bool(keep_with_next)
            if keep_together:
                paragraph.paragraph_format.keep_together = True
        except Exception:
            pass

    def _set_keep_rules_on_paragraph_xml(self, paragraph_xml, keep_next=False, keep_lines=True):
        """Setzt Keep-Regeln direkt auf ein w:p-Element."""
        if paragraph_xml is None:
            return

        try:
            p_pr = paragraph_xml.find(qn('w:pPr'))
            if p_pr is None:
                p_pr = OxmlElement('w:pPr')
                paragraph_xml.insert(0, p_pr)

            keep_next_node = p_pr.find(qn('w:keepNext'))
            if keep_next:
                if keep_next_node is None:
                    keep_next_node = OxmlElement('w:keepNext')
                    p_pr.append(keep_next_node)
            elif keep_next_node is not None:
                p_pr.remove(keep_next_node)

            keep_lines_node = p_pr.find(qn('w:keepLines'))
            if keep_lines:
                if keep_lines_node is None:
                    keep_lines_node = OxmlElement('w:keepLines')
                    p_pr.append(keep_lines_node)
            elif keep_lines_node is not None:
                p_pr.remove(keep_lines_node)
        except Exception:
            pass

    def _append_task_heading_with_points(self, doc, task, table=None):
        """Schreibt den Aufgabentitel als Heading 1 und Punkte rechtsbündig mit Rahmen."""
        title_text = str(task.get('title') or '').strip() or 'Ohne Titel'
        paragraph = doc.add_heading(level=1)
        paragraph.add_run(title_text)
        self._set_paragraph_keep_rules(paragraph, keep_with_next=True, keep_together=True)

        points_text = self._extract_points_text(task, table=table)
        if not points_text:
            return

        points_label = self._format_points_label(points_text)
        if not points_label:
            return

        try:
            section = doc.sections[-1]
            right_pos = section.page_width - section.left_margin - section.right_margin
            paragraph.paragraph_format.tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT)
        except Exception:
            pass

        paragraph.add_run('\t')
        points_run = paragraph.add_run(points_label)
        points_run.font.name = 'Aptos'
        points_run.bold = True
        self._apply_run_border(points_run)

        return paragraph

    def append_task_to_lek_document(self, doc, task):
        """Fügt eine Aufgabe als Heading1 (+ Punkte) und Inhalt in ein LEK-Dokument ein."""
        all_elements = task.get('all_elements') or []
        table = None
        if len(all_elements) == 1 and all_elements[0].get('type') == 'table':
            table = all_elements[0].get('content')

        self._append_task_heading_with_points(doc, task, table=table)
        # Nach Überschrift 1 immer eine Leerzeile einfügen.
        spacer = doc.add_paragraph()
        self._set_paragraph_keep_rules(spacer, keep_with_next=True, keep_together=True)
        self.append_task_content_for_lek(doc, task, include_title=False)

    def _set_table_full_width_xml(self, table_xml_element):
        """Setzt eine Tabellenbreite auf 100 % (OOXML: w:type='pct', w:w='5000')."""
        if table_xml_element is None:
            return

        try:
            tbl_pr = table_xml_element.find(qn('w:tblPr'))
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                table_xml_element.insert(0, tbl_pr)

            tbl_w = tbl_pr.find(qn('w:tblW'))
            if tbl_w is None:
                tbl_w = OxmlElement('w:tblW')
                tbl_pr.append(tbl_w)

            tbl_w.set(qn('w:type'), 'pct')
            tbl_w.set(qn('w:w'), '5000')
        except Exception:
            pass

    def append_task_content_for_lek(self, doc, task, include_title=True):
        """Fügt Aufgabeninhalt in LEK ein; strukturierte Tabellen als Fließtext ohne Tabellenlayout."""
        all_elements = task.get('all_elements') or []
        if len(all_elements) == 1 and all_elements[0].get('type') == 'table':
            table = all_elements[0].get('content')
            if self._is_structured_task_table_for_export(table):
                self._append_structured_task_as_flow_text(doc, task, table, include_title=include_title)
                return

        if all_elements:
            elements_to_copy = all_elements
            if not include_title:
                elements_to_copy = self._strip_leading_task_heading_element(all_elements, task.get('title', ''))

            if elements_to_copy:
                self._append_task_elements_for_lek(doc, task, elements_to_copy)
            return

        original_paragraphs = task.get('original_paragraphs') or []
        if original_paragraphs:
            self._copy_paragraphs_for_lek(doc, original_paragraphs)
            self._insert_external_table_reference(doc, task)
            return

        for content_line in task.get('content', []) or []:
            if str(content_line).strip():
                doc.add_paragraph(str(content_line))

        self._insert_external_table_reference(doc, task)

    def _cell_text_lines(self, cell):
        """Liest nicht-leere Textzeilen aus einer Tabellenzelle in stabiler Reihenfolge."""
        lines = []
        if cell is None:
            return lines

        for paragraph in cell.paragraphs:
            text = str(paragraph.text or '').strip()
            if text:
                lines.append(text)

        if not lines:
            raw = str(cell.text or '').strip()
            if raw:
                lines = [line.strip() for line in raw.splitlines() if line.strip()]

        return [line for line in lines if not self._parse_external_table_reference_line(line)]

    def build_task_flow_preview_lines(self, task):
        """
        Erzeugt eine textuelle Vorschau einer Aufgabe in derselben Reihenfolge wie der LEK-Export.

        Returns:
            list[str]: Vorschauzeilen für GUI-Preview/Review.
        """
        task = task or {}
        all_elements = task.get('all_elements') or []

        if len(all_elements) == 1 and all_elements[0].get('type') == 'table':
            table = all_elements[0].get('content')
            if self._is_structured_task_table_for_export(table):
                section_defs = []
                for section in self._task_flow_sections():
                    aliases = []
                    for field_key in section.get('fields', []):
                        aliases.extend(self._field_aliases(field_key))

                    fallback = task.get('title', '') if section.get('name') == 'title' else ''
                    section_defs.append((
                        section.get('label', ''),
                        aliases,
                        fallback,
                        bool(section.get('optional', False)),
                    ))

                lines = []
                written_sections = 0
                for label, aliases, fallback, optional in section_defs:
                    cell = self._structured_table_value_cell_by_aliases(table, aliases)
                    section_lines = self._cell_text_lines(cell)

                    if not section_lines and str(fallback or '').strip():
                        section_lines = [str(fallback).strip()]

                    if optional and not section_lines:
                        continue

                    if not section_lines:
                        continue

                    if written_sections > 0:
                        lines.append('')

                    lines.append(f"{label}:")
                    lines.extend(section_lines)
                    written_sections += 1

                if lines:
                    return lines

        # Fallback für nicht-strukturierte Aufgaben
        content_lines = [str(line).strip() for line in (task.get('content') or []) if str(line).strip()]
        external_table_label = self._task_external_table_label(task)
        if external_table_label:
            if content_lines:
                content_lines.append('')
            content_lines.append(external_table_label)
        if content_lines:
            return content_lines

        title = str(task.get('title') or '').strip()
        lines = [title] if title else ['(kein Inhalt verfügbar)']
        if external_table_label:
            lines.extend(['', external_table_label])
        return lines

    def analyze_task_flow_preview_delta(self, task):
        """
        Analysiert, welche optionale Blöcke in der strukturierten Exportreihenfolge fehlen.

        Returns:
            dict: {
                'is_structured': bool,
                'missing_optional_sections': list[str],
            }
        """
        task = task or {}
        all_elements = task.get('all_elements') or []

        result = {
            'is_structured': False,
            'missing_optional_sections': [],
        }

        if not (len(all_elements) == 1 and all_elements[0].get('type') == 'table'):
            return result

        table = all_elements[0].get('content')
        if not self._is_structured_task_table_for_export(table):
            return result

        result['is_structured'] = True

        if not bool(self.get_import_rule('preview_rules.show_optional_missing_sections', True)):
            return result

        optional_label_filter = {
            str(v).strip()
            for v in (self.get_import_rule('preview_rules.optional_sections', []) or [])
            if str(v).strip()
        }

        optional_sections = []
        for section in self._task_flow_sections():
            if not bool(section.get('optional', False)):
                continue

            label = str(section.get('label', '')).strip()
            if optional_label_filter and label not in optional_label_filter:
                continue

            aliases = []
            for field_key in section.get('fields', []):
                aliases.extend(self._field_aliases(field_key))

            optional_sections.append((label, aliases))

        missing = []
        for label, aliases in optional_sections:
            cell = self._structured_table_value_cell_by_aliases(table, aliases)
            if not self._cell_text_lines(cell):
                missing.append(label)

        result['missing_optional_sections'] = missing
        return result
    
    def _is_task_start(self, text, paragraph):
        """
        Erkennt, ob ein Text der Beginn einer neuen Aufgabe ist
        Priorisiert Überschrift 1 Format, fallback auf Textmuster
        
        Args:
            text (str): Zu prüfender Text
            paragraph: Word-Paragraph-Objekt
            
        Returns:
            bool: True wenn Aufgabenbeginn erkannt
        """
        # Primäre Erkennung: Überschrift 1 Formatvorlage
        try:
            if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Überschrift 1':
                return True
        except:
            # Falls Style-Zugriff fehlschlägt, weiter mit Textmustern
            pass
        
        # Sekundäre Erkennung: Verschiedene Textmuster für Aufgabenbeginne
        return self._matches_task_start_text(text)

    def _matches_task_start_text(self, text):
        """
        Prüft rein textbasiert, ob ein String wie ein Aufgabenstart aussieht.
        """
        patterns = [
            r'^Aufgabe\s*\d+',
            r'^\d+\.',
            r'^Übung\s*\d+',
            r'^Frage\s*\d+',
            r'^Problem\s*\d+',
            r'^Nr\.\s*\d+',
            r'^\w\)',  # a), b), etc.
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        return False

    def _extract_task_title(self, text):
        """
        Extrahiert den Titel einer Aufgabe
        
        Args:
            text (str): Text der ersten Zeile der Aufgabe
            
        Returns:
            str: Extrahierter Titel
        """
        # Entfernt Nummerierung und behält den Rest
        cleaned = re.sub(
            r'^(?:Aufgabe\s*)?\d+(?:[\.,]\d+)?\s*(?:[:\.]\s*|\)\s*|-\s*|$)|^(?:Übung\s*)?\d+(?:[\.,]\d+)?\s*(?:[:\.]\s*|\)\s*|-\s*|$)',
            '',
            text,
            flags=re.IGNORECASE,
        )
        
        # Falls der Text sehr lang ist, nur die ersten 50 Zeichen nehmen
        if len(cleaned) > 50:
            cleaned = cleaned[:47] + "..."
        
        return cleaned.strip() or "Ohne Titel"
    
    def _extract_difficulty(self, text):
        """
        Versucht den Schwierigkeitsgrad aus dem Text zu extrahieren
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            str: Schwierigkeitsgrad (Leicht, Mittel, Schwer)
        """
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['leicht', 'einfach', 'basic', 'anfänger']):
            return 'Leicht'
        elif any(word in text_lower for word in ['schwer', 'difficult', 'komplex', 'advanced']):
            return 'Schwer'
        elif any(word in text_lower for word in ['mittel', 'medium', 'intermediate']):
            return 'Mittel'
        
        # Standard-Schwierigkeitsgrad basierend auf Textlänge
        if len(text) > 200:
            return 'Schwer'
        elif len(text) > 100:
            return 'Mittel'
        else:
            return 'Leicht'
    
    def _extract_keywords(self, text):
        """
        Extrahiert relevante Schlüsselwörter aus dem Text
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            list: Liste von Schlüsselwörtern
        """
        # Stopwörter, die ignoriert werden sollen
        stop_words = {
            'der', 'die', 'das', 'ein', 'eine', 'und', 'oder', 'aber', 'von', 'zu', 'mit', 'bei', 'für',
            'ist', 'sind', 'war', 'waren', 'hat', 'haben', 'wird', 'werden', 'kann', 'können', 'soll',
            'sollte', 'auf', 'in', 'an', 'über', 'unter', 'durch', 'gegen', 'ohne', 'um', 'nach', 'vor'
        }
        
        # Text in Wörter aufteilen und bereinigen
        words = re.findall(r'\b\w+\b', text.lower())
        
        # Filtern: mindestens 3 Zeichen, keine Stopwörter, keine Zahlen
        keywords = [
            word for word in words 
            if len(word) >= 3 
            and word not in stop_words 
            and not word.isdigit()
        ]
        
        # Nur die häufigsten/relevantesten Keywords zurückgeben
        return list(set(keywords))[:10]
    
    def _extract_explicit_keywords(self, text):
        """
        Extrahiert explizite Schlüsselwörter aus "Schlüsselwörter:" Zeilen
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            list: Liste von explizit angegebenen Schlüsselwörtern oder leere Liste
        """
        import re
        
        # Suche nach "Schlüsselwörter:" gefolgt von kommagetrenner Liste
        patterns = [
            r'Schlüsselwörter:\s*([^\n\r]+)',
            r'Schlüsselwort:\s*([^\n\r]+)',
            r'Keywords:\s*([^\n\r]+)',
            r'Stichwörter:\s*([^\n\r]+)',
            r'Stichworte:\s*([^\n\r]+)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Nehme den ersten/längsten Treffer
                keywords_text = max(matches, key=len) if len(matches) > 1 else matches[0]
                
                # Trenne an Kommas und bereinige
                keywords = []
                for keyword in keywords_text.split(','):
                    cleaned = keyword.strip()
                    if cleaned and len(cleaned) >= 2:  # Mindestens 2 Zeichen
                        keywords.append(cleaned.lower())
                
                if keywords:
                    return keywords[:10]  # Maximal 10 Keywords
        
        return []  # Keine expliziten Keywords gefunden
    
    def _extract_explicit_difficulty(self, text):
        """
        Extrahiert explizite Schwierigkeitsangaben aus "Schwierigkeit:" Zeilen
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            str: Explizit angegebene Schwierigkeit oder None
        """
        import re
        
        # Suche nach "Schwierigkeit:" gefolgt von Schwierigkeitsgrad
        patterns = [
            r'Schwierigkeit:\s*([^\n\r]+)',
            r'Schwierigkeitsgrad:\s*([^\n\r]+)',
            r'Difficulty:\s*([^\n\r]+)',
            r'Level:\s*([^\n\r]+)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                difficulty_text = matches[0].strip().lower()
                normalized = self._normalize_difficulty_token(difficulty_text)
                if normalized:
                    return normalized.capitalize()
        
        return None  # Keine explizite Schwierigkeit gefunden

    def _difficulty_allowed_values(self):
        """Liefert erlaubte Schwierigkeitsgrade in Kleinbuchstaben."""
        values = self.get_import_rule('difficulty_rules.allowed_values', ['leicht', 'mittel', 'schwer']) or []
        normalized = [str(v).strip().lower() for v in values if str(v).strip()]
        return normalized or ['leicht', 'mittel', 'schwer']

    def _difficulty_aliases(self):
        """Liefert Alias-Mapping -> kanonischer Schwierigkeitsgrad (lowercase)."""
        defaults = {
            'easy': 'leicht',
            'einfach': 'leicht',
            'medium': 'mittel',
            'normal': 'mittel',
            'hard': 'schwer',
            'difficult': 'schwer',
            'komplex': 'schwer',
        }
        configured = self.get_import_rule('difficulty_rules.aliases', {}) or {}
        merged = dict(defaults)
        for key, value in configured.items():
            k = str(key).strip().lower()
            v = str(value).strip().lower()
            if k and v:
                merged[k] = v
        return merged

    def _normalize_difficulty_token(self, value):
        """Normalisiert einen Schwierigkeitsgrad-Token auf den kanonischen Lowercase-Wert."""
        val = str(value or '').strip().lower()
        if not val:
            return None

        allowed = set(self._difficulty_allowed_values())
        if val in allowed:
            return val

        aliases = self._difficulty_aliases()
        mapped = aliases.get(val)
        if mapped in allowed:
            return mapped

        return None

    def _has_inconsistent_difficulty(self, difficulty_raw):
        """Prüft, ob ein Rohwert mehrere Schwierigkeitsgrade gleichzeitig enthält."""
        if not difficulty_raw:
            return False

        value = str(difficulty_raw).lower()
        canonical_hits = set()

        allowed = self._difficulty_allowed_values()
        aliases = self._difficulty_aliases()

        token_map = {token: token for token in allowed}
        token_map.update(aliases)

        for token, canonical in token_map.items():
            pattern = r'\b' + re.escape(str(token).lower()) + r'\b'
            if re.search(pattern, value):
                normalized = self._normalize_difficulty_token(canonical)
                if normalized:
                    canonical_hits.add(normalized)

        return len(canonical_hits) > 1
    
    def create_document_from_tasks(self, tasks, output_path, lek_theme="", debug_context_report=False):
        """
        Erstellt ein neues Word-Dokument aus einer Liste von Aufgaben
        Verwendet Vorlagen aus dem Vorlagen-Ordner falls verfügbar
        
        Args:
            tasks (list): Liste von Aufgaben-Dictionaries
            output_path (str): Pfad für die neue Datei
            lek_theme (str): LEK-Thema für Vorlagen-Auswahl
            debug_context_report (bool): Gibt Kontext-Report zu Styles/Nummerierung aus
        """
        try:
            # Debug kann explizit oder per Umgebungsvariable aktiviert werden
            env_debug = os.getenv("LEK_DEBUG_EXPORT_CONTEXT", "").strip().lower() in {"1", "true", "yes", "on"}
            debug_context_report = bool(debug_context_report or env_debug)

            # Suche nach passender Vorlage
            template_path = self.template_manager.find_matching_template(lek_theme)
            
            if template_path:
                # Verwende Vorlage als Basis
                doc = self.template_manager.load_template_as_base(template_path)
                print(f"Verwende Vorlage: {os.path.basename(template_path)}")

                # Quellkontext (Styles/Nummerierung) in Ziel-Dokument übernehmen,
                # damit Listen/Formatierungen aus Aufgaben in der Vorlage korrekt bleiben.
                self._prepare_target_document_context(doc, tasks)
                if debug_context_report:
                    self._print_context_report()
                
                # Ersetze Platzhalter wie 'TitelFürThema' durch das LEK-Thema
                self.template_manager.replace_template_placeholders(doc, lek_theme)
                
                # Aufgaben ab Seite 2 einfügen
                self.template_manager.insert_tasks_from_page_2(doc, tasks, lek_theme)
                
            else:
                # Fallback: Erstelle neues Dokument mit programmatischem Deckblatt
                doc = Document()
                
                # Seitenränder einstellen
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1.5 / 2.54)      # 1,5 cm oben
                    section.bottom_margin = Inches(1.5 / 2.54)   # 1,5 cm unten  
                    section.left_margin = Inches(2.5 / 2.54)     # 2,5 cm links
                    section.right_margin = Inches(1.5 / 2.54)    # 1,5 cm rechts
                
                # Standard-Absatzformat anpassen
                styles = doc.styles
                normal_style = styles['Normal']
                normal_font = normal_style.font
                normal_font.name = 'Aptos'
                normal_font.size = Pt(11)
                
                normal_paragraph_format = normal_style.paragraph_format
                normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                normal_paragraph_format.line_spacing = 1.1
                normal_paragraph_format.space_before = Pt(0)
                normal_paragraph_format.space_after = Pt(0)
                
                # Deckblatt erstellen
                self._create_cover_page(doc, lek_theme)
                
                # Neue Seite für den Inhalt
                doc.add_page_break()
                
                # Seitennummerierung ab Seite 2 einrichten (beginnt mit 1)
                self._add_page_numbering_from_page_2(doc)
                
                # Titel der Aufgaben hinzufügen
                title = doc.add_heading('Ihre Lernerfolgskontrolle', 0)
                
                # LEK-Thema als Untertitel hinzufügen, falls vorhanden
                if lek_theme:
                    theme_paragraph = doc.add_paragraph()
                    theme_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    theme_run = theme_paragraph.add_run(lek_theme)
                    theme_run.font.name = 'Aptos'
                    theme_run.font.size = Pt(16)
                    theme_run.font.color.rgb = RGBColor(0, 164, 153)  # Primärfarbe
                    theme_run.bold = True
                    theme_paragraph.paragraph_format.space_after = Pt(12)
                
                # Jede Aufgabe hinzufügen
                for i, task in enumerate(tasks, 1):
                    # Aufgaben als Heading1 + Punkte und zentraler Inhaltslogik
                    try:
                        self.append_task_to_lek_document(doc, task)
                    except Exception as e:
                        # Notfall-Fallback: Verwende original_paragraphs wenn elements-Kopierung fehlschlägt
                        print(f"Warnung: Element-Kopierung fehlgeschlagen für Aufgabe {i}: {e}")
                        if task.get('original_paragraphs'):
                            self._copy_paragraphs_for_lek(doc, task['original_paragraphs'])
                        else:
                            # Allerletzter Fallback
                            content_list = task.get('content', [])
                            for content_line in content_list:
                                if content_line.strip():
                                    doc.add_paragraph(content_line)
                    
                    # Metadaten als Kommentar (optional)
                    if task.get('difficulty') or task.get('keywords'):
                        meta_info = []
                        if task.get('difficulty'):
                            meta_info.append(f"Schwierigkeit: {task['difficulty']}")
                        if task.get('keywords'):
                            meta_info.append(f"Schlüsselwörter: {', '.join(task['keywords'][:5])}")
                        
                        meta_paragraph = doc.add_paragraph()
                        meta_paragraph.style = normal_style
                        meta_run = meta_paragraph.add_run(' | '.join(meta_info))
                        meta_run.italic = True
                    
                    # Trennlinie zwischen Aufgaben
                    if i < len(tasks):
                        separator_paragraph = doc.add_paragraph("─" * 50)
                        separator_paragraph.style = normal_style
            
            # Dokument speichern
            doc.save(output_path)
            
        except Exception as e:
            raise Exception(f"Fehler beim Erstellen der Word-Datei: {str(e)}")

    def _prepare_target_document_context(self, target_doc, tasks):
        """
        Übernimmt benötigte Styles und Nummerierungsdefinitionen aus dem
        Quell-Dokument der Aufgaben in das Ziel-Dokument.

        Args:
            target_doc: Ziel-Word-Dokument
            tasks: Aufgabenliste mit all_elements
        """
        source_part = self._get_source_part_from_tasks(tasks)
        if source_part is None:
            self._num_id_map = {}
            self._last_context_report = {
                'used_style_ids': [],
                'used_num_ids': [],
                'missing_styles_before': [],
                'missing_styles_after': [],
                'num_id_map': {}
            }
            return

        used_style_ids, used_num_ids = self._collect_used_style_and_num_ids(tasks)
        missing_before = self._get_missing_style_ids(target_doc.part, used_style_ids)
        self._merge_required_styles(source_part, target_doc.part, used_style_ids)
        missing_after = self._get_missing_style_ids(target_doc.part, used_style_ids)
        self._num_id_map = self._merge_required_numbering(source_part, target_doc.part, used_num_ids)

        self._last_context_report = {
            'used_style_ids': sorted(used_style_ids),
            'used_num_ids': sorted(used_num_ids, key=lambda x: int(x) if str(x).isdigit() else 10**9),
            'missing_styles_before': sorted(missing_before),
            'missing_styles_after': sorted(missing_after),
            'num_id_map': dict(self._num_id_map)
        }

    def _get_missing_style_ids(self, target_part, used_style_ids):
        """Ermittelt, welche verwendeten Style-IDs im Ziel noch fehlen."""
        try:
            target_styles = target_part.styles.element
            target_ids = {
                s.get(qn('w:styleId'))
                for s in target_styles.xpath('./w:style')
                if s.get(qn('w:styleId'))
            }
        except Exception:
            return set(used_style_ids)

        return {style_id for style_id in used_style_ids if style_id not in target_ids}

    def _print_context_report(self):
        """Gibt den zuletzt erstellten Kontext-Report aus."""
        report = self._last_context_report or {}
        used_styles = report.get('used_style_ids', [])
        used_nums = report.get('used_num_ids', [])
        missing_before = report.get('missing_styles_before', [])
        missing_after = report.get('missing_styles_after', [])
        num_map = report.get('num_id_map', {})

        print("[Export-Kontext] Styles verwendet      :", len(used_styles))
        print("[Export-Kontext] NumIDs verwendet      :", len(used_nums))
        print("[Export-Kontext] Styles fehlend vorher :", len(missing_before))
        print("[Export-Kontext] Styles fehlend nachher:", len(missing_after))

        if missing_before:
            print("[Export-Kontext] Fehlend vorher (Beispiel):", ', '.join(missing_before[:10]))
        if num_map:
            mappings = [f"{k}->{v}" for k, v in sorted(num_map.items(), key=lambda kv: int(kv[0]) if str(kv[0]).isdigit() else 10**9)]
            print("[Export-Kontext] NumID-Mapping:", ', '.join(mappings[:15]))

    def get_last_context_report(self):
        """Liefert den zuletzt berechneten Export-Kontext-Report zurück."""
        return dict(self._last_context_report) if self._last_context_report else None

    def _get_source_part_from_tasks(self, tasks):
        """
        Ermittelt den Dokument-Part des Quell-Dokuments über die ersten
        verfügbaren Paragraph-/Table-Objekte in den Aufgaben.
        """
        for task in tasks or []:
            for element in task.get('all_elements', []):
                content = element.get('content')
                if hasattr(content, 'part'):
                    return content.part
        return None

    def _collect_used_style_and_num_ids(self, tasks):
        """
        Sammelt verwendete Style-IDs und Num-IDs aus den Aufgaben-Elementen.
        """
        style_ids = set()
        num_ids = set()

        def _xml_of(element):
            content = element.get('content')
            if hasattr(content, '_element'):
                return content._element
            return content

        for task in tasks or []:
            for element in task.get('all_elements', []):
                xml_elem = _xml_of(element)
                if xml_elem is None:
                    continue

                for n in xml_elem.xpath('.//w:pStyle|.//w:rStyle|.//w:tblStyle'):
                    val = n.get(qn('w:val'))
                    if val:
                        style_ids.add(val)

                for n in xml_elem.xpath('.//w:numPr/w:numId'):
                    val = n.get(qn('w:val'))
                    if val:
                        num_ids.add(val)

        return style_ids, num_ids

    def _merge_required_styles(self, source_part, target_part, used_style_ids):
        """
        Fügt fehlende Styles (inkl. Abhängigkeiten) aus Quelle in Ziel ein.
        """
        from copy import deepcopy

        if not used_style_ids:
            return

        try:
            source_styles = source_part.styles.element
            target_styles = target_part.styles.element
        except Exception:
            return

        src_by_id = {
            s.get(qn('w:styleId')): s
            for s in source_styles.xpath('./w:style')
            if s.get(qn('w:styleId'))
        }
        tgt_ids = {
            s.get(qn('w:styleId'))
            for s in target_styles.xpath('./w:style')
            if s.get(qn('w:styleId'))
        }

        pending = list(used_style_ids)
        visited = set()

        while pending:
            style_id = pending.pop()
            if not style_id or style_id in visited:
                continue
            visited.add(style_id)

            src_style = src_by_id.get(style_id)
            if src_style is None:
                continue

            # Abhängigkeiten zuerst aufnehmen
            for dep_xpath in ('./w:basedOn', './w:link', './w:next'):
                dep_nodes = src_style.xpath(dep_xpath)
                for dep in dep_nodes:
                    dep_id = dep.get(qn('w:val'))
                    if dep_id and dep_id not in visited:
                        pending.append(dep_id)

            if style_id not in tgt_ids:
                target_styles.append(deepcopy(src_style))
                tgt_ids.add(style_id)

    def _merge_required_numbering(self, source_part, target_part, used_num_ids):
        """
        Fügt benötigte Nummerierungsdefinitionen aus Quelle in Ziel ein.
        Bei ID-Konflikten werden neue IDs erzeugt und zurückgemappt.

        Returns:
            dict: Mapping alter numId -> neue/übernommene numId
        """
        from copy import deepcopy

        num_id_map = {}
        if not used_num_ids:
            return num_id_map

        try:
            src_numbering = source_part.numbering_part.element
            tgt_numbering = target_part.numbering_part.element
        except Exception:
            return num_id_map

        src_nums = {
            n.get(qn('w:numId')): n
            for n in src_numbering.xpath('./w:num')
            if n.get(qn('w:numId'))
        }
        tgt_nums = {
            n.get(qn('w:numId')): n
            for n in tgt_numbering.xpath('./w:num')
            if n.get(qn('w:numId'))
        }

        src_abs = {
            a.get(qn('w:abstractNumId')): a
            for a in src_numbering.xpath('./w:abstractNum')
            if a.get(qn('w:abstractNumId'))
        }
        tgt_abs = {
            a.get(qn('w:abstractNumId')): a
            for a in tgt_numbering.xpath('./w:abstractNum')
            if a.get(qn('w:abstractNumId'))
        }

        def _next_free_id(existing_ids):
            max_id = max([int(i) for i in existing_ids if str(i).isdigit()] + [0])
            return str(max_id + 1)

        for num_id in sorted(used_num_ids, key=lambda x: int(x) if str(x).isdigit() else 10**9):
            src_num = src_nums.get(num_id)
            if src_num is None:
                continue

            abs_nodes = src_num.xpath('./w:abstractNumId')
            if not abs_nodes:
                continue
            src_abs_id = abs_nodes[0].get(qn('w:val'))

            target_num_id = num_id
            if target_num_id in tgt_nums:
                # ID bereits belegt -> remappen
                target_num_id = _next_free_id(tgt_nums.keys())

            target_abs_id = src_abs_id
            if target_abs_id in tgt_abs:
                # Abstract-ID belegt -> remappen
                target_abs_id = _next_free_id(tgt_abs.keys())

            # AbstractNum sicherstellen
            src_abs_def = src_abs.get(src_abs_id)
            if src_abs_def is not None:
                new_abs = deepcopy(src_abs_def)
                new_abs.set(qn('w:abstractNumId'), target_abs_id)
                if target_abs_id not in tgt_abs:
                    tgt_numbering.append(new_abs)
                    tgt_abs[target_abs_id] = new_abs

            # Num hinzufügen
            new_num = deepcopy(src_num)
            new_num.set(qn('w:numId'), target_num_id)
            new_abs_nodes = new_num.xpath('./w:abstractNumId')
            if new_abs_nodes:
                new_abs_nodes[0].set(qn('w:val'), target_abs_id)

            if target_num_id not in tgt_nums:
                tgt_numbering.append(new_num)
                tgt_nums[target_num_id] = new_num

            num_id_map[num_id] = target_num_id

        return num_id_map

    def _remap_num_ids_in_element(self, xml_element, num_id_map):
        """
        Ersetzt numId-Werte innerhalb eines kopierten XML-Elements anhand Mapping.
        """
        if not num_id_map:
            return

        for node in xml_element.xpath('.//w:numPr/w:numId'):
            old_id = node.get(qn('w:val'))
            if old_id in num_id_map:
                node.set(qn('w:val'), num_id_map[old_id])
    
    def _add_page_numbering_from_page_2(self, doc):
        """
        Fügt Seitennummerierung ab Seite 2 hinzu (beginnend mit 1)
        Format: "Seite/Gesamtseiten" rechtsbündig in der Fußzeile
        
        Args:
            doc: Word-Dokument
        """
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement
        
        # Erste Sektion (Deckblatt) - keine Seitennummern
        first_section = doc.sections[0]
        first_section.different_first_page_header_footer = True
        
        # Neue Sektion für Inhalt ab Seite 2
        if len(doc.sections) == 1:
            # Falls nur eine Sektion existiert, füge eine neue hinzu
            new_section = doc.add_section()
        else:
            new_section = doc.sections[-1]
        
        # Fußzeile für die neue Sektion konfigurieren
        footer = new_section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Seitennummerierung hinzufügen (Feld-Code für PAGE und NUMPAGES)
        run = footer_paragraph.add_run()
        
        # XML für Seitennummerierung erstellen
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE \\* MERGEFORMAT'
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # " / " Trenner
        footer_paragraph.add_run(' / ')
        
        # Gesamtseitenzahl
        run2 = footer_paragraph.add_run()
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES \\* MERGEFORMAT'
        run2._r.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar4)
        
        # Formatierung der Fußzeile
        for run in footer_paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Aptos'
    
    def get_document_info(self, file_path):
        """
        Gibt grundlegende Informationen über ein Word-Dokument zurück
        
        Args:
            file_path (str): Pfad zur Word-Datei
            
        Returns:
            dict: Dokumentinformationen
        """
        try:
            doc = Document(file_path)
            
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Geschätzte Aufgabenanzahl basierend auf erkannten Mustern
            task_count = 0
            for paragraph in doc.paragraphs:
                if self._is_task_start(paragraph.text.strip()):
                    task_count += 1
            
            return {
                'paragraph_count': paragraph_count,
                'estimated_task_count': task_count,
                'file_size': f"{len(doc.paragraphs)} Absätze"
            }
            
        except Exception as e:
            raise Exception(f"Fehler beim Analysieren der Datei: {str(e)}")
    
    def _extract_table_text(self, table_element, doc):
        """
        Extrahiert Text aus einer Tabelle für Keyword-Analyse
        
        Args:
            table_element: XML-Element der Tabelle
            doc: Word-Dokument-Objekt
            
        Returns:
            str: Extrahierter Tabellentext
        """
        try:
            # Finde entsprechendes Table-Objekt
            for table in doc.tables:
                if table._element == table_element:
                    text_parts = []
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                text_parts.append(cell_text)
                    return ' '.join(text_parts)
            return ""
        except Exception:
            return ""