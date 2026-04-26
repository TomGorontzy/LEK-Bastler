"""
WordProcessor - Klasse für die Verarbeitung von Word-Dokumenten
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from template_manager import TemplateManager
import re
import os

class WordProcessor:
    """Klasse für das Lesen und Schreiben von Word-Dokumenten"""
    
    def __init__(self):
        self.template_manager = TemplateManager()
    
    def extract_tasks(self, file_path):
        """
        Extrahiert Aufgaben aus einer Word-Datei
        Findet Überschrift 1 und markiert alles bis zur nächsten Überschrift 1 als eine Aufgabe
        
        Args:
            file_path (str): Pfad zur Word-Datei
            
        Returns:
            list: Liste von Aufgaben-Dictionaries mit kompletter Struktur
        """
        try:
            doc = Document(file_path)
            tasks = []
            
            # Erweiterte Methode: Sammle alle Body-Elemente (Paragraphen UND Tabellen)
            all_elements = self._get_all_body_elements(doc)
            task_ranges = []
            
            # Schritt 1: Finde alle Überschrift 1 Positionen
            heading1_positions = []
            for i, element in enumerate(all_elements):
                if element['type'] == 'paragraph' and self._is_heading1(element['content']):
                    heading1_positions.append(i)
            
            # Schritt 2: Erstelle Aufgaben-Bereiche von Überschrift 1 bis zur nächsten
            for i, start_pos in enumerate(heading1_positions):
                end_pos = heading1_positions[i + 1] if i + 1 < len(heading1_positions) else len(all_elements)
                
                # Sammle alle Elemente von start_pos bis end_pos (exklusiv)
                task_elements = all_elements[start_pos:end_pos]
                
                if task_elements:
                    # Title aus erstem Paragraph extrahieren
                    first_element = task_elements[0]
                    task_title = first_element['content'].text.strip() if first_element['type'] == 'paragraph' else f"Aufgabe {i + 1}"
                    
                    # Sammle Text-Inhalt für Keyword/Difficulty-Extraktion (nur von Paragraphen)
                    content = []
                    for element in task_elements:
                        if element['type'] == 'paragraph' and element['content'].text.strip():
                            content.append(element['content'].text.strip())
                        elif element['type'] == 'table':
                            # Füge Tabellen-Text für Keyword-Analyse hinzu
                            table_text = self._extract_table_text_for_keywords(element['content'])
                            if table_text:
                                content.append(table_text)
                    
                    full_content = ' '.join(content)
                    
                    task = {
                        'number': i + 1,
                        'title': self._extract_task_title(task_title),
                        'content': content,
                        'all_elements': task_elements,  # NEUE: Komplette Element-Liste (Paragraphen + Tabellen)
                        'original_paragraphs': [elem['content'] for elem in task_elements if elem['type'] == 'paragraph'],  # Rückwärtskompatibilität
                        'difficulty': 'Mittel',  # Wird später aktualisiert
                        'keywords': []  # Wird später aktualisiert
                    }
                    
                    tasks.append(task)
            
            # Schritt 3: Metadaten für alle Aufgaben extrahieren
            for task in tasks:
                full_content = ' '.join(task.get('content', []))
                if full_content:
                    # Explizite Schwierigkeit suchen, dann automatische Extraktion
                    explicit_difficulty = self._extract_explicit_difficulty(full_content)
                    if explicit_difficulty:
                        task['difficulty'] = explicit_difficulty
                    else:
                        task['difficulty'] = self._extract_difficulty(full_content)
                    
                    # Explizite Keywords suchen, dann automatische Extraktion
                    explicit_keywords = self._extract_explicit_keywords(full_content)
                    if explicit_keywords:
                        task['keywords'] = explicit_keywords
                    else:
                        task['keywords'] = self._extract_keywords(full_content)
            
            return tasks
            
        except Exception as e:
            raise Exception(f"Fehler beim Lesen der Word-Datei: {str(e)}")
    
    def _get_heading_level(self, style_name):
        """
        Bestimmt das korrekte Heading-Level basierend auf dem Style-Namen
        
        Args:
            style_name: Name des Word-Styles
            
        Returns:
            int: Das entsprechende Heading-Level (1-9)
        """
        if not style_name:
            return 1
            
        style_lower = style_name.lower()
        
        # Direkte Zuordnung für deutsche und englische Styles
        heading_map = {
            'heading 1': 1, 'überschrift 1': 1,
            'heading 2': 2, 'überschrift 2': 2,
            'heading 3': 3, 'überschrift 3': 3,
            'heading 4': 4, 'überschrift 4': 4,
            'heading 5': 5, 'überschrift 5': 5,
            'heading 6': 6, 'überschrift 6': 6,
            'heading 7': 7, 'überschrift 7': 7,
            'heading 8': 8, 'überschrift 8': 8,
            'heading 9': 9, 'überschrift 9': 9
        }
        
        if style_lower in heading_map:
            return heading_map[style_lower]
        
        # Fallback: Versuche Zahl aus Style-Namen zu extrahieren
        import re
        match = re.search(r'\d+', style_name)
        if match:
            level = int(match.group())
            return min(max(level, 1), 9)  # Begrenze auf 1-9
        
        return 1  # Standard-Fallback

    def _is_heading1(self, paragraph):
        """
        Überprüft, ob ein Paragraph eine Überschrift 1 ist
        
        Args:
            paragraph: Word-Paragraph-Objekt
            
        Returns:
            bool: True wenn es eine Überschrift 1 ist
        """
        try:
            if paragraph.style:
                style_name = paragraph.style.name
                return (style_name == 'Heading 1' or 
                        style_name == 'Überschrift 1' or
                        style_name.lower() == 'heading 1' or
                        style_name.lower() == 'überschrift 1')
        except:
            # Fallback: Prüfe auf Textmuster wenn Style nicht verfügbar
            text = paragraph.text.strip()
            if text:
                return self._is_task_start(text, paragraph)
        
        return False

    def _get_all_body_elements(self, doc):
        """
        Sammelt alle Body-Elemente (Paragraphen und Tabellen) in der richtigen Reihenfolge
        
        Args:
            doc: Word-Dokument
            
        Returns:
            list: Liste von Element-Dictionaries mit 'type' und 'content'
        """
        elements = []
        
        # Nutze XML-Struktur für korrekte Reihenfolge
        body = doc._body._element
        
        for child in body:
            tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag_name == 'p':  # Paragraph
                # Finde entsprechendes Paragraph-Objekt
                for para in doc.paragraphs:
                    if para._element == child:
                        elements.append({
                            'type': 'paragraph',
                            'content': para
                        })
                        break
                        
            elif tag_name == 'tbl':  # Table
                # Finde entsprechendes Table-Objekt
                for table in doc.tables:
                    if table._element == child:
                        elements.append({
                            'type': 'table',
                            'content': table
                        })
                        break
        
        return elements

    def _extract_table_text_for_keywords(self, table):
        """
        Extrahiert Text aus Tabelle für Keyword-Analyse
        
        Args:
            table: Word-Table-Objekt
            
        Returns:
            str: Tabellen-Text zusammengefasst
        """
        try:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' '.join(row_text))
            return ' '.join(table_text)
        except:
            return ""
    
    def _create_cover_page(self, doc, lek_theme=""):
        """
        Erstellt ein Deckblatt (Semaphor) mit Farbschema basierend auf #00a499
        
        Args:
            doc: Das Word-Dokument
            lek_theme (str): LEK-Thema für das Deckblatt
        """
        # Farbpalette basierend auf #00a499
        primary_color = RGBColor(0, 164, 153)      # #00a499
        light_color = RGBColor(102, 204, 195)      # Hellere Version
        dark_color = RGBColor(0, 120, 112)         # Dunklere Version
        
        # Haupttitel
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run("LERNERFOLGSKONTROLLE")
        title_run.font.name = 'Aptos'
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = primary_color
        title_run.bold = True
        
        # Leerzeilen
        for _ in range(3):
            doc.add_paragraph()
        
        # Semaphor-Design (vereinfacht mit Text-Elementen)
        semaphor_paragraph = doc.add_paragraph()
        semaphor_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Oberer Balken
        top_run = semaphor_paragraph.add_run("━━━━━━━━━━━━━━━━━━━━")
        top_run.font.size = Pt(16)
        top_run.font.color.rgb = dark_color
        top_run.bold = True
        
        doc.add_paragraph()
        
        # Mittlerer Bereich
        middle_paragraph = doc.add_paragraph()
        middle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if lek_theme:
            middle_run = middle_paragraph.add_run(f"◆  {lek_theme.upper()}  ◆")
        else:
            middle_run = middle_paragraph.add_run("◆  AUFGABENSAMMLUNG  ◆")
            
        middle_run.font.name = 'Aptos'
        middle_run.font.size = Pt(18)
        middle_run.font.color.rgb = light_color
        middle_run.bold = True
        
        doc.add_paragraph()
        
        # Unterer Balken
        bottom_paragraph = doc.add_paragraph()
        bottom_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom_run = bottom_paragraph.add_run("━━━━━━━━━━━━━━━━━━━━")
        bottom_run.font.size = Pt(16)
        bottom_run.font.color.rgb = dark_color
        bottom_run.bold = True
        
        # Weitere Leerzeilen
        for _ in range(8):
            doc.add_paragraph()
        
        # Datum
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from datetime import datetime
        date_run = date_paragraph.add_run(f"Erstellt am {datetime.now().strftime('%d.%m.%Y')}")
        date_run.font.name = 'Aptos'
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = primary_color
        date_run.italic = True
    
    def _copy_elements_with_formatting(self, doc, elements):
        """
        Kopiert alle Arten von Word-Elementen (Paragraphen, Tabellen, etc.) mit Formatierung
        Verwendet python-docx API statt direkte XML-Manipulation für bessere Kompatibilität
        
        Args:
            doc: Ziel-Word-Dokument  
            elements: Liste von XML-Elementen
        """
        from docx import Document
        
        # Erstelle temporäres Dokument mit den Original-Elementen
        temp_doc = Document()
        
        for element in elements:
            try:
                if element.tag.endswith('p'):  # Paragraph
                    # Finde den entsprechenden Paragraph im Original-Dokument
                    # und kopiere ihn über die python-docx API
                    self._copy_paragraph_safely(doc, element)
                    
                elif element.tag.endswith('tbl'):  # Tabelle
                    # Tabelle über python-docx API kopieren
                    self._copy_table_safely(doc, element)
                
                else:
                    # Andere Strukturen als Paragraph mit Hinweis
                    doc.add_paragraph(f"[Struktur: {element.tag.split('}')[-1] if '}' in element.tag else element.tag}]")
                    
            except Exception as e:
                # Fallback: Fehlermeldung als Paragraph
                doc.add_paragraph(f"[Element konnte nicht kopiert werden: {str(e)[:50]}...]")
    
    def _copy_paragraph_safely(self, doc, element):
        """
        Kopiert einen Paragraph sicher über python-docx API
        """
        try:
            # Erstelle neuen Paragraph 
            new_para = doc.add_paragraph()
            
            # Kopiere Text-Runs und deren Formatierung
            for run_elem in element.xpath('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                text_content = ""
                for text_elem in run_elem.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    text_content += text_elem.text or ""
                
                if text_content:
                    run = new_para.add_run(text_content)
                    # Basis-Formatierung extrahieren
                    self._apply_run_formatting(run, run_elem)
            
            # Paragraph-Stil versuchen zu übernehmen
            self._apply_paragraph_formatting(new_para, element)
            
        except Exception as e:
            # Fallback: Nur Text ohne Formatierung
            doc.add_paragraph(f"[Text konnte nicht formatiert werden]")
    
    def _copy_table_safely(self, doc, element):
        """
        Kopiert eine Tabelle sicher über python-docx API
        """
        try:
            # Fallback: Tabelle als Beschreibung
            doc.add_paragraph("[Tabelle aus Original-Dokument]")
        except Exception as e:
            doc.add_paragraph("[Tabelle konnte nicht kopiert werden]")
    
    def _apply_run_formatting(self, run, run_element):
        """
        Wendet Basis-Formatierung auf einen Run an
        """
        try:
            # Bold
            bold_elem = run_element.xpath('.//w:b', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if bold_elem:
                run.bold = True
            
            # Italic  
            italic_elem = run_element.xpath('.//w:i', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if italic_elem:
                run.italic = True
                
        except:
            pass  # Formatierung fehlgeschlagen, aber Text ist da
    
    def _apply_paragraph_formatting(self, paragraph, para_element):
        """
        Wendet Basis-Paragraph-Formatierung an
        """
        try:
            # Heading-Stil erkennen
            style_elem = para_element.xpath('.//w:pStyle', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if style_elem and len(style_elem) > 0:
                style_val = style_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                if 'Heading' in style_val or 'Überschrift' in style_val:
                    # Versuche Heading-Level zu extrahieren
                    if '1' in style_val:
                        paragraph.style = 'Heading 1'
                    elif '2' in style_val:
                        paragraph.style = 'Heading 2'
                    elif '3' in style_val:
                        paragraph.style = 'Heading 3'
        except:
            pass  # Stil-Anwendung fehlgeschlagen

    def _copy_elements_with_formatting(self, doc, elements):
        """
        Kopiert alle Elemente (Paragraphen und Tabellen) mit ihrer ursprünglichen Formatierung
        
        Args:
            doc: Ziel-Word-Dokument
            elements: Liste von Element-Dictionaries mit 'type' und 'content'
        """
        for element in elements:
            if element['type'] == 'paragraph':
                # Verwende bestehende Paragraph-Kopierfunktion
                self._copy_paragraphs_with_formatting(doc, [element['content']])
                
            elif element['type'] == 'table':
                # Kopiere Tabelle mit vollständiger Formatierung
                self._copy_table_with_formatting(doc, element['content'])

    def _copy_elements_for_lek(self, doc, elements):
        """
        Spezielle Kopierfunktion für LEK-Erstellung mit angepassten Überschrift-Levels
        
        Args:
            doc: Ziel-Word-Dokument (LEK)
            elements: Liste von Element-Dictionaries mit 'type' und 'content'
        """
        for element in elements:
            if element['type'] == 'paragraph':
                # Verwende LEK-spezifische Paragraph-Kopierfunktion
                self._copy_paragraphs_for_lek(doc, [element['content']])
                
            elif element['type'] == 'table':
                # Kopiere Tabelle mit vollständiger Formatierung
                self._copy_table_with_formatting(doc, element['content'])

    def _copy_paragraphs_for_lek(self, doc, paragraphs):
        """
        Kopiert Paragraphen für LEK-Erstellung mit angepassten Überschrift-Levels
        
        Args:
            doc: Ziel-Word-Dokument (LEK)
            paragraphs: Liste von Quell-Paragraphen
        """
        for paragraph in paragraphs:
            if not paragraph.text.strip():
                # Leere Absätze als solche übernehmen
                doc.add_paragraph()
                continue
            
            # Prüfe ob es eine Überschrift ist
            if (paragraph.style and 
                (paragraph.style.name.startswith('Heading') or 
                 paragraph.style.name.startswith('Überschrift') or
                 'Heading' in paragraph.style.name or
                 'Überschrift' in paragraph.style.name)):
                
                # Extrahiere ursprüngliches Level und behalte es bei (KEIN Shift zu Level 2!)
                original_level = self._get_heading_level(paragraph.style.name)
                
                # Erstelle Heading mit ORIGINALEM Level
                new_paragraph = doc.add_heading(paragraph.text, level=original_level)
                
            else:
                # Normaler Absatz - kopiere mit bestehender Logik
                new_paragraph = doc.add_paragraph()
                
                # Absatzformatierung kopieren (vereinfacht für LEK)
                try:
                    if paragraph.style and not paragraph.style.name.startswith('Heading') and not paragraph.style.name.startswith('Überschrift'):
                        try:
                            new_paragraph.style = paragraph.style
                        except:
                            pass
                except:
                    pass
                
                # Runs mit Formatierung kopieren
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    try:
                        if run.font:
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                    except:
                        pass

    def _copy_table_with_formatting(self, doc, source_table):
        """
        Kopiert eine Tabelle mit vollständiger Formatierung
        
        Args:
            doc: Ziel-Word-Dokument
            source_table: Quell-Tabelle
        """
        try:
            # Erstelle neue Tabelle mit gleicher Struktur
            rows = len(source_table.rows)
            cols = len(source_table.columns) if source_table.columns else 1
            
            new_table = doc.add_table(rows=rows, cols=cols)
            
            # Kopiere Zellinhalt und Formatierung
            for row_idx, source_row in enumerate(source_table.rows):
                target_row = new_table.rows[row_idx]
                
                for cell_idx, source_cell in enumerate(source_row.cells):
                    if cell_idx < len(target_row.cells):
                        target_cell = target_row.cells[cell_idx]
                        
                        # Kopiere Text-Inhalt
                        target_cell.text = source_cell.text
                        
                        # Kopiere Zellenformatierung (Basic)
                        try:
                            # Paragraph-Formatierung in der Zelle übernehmen
                            if source_cell.paragraphs:
                                source_para = source_cell.paragraphs[0]
                                target_para = target_cell.paragraphs[0]
                                
                                # Alignment kopieren
                                if source_para.paragraph_format.alignment:
                                    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
                                
                                # Text-Formatierung kopieren
                                if source_para.runs:
                                    # Lösche bestehenden Text
                                    target_para.clear()
                                    for run in source_para.runs:
                                        new_run = target_para.add_run(run.text)
                                        if run.bold:
                                            new_run.bold = True
                                        if run.italic:
                                            new_run.italic = True
                                        if run.underline:
                                            new_run.underline = True
                                        if run.font.size:
                                            new_run.font.size = run.font.size
                        except:
                            # Formatierung fehlgeschlagen, aber Text ist da
                            pass
            
        except Exception as e:
            # Fallback: Einfache Tabellen-Ersetzung
            doc.add_paragraph(f"[Tabelle mit {len(source_table.rows)} Zeilen konnte nicht vollständig kopiert werden]")

    def _copy_paragraphs_with_formatting(self, doc, paragraphs):
        """
        Kopiert Paragraphen mit ihrer ursprünglichen Formatierung
        
        Args:
            doc: Ziel-Word-Dokument
            paragraphs: Liste von Quell-Paragraphen
        """
        for paragraph in paragraphs:
            if not paragraph.text.strip():
                # Leere Absätze als solche übernehmen
                doc.add_paragraph()
                continue
            
            # Neuen Absatz erstellen - prüfe zuerst ob es eine Überschrift ist
            if (paragraph.style and 
                (paragraph.style.name.startswith('Heading') or 
                 paragraph.style.name.startswith('Überschrift') or
                 'Heading' in paragraph.style.name or
                 'Überschrift' in paragraph.style.name)):
                
                # Extrahiere Heading-Level präzise basierend auf Style-Namen
                level = self._get_heading_level(paragraph.style.name)
                
                # Erstelle Heading mit Text
                new_paragraph = doc.add_heading(paragraph.text, level=level)
                
            else:
                # Normaler Absatz
                new_paragraph = doc.add_paragraph()
            
            # Absatzformatierung kopieren
            try:
                if paragraph.style:
                    # Versuche den Style zu übernehmen, außer bei Überschriften
                    if not paragraph.style.name.startswith('Heading') and not paragraph.style.name.startswith('Überschrift'):
                        try:
                            new_paragraph.style = paragraph.style
                        except:
                            # Falls Style nicht verfügbar, verwende Normal
                            pass
                
                # Absatz-Formatierung übernehmen
                if paragraph.paragraph_format:
                    try:
                        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
                        new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
                        new_paragraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
                        new_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
                        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
                        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
                        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
                        new_paragraph.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
                    except:
                        # Falls Formatierung nicht kopiert werden kann, ignorieren
                        pass
            except:
                # Falls Style-Zugriff fehlschlägt, ignorieren
                pass
            
            # Runs mit Formatierung kopieren (nur wenn es kein Heading ist, da dort Text schon gesetzt wurde)
            if not (paragraph.style and 
                    (paragraph.style.name.startswith('Heading') or 
                     paragraph.style.name.startswith('Überschrift') or
                     'Heading' in paragraph.style.name or
                     'Überschrift' in paragraph.style.name)):
                
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    
                    try:
                        # Font-Formatierung kopieren
                        if run.font:
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.color and run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                            
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            
                            # Weitere Font-Eigenschaften
                            try:
                                new_run.font.subscript = run.font.subscript
                                new_run.font.superscript = run.font.superscript
                                new_run.font.strike = run.font.strike
                                new_run.font.small_caps = run.font.small_caps
                                new_run.font.all_caps = run.font.all_caps
                            except:
                                # Ignoriere erweiterte Formatierungen falls nicht verfügbar
                                pass
                    except:
                        # Falls Font-Formatierung nicht kopiert werden kann, ignorieren
                        pass
    
    def _is_task_start(self, text, paragraph):
        """
        Erkennt, ob ein Text der Beginn einer neuen Aufgabe ist
        Priorisiert Überschrift 1 Format, fallback auf Textmuster
        
        Args:
            text (str): Zu prüfender Text
            paragraph: Word-Paragraph-Objekt
            
        Returns:
            bool: True wenn Aufgabenbeginn erkannt
        """
        # Primäre Erkennung: Überschrift 1 Formatvorlage
        try:
            if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Überschrift 1':
                return True
        except:
            # Falls Style-Zugriff fehlschlägt, weiter mit Textmustern
            pass
        
        # Sekundäre Erkennung: Verschiedene Textmuster für Aufgabenbeginne
        patterns = [
            r'^Aufgabe\s*\d+',
            r'^\d+\.',
            r'^Übung\s*\d+',
            r'^Frage\s*\d+',
            r'^Problem\s*\d+',
            r'^Nr\.\s*\d+',
            r'^\w\)',  # a), b), etc.
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _extract_task_title(self, text):
        """
        Extrahiert den Titel einer Aufgabe
        
        Args:
            text (str): Text der ersten Zeile der Aufgabe
            
        Returns:
            str: Extrahierter Titel
        """
        # Entfernt Nummerierung und behält den Rest
        cleaned = re.sub(r'^(Aufgabe\s*\d+[:\.]?\s*|^\d+\.\s*|^Übung\s*\d+[:\.]?\s*)', '', text, flags=re.IGNORECASE)
        
        # Falls der Text sehr lang ist, nur die ersten 50 Zeichen nehmen
        if len(cleaned) > 50:
            cleaned = cleaned[:47] + "..."
        
        return cleaned.strip() or "Ohne Titel"
    
    def _extract_difficulty(self, text):
        """
        Versucht den Schwierigkeitsgrad aus dem Text zu extrahieren
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            str: Schwierigkeitsgrad (Leicht, Mittel, Schwer)
        """
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['leicht', 'einfach', 'basic', 'anfänger']):
            return 'Leicht'
        elif any(word in text_lower for word in ['schwer', 'difficult', 'komplex', 'advanced']):
            return 'Schwer'
        elif any(word in text_lower for word in ['mittel', 'medium', 'intermediate']):
            return 'Mittel'
        
        # Standard-Schwierigkeitsgrad basierend auf Textlänge
        if len(text) > 200:
            return 'Schwer'
        elif len(text) > 100:
            return 'Mittel'
        else:
            return 'Leicht'
    
    def _extract_keywords(self, text):
        """
        Extrahiert relevante Schlüsselwörter aus dem Text
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            list: Liste von Schlüsselwörtern
        """
        # Stopwörter, die ignoriert werden sollen
        stop_words = {
            'der', 'die', 'das', 'ein', 'eine', 'und', 'oder', 'aber', 'von', 'zu', 'mit', 'bei', 'für',
            'ist', 'sind', 'war', 'waren', 'hat', 'haben', 'wird', 'werden', 'kann', 'können', 'soll',
            'sollte', 'auf', 'in', 'an', 'über', 'unter', 'durch', 'gegen', 'ohne', 'um', 'nach', 'vor'
        }
        
        # Text in Wörter aufteilen und bereinigen
        words = re.findall(r'\b\w+\b', text.lower())
        
        # Filtern: mindestens 3 Zeichen, keine Stopwörter, keine Zahlen
        keywords = [
            word for word in words 
            if len(word) >= 3 
            and word not in stop_words 
            and not word.isdigit()
        ]
        
        # Nur die häufigsten/relevantesten Keywords zurückgeben
        return list(set(keywords))[:10]
    
    def _extract_explicit_keywords(self, text):
        """
        Extrahiert explizite Schlüsselwörter aus "Schlüsselwörter:" Zeilen
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            list: Liste von explizit angegebenen Schlüsselwörtern oder leere Liste
        """
        import re
        
        # Suche nach "Schlüsselwörter:" gefolgt von kommagetrenner Liste
        patterns = [
            r'Schlüsselwörter:\s*([^\n\r]+)',
            r'Schlüsselwort:\s*([^\n\r]+)',
            r'Keywords:\s*([^\n\r]+)',
            r'Stichwörter:\s*([^\n\r]+)',
            r'Stichworte:\s*([^\n\r]+)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Nehme den ersten/längsten Treffer
                keywords_text = max(matches, key=len) if len(matches) > 1 else matches[0]
                
                # Trenne an Kommas und bereinige
                keywords = []
                for keyword in keywords_text.split(','):
                    cleaned = keyword.strip()
                    if cleaned and len(cleaned) >= 2:  # Mindestens 2 Zeichen
                        keywords.append(cleaned.lower())
                
                if keywords:
                    return keywords[:10]  # Maximal 10 Keywords
        
        return []  # Keine expliziten Keywords gefunden
    
    def _extract_explicit_difficulty(self, text):
        """
        Extrahiert explizite Schwierigkeitsangaben aus "Schwierigkeit:" Zeilen
        
        Args:
            text (str): Zu analysierender Text
            
        Returns:
            str: Explizit angegebene Schwierigkeit oder None
        """
        import re
        
        # Suche nach "Schwierigkeit:" gefolgt von Schwierigkeitsgrad
        patterns = [
            r'Schwierigkeit:\s*([^\n\r]+)',
            r'Schwierigkeitsgrad:\s*([^\n\r]+)',
            r'Difficulty:\s*([^\n\r]+)',
            r'Level:\s*([^\n\r]+)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                difficulty_text = matches[0].strip().lower()
                
                # Normalisiere auf Standard-Schwierigkeitsgrade
                if any(word in difficulty_text for word in ['leicht', 'einfach', 'easy', 'simple', '1']):
                    return 'Leicht'
                elif any(word in difficulty_text for word in ['schwer', 'difficult', 'hard', 'komplex', 'complex', '3']):
                    return 'Schwer'
                elif any(word in difficulty_text for word in ['mittel', 'medium', 'intermediate', '2']):
                    return 'Mittel'
                else:
                    # Fallback: Versuche direkte Zuordnung
                    if difficulty_text in ['leicht', 'mittel', 'schwer']:
                        return difficulty_text.capitalize()
        
        return None  # Keine explizite Schwierigkeit gefunden
    
    def create_document_from_tasks(self, tasks, output_path, lek_theme=""):
        """
        Erstellt ein neues Word-Dokument aus einer Liste von Aufgaben
        Verwendet Vorlagen aus dem Vorlagen-Ordner falls verfügbar
        
        Args:
            tasks (list): Liste von Aufgaben-Dictionaries
            output_path (str): Pfad für die neue Datei
            lek_theme (str): LEK-Thema für Vorlagen-Auswahl
        """
        try:
            # Suche nach passender Vorlage
            template_path = self.template_manager.find_matching_template(lek_theme)
            
            if template_path:
                # Verwende Vorlage als Basis
                doc = self.template_manager.load_template_as_base(template_path)
                print(f"Verwende Vorlage: {os.path.basename(template_path)}")
                
                # Ersetze Platzhalter wie 'TitelFürThema' durch das LEK-Thema
                self.template_manager.replace_template_placeholders(doc, lek_theme)
                
                # Aufgaben ab Seite 3 einfügen
                self.template_manager.insert_tasks_from_page_3(doc, tasks, lek_theme)
                
            else:
                # Fallback: Erstelle neues Dokument mit programmatischem Deckblatt
                doc = Document()
                
                # Seitenränder einstellen
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1.5 / 2.54)      # 1,5 cm oben
                    section.bottom_margin = Inches(1.5 / 2.54)   # 1,5 cm unten  
                    section.left_margin = Inches(2.5 / 2.54)     # 2,5 cm links
                    section.right_margin = Inches(1.5 / 2.54)    # 1,5 cm rechts
                
                # Standard-Absatzformat anpassen
                styles = doc.styles
                normal_style = styles['Normal']
                normal_font = normal_style.font
                normal_font.name = 'Aptos'
                normal_font.size = Pt(11)
                
                normal_paragraph_format = normal_style.paragraph_format
                normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                normal_paragraph_format.line_spacing = 1.1
                normal_paragraph_format.space_before = Pt(0)
                normal_paragraph_format.space_after = Pt(0)
                
                # Deckblatt erstellen
                self._create_cover_page(doc, lek_theme)
                
                # Neue Seite für den Inhalt
                doc.add_page_break()
                
                # Seitennummerierung ab Seite 2 einrichten (beginnt mit 1)
                self._add_page_numbering_from_page_2(doc)
                
                # Titel der Aufgaben hinzufügen
                title = doc.add_heading('Ihre Lernerfolgskontrolle', 0)
                
                # LEK-Thema als Untertitel hinzufügen, falls vorhanden
                if lek_theme:
                    theme_paragraph = doc.add_paragraph()
                    theme_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    theme_run = theme_paragraph.add_run(lek_theme)
                    theme_run.font.name = 'Aptos'
                    theme_run.font.size = Pt(16)
                    theme_run.font.color.rgb = RGBColor(0, 164, 153)  # Primärfarbe
                    theme_run.bold = True
                    theme_paragraph.paragraph_format.space_after = Pt(12)
                
                # Jede Aufgabe hinzufügen
                for i, task in enumerate(tasks, 1):
                    # Aufgabeninhalt mit vollständiger Struktur - verwende LEK-spezifische Kopier-Logik
                    try:
                        if task.get('all_elements'):
                            # Beste Option: Kopiere alle Elemente mit LEK-angepassten Überschrift-Levels
                            self._copy_elements_for_lek(doc, task['all_elements'])
                        elif task.get('original_paragraphs'):
                            # Sichere Fallback: Kopiere nur Paragraphen mit LEK-Anpassungen
                            self._copy_paragraphs_for_lek(doc, task['original_paragraphs'])
                        else:
                            # Letzter Fallback: Erstelle aus content Array
                            content_list = task.get('content', [])
                            for content_line in content_list:
                                if content_line.strip():
                                    doc.add_paragraph(content_line)
                    except Exception as e:
                        # Notfall-Fallback: Verwende original_paragraphs wenn elements-Kopierung fehlschlägt
                        print(f"Warnung: Element-Kopierung fehlgeschlagen für Aufgabe {i}: {e}")
                        if task.get('original_paragraphs'):
                            self._copy_paragraphs_for_lek(doc, task['original_paragraphs'])
                        else:
                            # Allerletzter Fallback
                            content_list = task.get('content', [])
                            for content_line in content_list:
                                if content_line.strip():
                                    doc.add_paragraph(content_line)
                    
                    # Metadaten als Kommentar (optional)
                    if task.get('difficulty') or task.get('keywords'):
                        meta_info = []
                        if task.get('difficulty'):
                            meta_info.append(f"Schwierigkeit: {task['difficulty']}")
                        if task.get('keywords'):
                            meta_info.append(f"Schlüsselwörter: {', '.join(task['keywords'][:5])}")
                        
                        meta_paragraph = doc.add_paragraph()
                        meta_paragraph.style = normal_style
                        meta_run = meta_paragraph.add_run(' | '.join(meta_info))
                        meta_run.italic = True
                    
                    # Trennlinie zwischen Aufgaben
                    if i < len(tasks):
                        separator_paragraph = doc.add_paragraph("─" * 50)
                        separator_paragraph.style = normal_style
            
            # Dokument speichern
            doc.save(output_path)
            
        except Exception as e:
            raise Exception(f"Fehler beim Erstellen der Word-Datei: {str(e)}")
    
    def _add_page_numbering_from_page_2(self, doc):
        """
        Fügt Seitennummerierung ab Seite 2 hinzu (beginnend mit 1)
        Format: "Seite/Gesamtseiten" rechtsbündig in der Fußzeile
        
        Args:
            doc: Word-Dokument
        """
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement
        
        # Erste Sektion (Deckblatt) - keine Seitennummern
        first_section = doc.sections[0]
        first_section.different_first_page_header_footer = True
        
        # Neue Sektion für Inhalt ab Seite 2
        if len(doc.sections) == 1:
            # Falls nur eine Sektion existiert, füge eine neue hinzu
            new_section = doc.add_section()
        else:
            new_section = doc.sections[-1]
        
        # Fußzeile für die neue Sektion konfigurieren
        footer = new_section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Seitennummerierung hinzufügen (Feld-Code für PAGE und NUMPAGES)
        run = footer_paragraph.add_run()
        
        # XML für Seitennummerierung erstellen
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE \\* MERGEFORMAT'
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # " / " Trenner
        footer_paragraph.add_run(' / ')
        
        # Gesamtseitenzahl
        run2 = footer_paragraph.add_run()
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES \\* MERGEFORMAT'
        run2._r.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar4)
        
        # Formatierung der Fußzeile
        for run in footer_paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Aptos'
    
    def get_document_info(self, file_path):
        """
        Gibt grundlegende Informationen über ein Word-Dokument zurück
        
        Args:
            file_path (str): Pfad zur Word-Datei
            
        Returns:
            dict: Dokumentinformationen
        """
        try:
            doc = Document(file_path)
            
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Geschätzte Aufgabenanzahl basierend auf erkannten Mustern
            task_count = 0
            for paragraph in doc.paragraphs:
                if self._is_task_start(paragraph.text.strip()):
                    task_count += 1
            
            return {
                'paragraph_count': paragraph_count,
                'estimated_task_count': task_count,
                'file_size': f"{len(doc.paragraphs)} Absätze"
            }
            
        except Exception as e:
            raise Exception(f"Fehler beim Analysieren der Datei: {str(e)}")
    
    def _extract_table_text(self, table_element, doc):
        """
        Extrahiert Text aus einer Tabelle für Keyword-Analyse
        
        Args:
            table_element: XML-Element der Tabelle
            doc: Word-Dokument-Objekt
            
        Returns:
            str: Extrahierter Tabellentext
        """
        try:
            # Finde entsprechendes Table-Objekt
            for table in doc.tables:
                if table._element == table_element:
                    text_parts = []
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                text_parts.append(cell_text)
                    return ' '.join(text_parts)
            return ""
        except Exception:
            return ""