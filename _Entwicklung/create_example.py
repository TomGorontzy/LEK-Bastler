"""
Beispielskript zum Erstellen einer Test-Word-Datei mit Aufgaben
"""

from docx import Document
import os

def create_sample_document():
    """Erstellt eine Beispiel-Word-Datei mit verschiedenen Aufgaben"""
    
    doc = Document()
    
    # Titel
    doc.add_heading('Beispiel-Aufgabensammlung', 0)
    doc.add_paragraph('Diese Datei enthält verschiedene Beispielaufgaben zum Testen des LEK-Bastlers.')
    
    # Aufgabe 1 - Mathematik, Leicht
    doc.add_heading('Aufgabe 1: Einfache Addition', 1)
    doc.add_paragraph('Berechnen Sie die folgenden Additionen:')
    doc.add_paragraph('a) 15 + 27 = ?')
    doc.add_paragraph('b) 43 + 56 = ?')
    doc.add_paragraph('c) 89 + 12 = ?')
    doc.add_paragraph('Schwierigkeit: Leicht')
    doc.add_paragraph('Schlüsselwörter: Mathematik, Addition, Grundrechenarten')
    
    # Aufgabe 2 - Deutsch, Mittel
    doc.add_heading('Aufgabe 2: Textanalyse', 1)
    doc.add_paragraph('Analysieren Sie den folgenden Text und beantworten Sie die Fragen:')
    doc.add_paragraph('"Der Wind heulte durch die kahlen Bäume des Waldes. Maria zog ihren Mantel enger um sich und beschleunigte ihre Schritte. Die Dunkelheit schien sie zu verfolgen."')
    doc.add_paragraph('Fragen:')
    doc.add_paragraph('1. Welche Stimmung wird in diesem Text erzeugt?')
    doc.add_paragraph('2. Welche sprachlichen Mittel werden verwendet?')
    doc.add_paragraph('3. Beschreiben Sie Marias Gemütszustand.')
    doc.add_paragraph('Schwierigkeit: Mittel')
    doc.add_paragraph('Schlüsselwörter: Deutsch, Textanalyse, Literatur, Interpretation')
    
    # Aufgabe 3 - Physik, Schwer
    doc.add_heading('Aufgabe 3: Mechanik Problem', 1)
    doc.add_paragraph('Ein Auto fährt mit konstanter Geschwindigkeit von 60 km/h auf einer geraden Strecke. Plötzlich bremst der Fahrer mit einer konstanten Verzögerung von 4 m/s².')
    doc.add_paragraph('Berechnen Sie:')
    doc.add_paragraph('a) Den Bremsweg bis zum Stillstand')
    doc.add_paragraph('b) Die Bremszeit')
    doc.add_paragraph('c) Die zurückgelegte Strecke in den ersten 2 Sekunden des Bremsvorgangs')
    doc.add_paragraph('Hinweis: Berücksichtigen Sie die Umrechnung von km/h in m/s.')
    doc.add_paragraph('Schwierigkeit: Schwer')
    doc.add_paragraph('Schlüsselwörter: Physik, Mechanik, Kinematik, Bremsvorgang')
    
    # Aufgabe 4 - Geschichte, Leicht
    doc.add_heading('Aufgabe 4: Historische Daten', 1)
    doc.add_paragraph('Ordnen Sie die folgenden historischen Ereignisse chronologisch:')
    doc.add_paragraph('- Fall der Berliner Mauer')
    doc.add_paragraph('- Gründung der Bundesrepublik Deutschland') 
    doc.add_paragraph('- Erste Mondlandung')
    doc.add_paragraph('- Beginn des Ersten Weltkriegs')
    doc.add_paragraph('- Entdeckung Amerikas durch Kolumbus')
    doc.add_paragraph('Schwierigkeit: Leicht')
    doc.add_paragraph('Schlüsselwörter: Geschichte, Chronologie, Deutschland, Weltgeschichte')
    
    # Aufgabe 5 - Informatik, Mittel  
    doc.add_heading('Aufgabe 5: Programmierung Basics', 1)
    doc.add_paragraph('Schreiben Sie ein Programm in Python, das:')
    doc.add_paragraph('1. Eine Liste von Zahlen einliest')
    doc.add_paragraph('2. Den Durchschnitt berechnet')
    doc.add_paragraph('3. Die größte und kleinste Zahl findet')
    doc.add_paragraph('4. Das Ergebnis formatiert ausgibt')
    doc.add_paragraph('Beispiel-Eingabe: [15, 23, 8, 42, 16, 4, 28]')
    doc.add_paragraph('Schwierigkeit: Mittel')
    doc.add_paragraph('Schlüsselwörter: Programmierung, Python, Algorithmus, Listen')
    
    # Aufgabe 6 - Chemie, Schwer
    doc.add_heading('Aufgabe 6: Chemische Reaktion', 1)
    doc.add_paragraph('Bei der Verbrennung von Methan (CH₄) entsteht Kohlendioxid und Wasser.')
    doc.add_paragraph('Aufgaben:')
    doc.add_paragraph('a) Stellen Sie die ausgeglichene Reaktionsgleichung auf')
    doc.add_paragraph('b) Berechnen Sie, wie viel CO₂ bei der Verbrennung von 100g Methan entsteht')
    doc.add_paragraph('c) Welches Volumen nimmt dieses CO₂ bei Standardbedingungen ein?')
    doc.add_paragraph('d) Diskutieren Sie die Umweltauswirkungen dieser Reaktion')
    doc.add_paragraph('Schwierigkeit: Schwer')
    doc.add_paragraph('Schlüsselwörter: Chemie, Verbrennung, Stöchiometrie, Umwelt')
    
    # Speichern
    output_path = os.path.join('examples', 'beispiel_aufgaben.docx')
    doc.save(output_path)
    print(f"Beispiel-Dokument wurde erstellt: {output_path}")

if __name__ == "__main__":
    create_sample_document()