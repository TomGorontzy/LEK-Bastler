"""
Erstellt eine Test-Word-Datei mit Tabellen und Listen für die Strukturerhaltung
"""

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Neues Dokument erstellen
doc = Document()

# Titel
title = doc.add_heading('Aufgabensammlung mit komplexen Strukturen', 0)

# Aufgabe 1 mit Tabelle
task1_title = doc.add_heading('Aufgabe 1: Kostenvergleich', 1)

doc.add_paragraph('Vergleichen Sie die Kosten der verschiedenen Materialien anhand der folgenden Tabelle:')

# Tabelle erstellen
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

# Header-Zeile
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Material'
hdr_cells[1].text = 'Kosten pro m²'
hdr_cells[2].text = 'Verfügbarkeit'

# Datenzeilen
row1_cells = table.rows[1].cells
row1_cells[0].text = 'Stahl'
row1_cells[1].text = '45,00 €'
row1_cells[2].text = 'Sofort lieferbar'

row2_cells = table.rows[2].cells
row2_cells[0].text = 'Aluminium'
row2_cells[1].text = '32,50 €'
row2_cells[2].text = '2 Wochen'

row3_cells = table.rows[3].cells
row3_cells[0].text = 'Kunststoff'
row3_cells[1].text = '18,75 €'
row3_cells[2].text = 'Sofort lieferbar'

doc.add_paragraph('Berechnen Sie die Gesamtkosten für ein Projekt mit 150 m² Materialbedarf.')

# Aufgabe 2 mit Aufzählung
task2_title = doc.add_heading('Aufgabe 2: Koordinationsaufgaben', 1)

doc.add_paragraph('Bei der Projektkoordination sind folgende Aspekte zu berücksichtigen:')

# Aufzählung erstellen
doc.add_paragraph('• Zeitplanung und Terminüberwachung', style='List Bullet')
doc.add_paragraph('• Ressourcenzuteilung und -optimierung', style='List Bullet')
doc.add_paragraph('• Kommunikation zwischen den Beteiligten', style='List Bullet')
doc.add_paragraph('• Qualitätskontrolle und Dokumentation', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Zusätzlich sind folgende Schritte in der angegebenen Reihenfolge auszuführen:')

# Nummerierte Liste
doc.add_paragraph('1. Stakeholder-Analyse durchführen', style='List Number')
doc.add_paragraph('2. Projektstruktur definieren', style='List Number')
doc.add_paragraph('3. Meilensteine festlegen', style='List Number')
doc.add_paragraph('4. Risikobewertung vornehmen', style='List Number')

# Aufgabe 3 mit gemischten Strukturen
task3_title = doc.add_heading('Aufgabe 3: Komplexe Strukturen', 1)

doc.add_paragraph('Diese Aufgabe kombiniert verschiedene Strukturelemente:')

# Verschachtelte Liste
doc.add_paragraph('A. Planungsphase')
doc.add_paragraph('   1. Bedarfsanalyse', style='List Number')
doc.add_paragraph('   2. Machbarkeitsstudie', style='List Number')
doc.add_paragraph('B. Umsetzungsphase')
doc.add_paragraph('   1. Prototypentwicklung', style='List Number')
doc.add_paragraph('   2. Testverfahren', style='List Number')

doc.add_paragraph('')
doc.add_paragraph('Priorisierung der Aufgaben:')

# Kleine Tabelle
small_table = doc.add_table(rows=3, cols=2)
small_table.style = 'Light List Accent 1'

priority_cells = small_table.rows[0].cells
priority_cells[0].text = 'Priorität'
priority_cells[1].text = 'Aufgabe'

high_cells = small_table.rows[1].cells
high_cells[0].text = 'Hoch'
high_cells[1].text = 'Sicherheitsanalyse'

medium_cells = small_table.rows[2].cells
medium_cells[0].text = 'Mittel'
medium_cells[1].text = 'Kostenschätzung'

doc.add_paragraph('')
para = doc.add_paragraph()
run = para.add_run('Wichtiger Hinweis: ')
run.bold = True
run = para.add_run('Alle Strukturen sollten ')
run = para.add_run('vollständig erhalten')
run.italic = True
run = para.add_run(' bleiben.')

# Dokument speichern
doc.save('d:/OneDrive/Git-Projekte/LEK-Bastler/examples/Aufgaben_mit_Strukturen.docx')
print("Test-Datei mit komplexen Strukturen wurde erstellt: examples/Aufgaben_mit_Strukturen.docx")