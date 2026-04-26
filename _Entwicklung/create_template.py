"""
Erstellt eine Beispiel-Vorlage für Auftragssteuerung und -koordination
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template():
    """Erstellt eine Beispiel-Vorlage für LEK Auftragssteuerung"""
    
    doc = Document()
    
    # Seitenränder einstellen
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5 / 2.54)
        section.bottom_margin = Inches(1.5 / 2.54)
        section.left_margin = Inches(2.5 / 2.54)
        section.right_margin = Inches(1.5 / 2.54)
    
    # === SEITE 1: DECKBLATT ===
    # Haupttitel
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("LERNERFOLGSKONTROLLE")
    title_run.font.name = 'Aptos'
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 164, 153)
    title_run.bold = True
    
    # Leerzeilen
    for _ in range(4):
        doc.add_paragraph()
    
    # Fachbereich
    subject_p = doc.add_paragraph()
    subject_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_run = subject_p.add_run("Auftragssteuerung und -koordination")
    subject_run.font.name = 'Aptos'
    subject_run.font.size = Pt(20)
    subject_run.font.color.rgb = RGBColor(0, 120, 112)
    subject_run.bold = True
    
    # Leerzeilen
    for _ in range(8):
        doc.add_paragraph()
    
    # Informationsfelder
    info_fields = [
        "Name: _______________________________",
        "",
        "Datum: ______________________________",
        "",
        "Bearbeitungszeit: ____________________",
        "",
        "Erreichte Punkte: ______ von ______ Punkten"
    ]
    
    for field in info_fields:
        if field.strip():  # Nur wenn nicht leer
            field_p = doc.add_paragraph(field)
            field_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            field_run = field_p.runs[0]
            field_run.font.name = 'Aptos'
            field_run.font.size = Pt(12)
        else:
            doc.add_paragraph()  # Leerer Absatz
    
    # === SEITE 2: HINWEISE/ANWEISUNGEN ===
    doc.add_page_break()
    
    # Überschrift Seite 2
    instructions_title = doc.add_heading("Hinweise zur Bearbeitung", 1)
    instructions_title.style.font.color.rgb = RGBColor(0, 164, 153)
    
    # Bearbeitungshinweise
    instructions = [
        "• Lesen Sie alle Aufgaben sorgfältig durch.",
        "• Bearbeiten Sie die Aufgaben in beliebiger Reihenfolge.",
        "• Verwenden Sie für Ihre Antworten einen schwarzen oder blauen Stift.",
        "• Hilfsmittel: Taschenrechner, Formelsammlung (nach Absprache)",
        "• Bei Unklarheiten fragen Sie bitte nach.",
        "",
        "Viel Erfolg bei der Bearbeitung!"
    ]
    
    for instruction in instructions:
        if instruction:
            inst_p = doc.add_paragraph(instruction)
            inst_p.style.font.name = 'Aptos'
            inst_p.style.font.size = Pt(11)
        else:
            doc.add_paragraph()
    
    # Bewertungsschema
    doc.add_paragraph()
    doc.add_paragraph()
    
    grading_title = doc.add_heading("Bewertungsschema", 2)
    grading_title.style.font.color.rgb = RGBColor(0, 164, 153)
    
    grading_info = [
        "100% - 92%:  Note 1 (sehr gut)",
        " 91% - 81%:  Note 2 (gut)", 
        " 80% - 67%:  Note 3 (befriedigend)",
        " 66% - 50%:  Note 4 (ausreichend)",
        " 49% - 30%:  Note 5 (mangelhaft)",
        " 29% -  0%:  Note 6 (ungenügend)"
    ]
    
    for grade in grading_info:
        grade_p = doc.add_paragraph(grade)
        grade_p.style.font.name = 'Aptos'
        grade_p.style.font.size = Pt(11)
    
    # Speichern
    output_path = os.path.join("Vorlagen", "LEK-Vorlage_Auftragssteuerung-Koordination.docx")
    doc.save(output_path)
    print(f"Vorlage erstellt: {output_path}")

if __name__ == "__main__":
    create_template()